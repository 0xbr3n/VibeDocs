"""
Purchase Order (project-scoped) and Request For Information (report-scoped)
document storage.

We don't add new tables for these — Project.details and Report.details are
JSON columns that already hold engagement-specific metadata. We append to a
`po_documents` / `rfi_documents` list on each row, where each entry is:

    {
      "id":              "<uuid-hex>",
      "original_name":   "Acme PO 12345.pdf",
      "stored_path":     "/data/uploads/po/12/<uuid>.pdf",
      "content_type":    "application/pdf",
      "size_bytes":      83214,
      "uploaded_by_id":  17,
      "uploaded_at":     "2026-05-13T08:42:11Z",
      "extracted_text":  "...best-effort plain text preview..."
    }

Endpoints (all auth-required; PO requires project visibility, RFI requires
report `view` access for read and `edit` for write):

    POST    /api/projects/{pid}/po
    GET     /api/projects/{pid}/po
    GET     /api/projects/{pid}/po/{doc_id}/download
    DELETE  /api/projects/{pid}/po/{doc_id}

    POST    /api/reports/{rid}/rfi
    GET     /api/reports/{rid}/rfi
    GET     /api/reports/{rid}/rfi/{doc_id}/download
    DELETE  /api/reports/{rid}/rfi/{doc_id}
"""
from __future__ import annotations
import shutil
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy.orm.attributes import flag_modified

from ..database import get_db
from ..models import (
    Project, Report, User, Role, AccessLevel, AuditLog,
)
from ..auth import get_current_user
from ..config import settings
from .permissions import effective_access, require_access
from ..services import doc_extract


router = APIRouter(tags=["engagement-docs"])

# Files we accept. Anything else gets 400.
ALLOWED_EXT = {".pdf", ".xlsx", ".xls", ".xlsm", ".csv", ".docx", ".doc"}
MAX_BYTES = 25 * 1024 * 1024   # 25 MB — generous for client RFIs


# ============================================================
# Helpers
# ============================================================

def _new_doc_id() -> str:
    return uuid.uuid4().hex


def _save_upload(file: UploadFile, subdir: Path) -> tuple[Path, int]:
    """Stream the upload to disk, enforcing the size limit. All file types accepted."""
    subdir.mkdir(parents=True, exist_ok=True)
    suffix = Path(file.filename or "").suffix.lower()
    out = subdir / f"{_new_doc_id()}{suffix}"
    written = 0
    with out.open("wb") as fh:
        while True:
            chunk = file.file.read(1024 * 64)
            if not chunk:
                break
            written += len(chunk)
            if written > MAX_BYTES:
                fh.close()
                out.unlink(missing_ok=True)
                raise HTTPException(413, f"File too large (limit {MAX_BYTES // (1024*1024)} MB).")
            fh.write(chunk)
    return out, written


def _entry_from_upload(file: UploadFile, out: Path, size: int, user: User) -> dict:
    return {
        "id": out.stem,
        "original_name": (file.filename or out.name)[:255],
        "stored_path": str(out),
        "content_type": file.content_type or "application/octet-stream",
        "size_bytes": size,
        "uploaded_by_id": user.id,
        "uploaded_at": datetime.utcnow().isoformat() + "Z",
        "extracted_text": doc_extract.extract_text(out),
    }


def _strip_secrets(entries: list[dict]) -> list[dict]:
    """Strip the full path before sending to the client — they only need the
    id (for downloads) and the metadata."""
    out = []
    for e in entries or []:
        out.append({
            "id": e.get("id"),
            "original_name": e.get("original_name"),
            "content_type": e.get("content_type"),
            "size_bytes": e.get("size_bytes"),
            "uploaded_by_id": e.get("uploaded_by_id"),
            "uploaded_at": e.get("uploaded_at"),
            "has_text_preview": bool(e.get("extracted_text")),
            "category": e.get("category", "po"),
        })
    return out


# ============================================================
# PO (project-level)
# ============================================================

def _project_or_404(db: Session, pid: int) -> Project:
    p = db.get(Project, pid)
    if not p:
        raise HTTPException(404, "Project not found")
    return p


def _project_can_write(db: Session, project: Project, user: User) -> bool:
    """Project edits are gated by lead/admin/senior. (We don't yet have a
    full project-level ACL; lead/admin is the existing convention.)"""
    if user.role in (Role.admin, Role.senior):
        return True
    return project.lead_id == user.id


# Shared by both PO and RFI preview endpoints.
_PREVIEWABLE_AS_PDF = {".xlsx", ".xls", ".xlsm", ".doc", ".docx", ".odt", ".ods"}
_PREVIEWABLE_AS_HTML = {".csv", ".xlsx", ".xls", ".xlsm", ".docx"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
_IMAGE_MIME = {
    ".png":  "image/png",  ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg", ".gif":  "image/gif",
    ".webp": "image/webp",
}


# Inline CSS for HTML fallback previews — themed neutral so the iframe
# preview reads under any theme without inheriting the host page's
# stylesheet. Kept compact.
_HTML_PREVIEW_HEAD = (
    "<style>"
    "body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;"
    "margin:14px;color:#111;background:#fff;font-size:13px;line-height:1.5}"
    "h1,h2,h3,h4,h5{margin:.8em 0 .3em;color:#0f172a}"
    "h1{font-size:1.5em}h2{font-size:1.25em}h3{font-size:1.1em}"
    ".sheet-title{font-size:13px;font-weight:600;margin:14px 0 6px;color:#374151}"
    "table{border-collapse:collapse;font-size:12px;margin-bottom:18px;max-width:100%}"
    "th,td{border:1px solid #d1d5db;padding:5px 8px;text-align:left;vertical-align:top;"
    "white-space:pre-wrap;word-break:break-word;max-width:380px;overflow:hidden}"
    "th{background:#f3f4f6;font-weight:600}"
    "tr:nth-child(even) td{background:#fafafa}"
    "p{margin:.4em 0}ul,ol{margin:.4em 0 .4em 22px;padding:0}"
    "pre,code{font-family:ui-monospace,Consolas,monospace;font-size:12px}"
    "</style>"
)


def _html_preview_csv(src: Path) -> str:
    """Render a CSV file as an HTML table. Caps rows + columns so a
    multi-megabyte tracker doesn't blow up the iframe."""
    import csv
    MAX_ROWS, MAX_COLS = 500, 40
    rows: list[list[str]] = []
    with src.open("r", encoding="utf-8-sig", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i >= MAX_ROWS:
                rows.append([f"… ({MAX_ROWS}+ rows truncated)"])
                break
            rows.append(row[:MAX_COLS])
    return _rows_to_table_html(src.name, [(None, rows)], truncated_cols=any(len(r) > MAX_COLS for r in rows))


def _html_preview_xlsx(src: Path) -> str:
    """Render an XLSX / XLSM workbook as a series of HTML tables, one per
    sheet. Uses openpyxl (already a dependency for the tracker import).
    Caps each sheet to keep the preview snappy.
    """
    from openpyxl import load_workbook
    MAX_ROWS, MAX_COLS = 300, 40
    wb = load_workbook(src, read_only=True, data_only=True)
    sheets: list[tuple[str, list[list[str]]]] = []
    for ws in wb.worksheets:
        rows: list[list[str]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= MAX_ROWS:
                rows.append([f"… ({MAX_ROWS}+ rows truncated)"])
                break
            cells = [("" if c is None else str(c)) for c in row[:MAX_COLS]]
            rows.append(cells)
        sheets.append((ws.title, rows))
    return _rows_to_table_html(src.name, sheets)


def _html_preview_docx(src: Path) -> str:
    """Render a DOCX as best-effort HTML using python-docx. Heading levels
    are mapped to <h1>-<h6>, paragraphs to <p>, tables to <table>. Lists
    use Word's `numId` heuristic — if a paragraph has list formatting we
    wrap it in <ul>. This is intentionally simple; the real DOCX viewer
    is the LibreOffice→PDF path, this is a fallback when LibreOffice
    isn't available or the file fails conversion."""
    from docx import Document
    from html import escape
    doc = Document(str(src))
    out: list[str] = []

    def _para_html(p) -> str:
        # Heading?
        style = (p.style.name if p.style else "") or ""
        text = escape(p.text or "")
        if not (p.text or "").strip():
            return "<p>&nbsp;</p>"
        if style.startswith("Heading "):
            try:
                lvl = int(style.split(" ", 1)[1])
                lvl = max(1, min(lvl, 6))
            except (ValueError, IndexError):
                lvl = 3
            return f"<h{lvl}>{text}</h{lvl}>"
        return f"<p>{text}</p>"

    for el in doc.element.body.iterchildren():
        tag = el.tag.split("}", 1)[-1]
        if tag == "p":
            # Find the matching python-docx Paragraph wrapper for styled output.
            for p in doc.paragraphs:
                if p._element is el:
                    out.append(_para_html(p))
                    break
        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._element is el:
                    out.append("<table>")
                    for row in tbl.rows:
                        out.append("<tr>")
                        for cell in row.cells:
                            out.append(f"<td>{escape(cell.text)}</td>")
                        out.append("</tr>")
                    out.append("</table>")
                    break
    body = "\n".join(out) or "<p class='muted'>(empty document)</p>"
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        + _HTML_PREVIEW_HEAD
        + f"<title>{escape(src.name)}</title></head><body>"
        + f"<h2>{escape(src.name)}</h2>"
        + body
        + "</body></html>"
    )


def _rows_to_table_html(filename: str, sheets: list[tuple[Optional[str], list[list[str]]]],
                       truncated_cols: bool = False) -> str:
    """Wrap one or more (sheet_name, rows) chunks into a single HTML doc.
    `sheet_name=None` collapses the per-sheet header (CSV has no sheets)."""
    from html import escape
    body_parts: list[str] = [f"<h2>{escape(filename)}</h2>"]
    for sheet_name, rows in sheets:
        if sheet_name:
            body_parts.append(f"<div class='sheet-title'>{escape(sheet_name)}</div>")
        if not rows:
            body_parts.append("<p class='muted'>(empty sheet)</p>")
            continue
        body_parts.append("<table>")
        for i, row in enumerate(rows):
            cell_tag = "th" if i == 0 else "td"
            body_parts.append("<tr>")
            for cell in row:
                body_parts.append(f"<{cell_tag}>{escape(str(cell))}</{cell_tag}>")
            body_parts.append("</tr>")
        body_parts.append("</table>")
    if truncated_cols:
        body_parts.append("<p class='muted'><em>Some rows were wider than the preview limit and have been trimmed.</em></p>")
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        + _HTML_PREVIEW_HEAD
        + f"<title>{escape(filename)}</title></head><body>"
        + "\n".join(body_parts)
        + "</body></html>"
    )


def _html_preview(src: Path):
    """Best-effort browser-renderable HTML preview for files that
    LibreOffice can't / shouldn't convert (CSV always, XLSX/DOCX as a
    fallback path)."""
    from fastapi.responses import HTMLResponse
    suffix = src.suffix.lower()
    try:
        if suffix == ".csv":
            html = _html_preview_csv(src)
        elif suffix in (".xlsx", ".xls", ".xlsm"):
            html = _html_preview_xlsx(src)
        elif suffix == ".docx":
            html = _html_preview_docx(src)
        else:
            return None
    except Exception as e:                          # pragma: no cover
        # Don't 500 the page — let the caller fall through to the
        # "download instead" message.
        import logging
        logging.getLogger(__name__).warning(
            "HTML-fallback preview failed for %s: %s", src.name, e
        )
        return None
    return HTMLResponse(html)


def _preview_file(src: Path):
    """Return a Response the browser can render inline.

    - .pdf  → pass-through with Content-Disposition inline.
    - .xlsx / .xls / .doc / .docx / .odt / .ods → LibreOffice → PDF
      (result cached next to the source as `<stem>.preview.pdf` and
      invalidated when the source mtime moves). On LibreOffice failure
      we fall back to an HTML preview built from openpyxl /
      python-docx, so the user sees the data instead of an opaque
      "LibreOffice did not produce a PDF" error.
    - .csv → always rendered as an HTML table (no LibreOffice round-trip).
    - .png / .jpg / .jpeg / .gif / .webp → pass-through as image.
    - Anything else → 415 so the UI can fall back to download.
    """
    suffix = src.suffix.lower()
    if suffix == ".pdf":
        return FileResponse(
            src, media_type="application/pdf", filename=src.name,
            headers={"Content-Disposition": f'inline; filename="{src.name}"'},
        )
    if suffix in _IMAGE_EXTS:
        return FileResponse(
            src, media_type=_IMAGE_MIME[suffix], filename=src.name,
            headers={"Content-Disposition": f'inline; filename="{src.name}"'},
        )
    # CSV always goes the HTML route — LibreOffice would fight over
    # delimiter / locale anyway, and the HTML table is easier to read.
    if suffix == ".csv":
        resp = _html_preview(src)
        if resp is not None:
            return resp
        raise HTTPException(500, "Could not render CSV preview.")
    if suffix in _PREVIEWABLE_AS_PDF:
        cached_pdf = src.with_suffix(".preview.pdf")
        if not cached_pdf.exists() or cached_pdf.stat().st_mtime < src.stat().st_mtime:
            from ..services.docx_generator import convert_to_pdf
            produced: Optional[Path] = None
            try:
                produced = convert_to_pdf(src, src.parent)
            except RuntimeError:
                # LibreOffice may exit 0 but write the PDF under an
                # unexpected name (filename normalisation, etc). Glob for
                # any fresh .pdf in the same dir as a fallback before
                # giving up.
                fresh_pdfs = [
                    p for p in src.parent.glob("*.pdf")
                    if p != cached_pdf and p.stat().st_mtime >= src.stat().st_mtime
                ]
                if fresh_pdfs:
                    produced = max(fresh_pdfs, key=lambda p: p.stat().st_mtime)
            except Exception:
                produced = None
            if produced is None or not produced.exists():
                # Fall back to the HTML preview path if we can produce
                # one for this file type. Beats a 500 because at least
                # the user can read the data.
                if suffix in _PREVIEWABLE_AS_HTML:
                    resp = _html_preview(src)
                    if resp is not None:
                        return resp
                raise HTTPException(
                    500,
                    "Preview unavailable: LibreOffice did not produce a PDF "
                    "for this file, and the in-browser fallback could not "
                    "render it either. Try downloading it instead.",
                )
            try:
                produced.replace(cached_pdf)
            except Exception:
                return FileResponse(
                    produced, media_type="application/pdf", filename=produced.name,
                    headers={"Content-Disposition":
                             f'inline; filename="{produced.name}"'},
                )
        return FileResponse(
            cached_pdf, media_type="application/pdf", filename=cached_pdf.name,
            headers={"Content-Disposition": f'inline; filename="{cached_pdf.name}"'},
        )
    raise HTTPException(
        415,
        f"Preview not supported for {suffix or 'this file type'}. "
        "Use Download instead.",
    )


@router.post("/api/projects/{pid}/po")
def po_upload(pid: int, file: UploadFile = File(...),
              category: str = Form("po"),
              db: Session = Depends(get_db),
              user: User = Depends(get_current_user)):
    project = _project_or_404(db, pid)
    if not _project_can_write(db, project, user):
        raise HTTPException(403, "Only project lead or senior+ can upload POs")
    subdir = Path(settings.UPLOAD_DIR) / "po" / str(pid)
    out, size = _save_upload(file, subdir)
    entry = _entry_from_upload(file, out, size, user)
    entry["category"] = category if category in ("po", "logs") else "po"
    details = dict(project.details or {})
    docs = list(details.get("po_documents") or [])
    docs.append(entry)
    details["po_documents"] = docs
    project.details = details
    flag_modified(project, "details")
    db.add(AuditLog(actor_id=user.id, action="project.po.upload",
                    object_type="project", object_id=pid,
                    detail={"doc_id": entry["id"], "filename": entry["original_name"]}))
    db.commit()
    return _strip_secrets([entry])[0]


@router.get("/api/projects/{pid}/po")
def po_list(pid: int, db: Session = Depends(get_db),
            user: User = Depends(get_current_user)):
    from .permissions import require_project_visibility
    project = _project_or_404(db, pid)
    require_project_visibility(db, user, project)
    return {"project_id": pid,
            "items": _strip_secrets((project.details or {}).get("po_documents") or [])}


def _find_doc(entries: list[dict], doc_id: str) -> Optional[dict]:
    for e in entries or []:
        if e.get("id") == doc_id:
            return e
    return None


@router.get("/api/projects/{pid}/po/{doc_id}/download")
def po_download(pid: int, doc_id: str, db: Session = Depends(get_db),
                user: User = Depends(get_current_user)):
    from .permissions import require_project_visibility
    project = _project_or_404(db, pid)
    require_project_visibility(db, user, project)
    entry = _find_doc((project.details or {}).get("po_documents"), doc_id)
    if not entry:
        raise HTTPException(404, "PO document not found")
    path = Path(entry["stored_path"])
    if not path.exists():
        raise HTTPException(410, "File missing on disk")
    return FileResponse(path, filename=entry.get("original_name") or path.name,
                        media_type=entry.get("content_type") or "application/octet-stream")


@router.get("/api/projects/{pid}/po/{doc_id}/preview")
def po_preview(pid: int, doc_id: str, db: Session = Depends(get_db),
               user: User = Depends(get_current_user)):
    """Inline-renderable preview of a Purchase Order document. PDFs are
    served as-is; Excel files go through LibreOffice and the resulting
    PDF is cached next to the source. Mirrors the RFI preview endpoint so
    the project page can offer a Preview button alongside Download —
    consultants don't have to download every PO just to skim it."""
    from .permissions import require_project_visibility
    project = _project_or_404(db, pid)
    require_project_visibility(db, user, project)
    entry = _find_doc((project.details or {}).get("po_documents"), doc_id)
    if not entry:
        raise HTTPException(404, "PO document not found")
    src = Path(entry["stored_path"])
    if not src.exists():
        raise HTTPException(410, "File missing on disk")
    return _preview_file(src)


@router.delete("/api/projects/{pid}/po/{doc_id}")
def po_delete(pid: int, doc_id: str, db: Session = Depends(get_db),
              user: User = Depends(get_current_user)):
    project = _project_or_404(db, pid)
    if not _project_can_write(db, project, user):
        raise HTTPException(403, "Only project lead or senior+ can delete POs")
    details = dict(project.details or {})
    docs = list(details.get("po_documents") or [])
    keep, dropped = [], None
    for e in docs:
        if e.get("id") == doc_id:
            dropped = e
        else:
            keep.append(e)
    if not dropped:
        raise HTTPException(404, "PO document not found")
    details["po_documents"] = keep
    project.details = details
    flag_modified(project, "details")
    db.add(AuditLog(actor_id=user.id, action="project.po.delete",
                    object_type="project", object_id=pid,
                    detail={"doc_id": doc_id, "filename": dropped.get("original_name")}))
    db.commit()
    # Best-effort delete on disk; if it fails the audit row still preserves intent
    try: Path(dropped["stored_path"]).unlink(missing_ok=True)
    except Exception: pass
    return {"ok": True, "deleted_id": doc_id}


# ============================================================
# RFI (report-level)
# ============================================================

def _report_or_404(db: Session, rid: int) -> Report:
    r = db.get(Report, rid)
    if not r:
        raise HTTPException(404, "Report not found")
    return r


@router.post("/api/reports/{rid}/rfi")
def rfi_upload(rid: int, file: UploadFile = File(...),
               db: Session = Depends(get_db),
               user: User = Depends(get_current_user)):
    report = _report_or_404(db, rid)
    require_access(db, user, report, need=AccessLevel.edit)
    subdir = Path(settings.UPLOAD_DIR) / "rfi" / str(rid)
    out, size = _save_upload(file, subdir)
    entry = _entry_from_upload(file, out, size, user)
    details = dict(report.details or {})
    docs = list(details.get("rfi_documents") or [])
    docs.append(entry)
    details["rfi_documents"] = docs
    report.details = details
    flag_modified(report, "details")
    db.add(AuditLog(actor_id=user.id, action="report.rfi.upload",
                    object_type="report", object_id=rid,
                    detail={"doc_id": entry["id"], "filename": entry["original_name"]}))
    db.commit()
    return _strip_secrets([entry])[0]


@router.get("/api/reports/{rid}/rfi")
def rfi_list(rid: int, db: Session = Depends(get_db),
             user: User = Depends(get_current_user)):
    report = _report_or_404(db, rid)
    require_access(db, user, report, need=AccessLevel.view)
    return {"report_id": rid,
            "items": _strip_secrets((report.details or {}).get("rfi_documents") or [])}


@router.get("/api/reports/{rid}/rfi/{doc_id}/download")
def rfi_download(rid: int, doc_id: str, db: Session = Depends(get_db),
                 user: User = Depends(get_current_user)):
    report = _report_or_404(db, rid)
    require_access(db, user, report, need=AccessLevel.view)
    entry = _find_doc((report.details or {}).get("rfi_documents"), doc_id)
    if not entry:
        raise HTTPException(404, "RFI document not found")
    path = Path(entry["stored_path"])
    if not path.exists():
        raise HTTPException(410, "File missing on disk")
    return FileResponse(path, filename=entry.get("original_name") or path.name,
                        media_type=entry.get("content_type") or "application/octet-stream")


@router.get("/api/reports/{rid}/rfi/{doc_id}/preview")
def rfi_preview(rid: int, doc_id: str, db: Session = Depends(get_db),
                user: User = Depends(get_current_user)):
    """Render the RFI for inline viewing.

    PDFs pass through as-is (the browser's native viewer handles them in
    an iframe). Excel files go through the same LibreOffice pipeline the
    DOCX preview uses — `convert_to_pdf` writes a PDF next to the source
    so subsequent views are instant. Anything else gets a 415 so the
    frontend can fall back to "download to view".

    The Content-Disposition is `inline` so the browser embeds rather than
    triggers a download.
    """
    report = _report_or_404(db, rid)
    require_access(db, user, report, need=AccessLevel.view)
    entry = _find_doc((report.details or {}).get("rfi_documents"), doc_id)
    if not entry:
        raise HTTPException(404, "RFI document not found")
    src = Path(entry["stored_path"])
    if not src.exists():
        raise HTTPException(410, "File missing on disk")

    return _preview_file(src)


@router.delete("/api/reports/{rid}/rfi/{doc_id}")
def rfi_delete(rid: int, doc_id: str, db: Session = Depends(get_db),
               user: User = Depends(get_current_user)):
    report = _report_or_404(db, rid)
    require_access(db, user, report, need=AccessLevel.edit)
    details = dict(report.details or {})
    docs = list(details.get("rfi_documents") or [])
    keep, dropped = [], None
    for e in docs:
        if e.get("id") == doc_id:
            dropped = e
        else:
            keep.append(e)
    if not dropped:
        raise HTTPException(404, "RFI document not found")
    details["rfi_documents"] = keep
    report.details = details
    flag_modified(report, "details")
    db.add(AuditLog(actor_id=user.id, action="report.rfi.delete",
                    object_type="report", object_id=rid,
                    detail={"doc_id": doc_id, "filename": dropped.get("original_name")}))
    db.commit()
    try: Path(dropped["stored_path"]).unlink(missing_ok=True)
    except Exception: pass
    return {"ok": True, "deleted_id": doc_id}
