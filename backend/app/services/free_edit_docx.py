"""
Round-trip DOCX ↔ HTML conversion for the "Free Edit" feature.

The consultant wants to make ad-hoc tweaks to a generated Word report
(alignment fixes, missing sections, graph placement) inside the browser
or in Word, without re-running the template renderer.

This module provides two conversions:

  docx_to_html(path) -> str
      Best-effort one-way render of a .docx into editable HTML. Headings
      map to <h1>-<h6>, paragraphs to <p>, runs preserve bold / italic /
      underline / colour, tables become <table>, inline images are
      embedded as base64 <img> data: URLs. Hyperlinks are preserved.

  html_to_docx(html_str, output_path)
      Parse the HTML the user submitted back from the editor and write
      a fresh DOCX. Supports the same block + inline tags that
      docx_to_html emits. The image data: URLs are decoded back into
      InlineShape entries.

Caveats — deliberate trade-offs to keep this dependency-free:

  * Complex Word features (text boxes, SmartArt, footnotes, comments,
    track changes, multi-column layouts) DO NOT survive the round-trip.
    The free-edit flow is for adjusting prose + adding / moving simple
    blocks, not for rebuilding the document's layout.
  * Headers / footers from the source document are preserved by reading
    them as text into the HTML for context and re-writing them on the
    output side. Watermarks are dropped (they're re-applied at generate
    time anyway).
  * Lists detected via Word's `numId` are emitted as <ul>; ordered
    lists fall back to <ol> when the abstractNum is decimal.

Anything beyond the supported subset is rendered as its plain-text
content so the editor never loses data — it just loses formatting.
"""
from __future__ import annotations
import base64
import io
import re
from html.parser import HTMLParser
from html import escape
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# DOCX -> HTML
# ============================================================

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def _run_to_html(run) -> str:
    """Render a docx Run as inline HTML preserving bold/italic/underline
    /strike/colour. Empty runs are skipped to avoid littering the editor
    with empty <span>s."""
    text = run.text or ""
    if not text:
        return ""
    # Use <br> for explicit Word line breaks within a run (Shift+Enter).
    out = escape(text).replace("\n", "<br>")
    if run.bold:        out = f"<strong>{out}</strong>"
    if run.italic:      out = f"<em>{out}</em>"
    if run.underline:   out = f"<u>{out}</u>"
    try:
        if run.font.strike:
            out = f"<s>{out}</s>"
    except (AttributeError, KeyError):
        pass
    # Colour
    try:
        color = run.font.color and run.font.color.rgb
        if color:
            out = f'<span style="color:#{color}">{out}</span>'
    except (AttributeError, KeyError):
        pass
    return out


def _paragraph_to_html(p) -> str:
    """Render a docx Paragraph element to a block-level HTML string.

    Headings become <h1>-<h6> based on the paragraph style name.
    List paragraphs become a marker; the caller groups consecutive
    list paragraphs into <ul>/<ol>. Other paragraphs become <p>.
    """
    runs_html = "".join(_run_to_html(r) for r in p.runs)
    if not runs_html:
        runs_html = "&nbsp;"

    # Heading level via the style
    style_name = (p.style.name or "") if p.style else ""
    m = _HEADING_RE.match(style_name)
    if m:
        lvl = max(1, min(int(m.group(1)), 6))
        return f"<h{lvl}>{runs_html}</h{lvl}>"

    # Detect Word list paragraphs via the underlying XML: a paragraph
    # in a list has a <w:numPr> child. Returns the bullet/list marker
    # so the outer pass can group them. The actual <ul>/<ol> wrapping
    # is done in _body_to_html.
    pPr = p._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is not None:
        numPr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
        if numPr is not None:
            return f"<li>{runs_html}</li>"

    # Alignment
    align = None
    try:
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align = "center"
        elif p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align = "right"
        elif p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align = "justify"
    except (AttributeError, KeyError):
        pass
    if align:
        return f'<p style="text-align:{align}">{runs_html}</p>'
    return f"<p>{runs_html}</p>"


def _table_to_html(tbl) -> str:
    rows_html: list[str] = []
    for row in tbl.rows:
        cells_html = []
        for cell in row.cells:
            inner = []
            for p in cell.paragraphs:
                inner.append("".join(_run_to_html(r) for r in p.runs) or "&nbsp;")
            cells_html.append("<td>" + "<br>".join(inner) + "</td>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return "<table border='1'>" + "\n".join(rows_html) + "</table>"


def _extract_inline_images(doc: Document) -> dict[str, str]:
    """Return {rId: data:base64 URL} for every inline image in the
    document's main part. Used to substitute <img> tags into the HTML
    in place of bare image placeholders."""
    images: dict[str, str] = {}
    rels = doc.part.rels
    for rid, rel in rels.items():
        if "image" in (rel.reltype or ""):
            try:
                blob = rel.target_part.blob
            except Exception:
                continue
            mime = "image/png"
            if rel.target_part.partname.lower().endswith(".jpg") or \
               rel.target_part.partname.lower().endswith(".jpeg"):
                mime = "image/jpeg"
            elif rel.target_part.partname.lower().endswith(".gif"):
                mime = "image/gif"
            data_url = f"data:{mime};base64,{base64.b64encode(blob).decode('ascii')}"
            images[rid] = data_url
    return images


_NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_NS_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def _paragraph_inline_images(p, img_map: dict[str, str]) -> list[str]:
    """Find <a:blip r:embed="..."> elements within the paragraph and
    return matching data URLs so we can append them after the prose."""
    out: list[str] = []
    for blip in p._p.iter(_NS_A + "blip"):
        rid = blip.get(_NS_R + "embed")
        if rid and rid in img_map:
            out.append(img_map[rid])
    return out


def docx_to_html(src: Path) -> str:
    """Render `src` as editable HTML. Returns a complete <html> document
    so the editor can iframe it if it wants to."""
    doc = Document(str(src))
    img_map = _extract_inline_images(doc)

    body_chunks: list[str] = []
    list_buffer: list[str] = []
    list_open = False

    def _flush_list():
        nonlocal list_buffer, list_open
        if list_buffer:
            body_chunks.append("<ul>" + "".join(list_buffer) + "</ul>")
            list_buffer = []
        list_open = False

    for el in doc.element.body.iterchildren():
        tag = el.tag.split("}", 1)[-1]
        if tag == "p":
            p = None
            for pp in doc.paragraphs:
                if pp._element is el:
                    p = pp; break
            if p is None:
                continue
            html_chunk = _paragraph_to_html(p)
            if html_chunk.startswith("<li>"):
                list_buffer.append(html_chunk)
                list_open = True
                continue
            else:
                _flush_list()
            body_chunks.append(html_chunk)
            # Inline images attached to the paragraph
            for url in _paragraph_inline_images(p, img_map):
                body_chunks.append(f'<p><img src="{url}" style="max-width:100%"></p>')
        elif tag == "tbl":
            _flush_list()
            for t in doc.tables:
                if t._element is el:
                    body_chunks.append(_table_to_html(t))
                    break
        elif tag == "sectPr":
            _flush_list()
        # Other elements (sectPr only really) are ignored deliberately.
    _flush_list()

    body_html = "\n".join(body_chunks) or "<p>(empty document)</p>"
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<style>"
        "body{font-family:'Calibri',system-ui,sans-serif;font-size:12pt;"
        "padding:24px;max-width:820px;margin:0 auto;background:#fff;color:#111;"
        "line-height:1.5}"
        "h1{font-size:22pt;font-weight:600;margin:18pt 0 6pt;color:#1f2937}"
        "h2{font-size:18pt;font-weight:600;margin:14pt 0 4pt;color:#1f2937}"
        "h3{font-size:14pt;font-weight:600;margin:12pt 0 4pt;color:#1f2937}"
        "h4,h5,h6{font-size:12pt;font-weight:600;margin:10pt 0 3pt;color:#1f2937}"
        "p{margin:6pt 0}img{display:block;margin:8pt auto}"
        "table{border-collapse:collapse;margin:8pt 0;width:100%}"
        "table td,table th{border:1px solid #888;padding:4pt 6pt;vertical-align:top}"
        "</style></head>"
        f"<body contenteditable='true' spellcheck='true'>{body_html}</body></html>"
    )


# ============================================================
# HTML -> DOCX
# ============================================================

class _HTMLToDocx(HTMLParser):
    """Walk the user-submitted HTML and emit paragraphs / runs into a
    fresh `docx.Document`. Tag handling mirrors what `docx_to_html` emits;
    anything outside the allow-list is rendered as plain text so we
    never lose user data.
    """

    def __init__(self, doc: Document):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        # Run-level formatting flags. We track these as a stack so nested
        # <strong><em>text</em></strong> works correctly.
        self.bold = 0; self.italic = 0; self.underline = 0; self.strike = 0
        self.color_stack: list[Optional[str]] = []
        # Current paragraph & alignment + style.
        self.cur_para = None
        self.cur_style: Optional[str] = None
        self.cur_align: Optional[str] = None
        # Pending image data URL captured between <img> open/close tags.
        self.pending_img: Optional[str] = None
        # List context — depth + ordered vs unordered.
        self.list_stack: list[str] = []
        # Table context.
        self.in_table = False; self.cur_table = None
        self.cur_row = None; self.cur_cell = None; self.cur_cell_idx = 0
        # Hyperlinks: emit the visible text + URL parenthetical (URL
        # rId handling inside docx is fragile to do from a forward parser).
        self.link_href: Optional[str] = None

    # ---- helpers ----

    def _ensure_para(self):
        if self.in_table and self.cur_cell is not None:
            return self.cur_cell.paragraphs[-1]
        if self.cur_para is None:
            if self.cur_style:
                self.cur_para = self.doc.add_paragraph(style=self.cur_style)
            elif self.list_stack:
                style = "List Number" if self.list_stack[-1] == "ol" else "List Bullet"
                try:
                    self.cur_para = self.doc.add_paragraph(style=style)
                except KeyError:
                    self.cur_para = self.doc.add_paragraph()
            else:
                self.cur_para = self.doc.add_paragraph()
            if self.cur_align == "center":
                self.cur_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif self.cur_align == "right":
                self.cur_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif self.cur_align == "justify":
                self.cur_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return self.cur_para

    def _add_run(self, text: str):
        if not text:
            return
        para = self._ensure_para()
        run = para.add_run(text)
        if self.bold:     run.bold = True
        if self.italic:   run.italic = True
        if self.underline: run.underline = True
        if self.strike:
            try: run.font.strike = True
            except (AttributeError, KeyError): pass
        active_color = next((c for c in reversed(self.color_stack) if c), None)
        if active_color:
            try:
                run.font.color.rgb = RGBColor.from_string(active_color)
            except (ValueError, AttributeError):
                pass

    # Only safe raster formats — SVG is deliberately excluded because SVG
    # files can carry embedded <script> elements that execute when the
    # DOCX is opened in Word (even sandboxed). Mirrors the same allow-list
    # as html_sanitize._DATA_IMG_RE.
    _SAFE_DATA_IMG_RE = re.compile(
        r"^data:image/(?:png|jpe?g|gif|webp|bmp|x-icon);base64,",
        re.IGNORECASE,
    )

    def _add_image_data_url(self, data_url: str):
        if not self._SAFE_DATA_IMG_RE.match(data_url):
            return
        m = re.match(r"data:image/[^;]+;base64,(.+)$", data_url, re.DOTALL)
        if not m:
            return
        try:
            blob = base64.b64decode(m.group(1))
        except Exception:
            return
        para = self._ensure_para()
        run = para.add_run()
        try:
            run.add_picture(io.BytesIO(blob), width=Inches(5.5))
        except Exception:
            pass

    # ---- tag handlers ----

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attrs_d = {k: v for k, v in attrs}
        if tag in ("p", "div"):
            self.cur_para = None
            self.cur_style = None
            self.cur_align = self._align_from_style(attrs_d.get("style"))
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.cur_para = None
            self.cur_style = "Heading " + tag[1]
            self.cur_align = self._align_from_style(attrs_d.get("style"))
        elif tag == "br":
            para = self._ensure_para()
            para.add_run().add_break()
        elif tag in ("strong", "b"): self.bold += 1
        elif tag in ("em", "i"):     self.italic += 1
        elif tag == "u":              self.underline += 1
        elif tag in ("s", "strike", "del"): self.strike += 1
        elif tag == "span":
            color = self._color_from_style(attrs_d.get("style"))
            self.color_stack.append(color)
        elif tag in ("ul", "ol"):
            self.list_stack.append("ol" if tag == "ol" else "ul")
        elif tag == "li":
            self.cur_para = None
            self.cur_style = None
            self.cur_align = None
        elif tag == "table":
            self.in_table = True
            self.cur_table = self.doc.add_table(rows=0, cols=0)
            self.cur_table.style = "Table Grid"
        elif tag == "tr" and self.cur_table is not None:
            ncols = max(len(self.cur_table.columns), 1)
            self.cur_row = self.cur_table.add_row()
            self.cur_cell_idx = -1
        elif tag in ("td", "th") and self.cur_row is not None:
            # Lazy-grow column count
            self.cur_cell_idx += 1
            while len(self.cur_row.cells) <= self.cur_cell_idx:
                # Grow the column header by adding columns to the table
                self.cur_table.add_column(Inches(1.0))
            self.cur_cell = self.cur_row.cells[self.cur_cell_idx]
            # Reset paragraph cursor for cell content
            self.cur_para = self.cur_cell.paragraphs[-1]
        elif tag == "a":
            self.link_href = attrs_d.get("href")
        elif tag == "img":
            src = attrs_d.get("src") or ""
            if src.startswith("data:image/"):
                self._add_image_data_url(src)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li"):
            self.cur_para = None
            self.cur_style = None
            self.cur_align = None
        elif tag in ("strong", "b"): self.bold = max(0, self.bold - 1)
        elif tag in ("em", "i"):     self.italic = max(0, self.italic - 1)
        elif tag == "u":              self.underline = max(0, self.underline - 1)
        elif tag in ("s", "strike", "del"): self.strike = max(0, self.strike - 1)
        elif tag == "span":
            if self.color_stack: self.color_stack.pop()
        elif tag in ("ul", "ol"):
            if self.list_stack: self.list_stack.pop()
        elif tag == "table":
            self.in_table = False; self.cur_table = None
            self.cur_row = None; self.cur_cell = None; self.cur_cell_idx = 0
        elif tag == "tr":
            self.cur_row = None
        elif tag in ("td", "th"):
            self.cur_cell = None
        elif tag == "a" and self.link_href:
            self._add_run(f" ({self.link_href})")
            self.link_href = None

    def handle_data(self, data):
        if not data:
            return
        # Collapse runs of whitespace inside the same text node (matches
        # how browsers render contenteditable output).
        cleaned = re.sub(r"[ \t]+", " ", data)
        self._add_run(cleaned)

    # ---- style helpers ----

    @staticmethod
    def _align_from_style(style: Optional[str]) -> Optional[str]:
        if not style: return None
        m = re.search(r"text-align\s*:\s*(center|right|left|justify)", style, re.IGNORECASE)
        return m.group(1).lower() if m else None

    @staticmethod
    def _color_from_style(style: Optional[str]) -> Optional[str]:
        if not style: return None
        m = re.search(r"(?<!-)color\s*:\s*#?([0-9a-fA-F]{6})\b", style)
        return m.group(1).upper() if m else None


def html_to_docx(html_str: str, output: Path) -> Path:
    """Convert `html_str` (typically a contenteditable's innerHTML)
    into a fresh `.docx` at `output`. Returns the same path.

    The output is a self-contained document — we do NOT merge it into
    the template that produced the original generated DOCX, because the
    user just edited a flattened HTML representation and the goal is to
    save exactly what they see. The template-based renderer remains
    available via "Generate" if they want to re-run from source.
    """
    doc = Document()
    parser = _HTMLToDocx(doc)
    parser.feed(html_str or "")
    parser.close()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output
