"""HCR custom-benchmark → CIS-benchmark mapping extractor.

A Host Configuration Review (HCR) engagement is usually run against a
*client's own* hardening standard — a Word doc or PDF that, control by
control, says "this is derived from / equivalent to CIS <Benchmark>
control <id>". Building the cross-reference table by hand is slow and
error-prone. This module automates it:

  1. Pull the raw text out of the uploaded document (.docx via
     python-docx, .pdf via pypdf). Tables in .docx are flattened
     row-by-row so a "Control | Description | CIS Ref" table still
     yields linear text the extractor can pair up.
  2. Detect every CIS Benchmark *title* reference, e.g.
        "CIS Microsoft Windows Server 2019 Benchmark v1.3.0"
        "CIS Ubuntu Linux 22.04 LTS Benchmark v1.0.0 Level 1"
  3. Detect every CIS *control id*, e.g. "1.1.1", "2.3.10.5",
     "CIS 18.9.45.1", "Recommendation 5.2.3".
  4. Detect the client's own custom control ids (e.g. "HCR-001",
     "STD-1.4", "SEC.3.2") so each mapping row is anchored to the
     client's numbering, not just ours.
  5. Emit a styled .xlsx mapping workbook + a JSON preview.

The matching is line/paragraph-windowed: a CIS control id is paired
with the nearest preceding custom-control id (same paragraph first,
then the most recent one seen while walking the document top-to-
bottom). This mirrors how these documents are actually written — the
CIS ref sits inside or right after the control it annotates.

Conservative by design: when no custom id is in scope we still record
the CIS reference (custom id blank) so nothing is silently dropped;
the consultant reviews + adjusts the workbook.

Public surface
--------------
  * ``extract_text(filename, data) -> list[str]``  — ordered text
    segments (paragraphs / table rows / pdf lines).
  * ``find_cis_references(segments) -> MappingResult`` — the parsed
    mapping rows + detected benchmark titles + diagnostics.
  * ``build_xlsx(result) -> bytes`` — styled workbook.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field


# ── CIS benchmark *title* — "CIS <product/os ...> Benchmark vX.Y.Z"
# Captures the whole title incl. optional version + optional
# "Level 1/2" so the consultant sees exactly which edition was cited.
_CIS_TITLE_RE = re.compile(
    r"CIS\s+[A-Za-z0-9][\w .,/&()+-]{2,90}?\s+Benchmark"
    r"(?:\s+v?\d+(?:\.\d+){0,3})?"
    r"(?:\s*(?:[-,]?\s*Level\s*[12]))?",
    re.IGNORECASE,
)

# ── CIS *control id*. Three accepted shapes:
#   * dotted numeric, 2-6 segments:  1.1.1   2.3.10.5   18.9.45.1
#   * prefixed:  "CIS 1.1.1" / "CIS Control 5.2"
#   * "Recommendation 1.1.1" / "Rec. 1.1.1"
# The dotted-numeric core requires at least ONE dot so we don't grab
# bare section numbers like "5" or years like "2019".
_CIS_ID_CORE = r"\d{1,2}(?:\.\d{1,3}){1,5}"
_CIS_ID_RE = re.compile(
    r"(?:(?:CIS|CIS\s+Control|Recommendation|Rec\.?|Control)\s+)?"
    rf"(?<![\w.])({_CIS_ID_CORE})(?![\w.])",
    re.IGNORECASE,
)
# A *strong* CIS-id signal: the dotted id is explicitly prefixed with
# CIS / Recommendation / Control. Used to avoid pairing unrelated
# dotted numbers (e.g. a software version "1.2.3") when there is NO
# CIS title anywhere in the document.
_CIS_ID_STRONG_RE = re.compile(
    r"(?:CIS|CIS\s+Control|Recommendation|Rec\.?|Control)\s+"
    rf"(?<![\w.])({_CIS_ID_CORE})(?![\w.])",
    re.IGNORECASE,
)

# ── Client custom-control id. Common shapes seen on real client
# hardening standards. Deliberately broad but anchored so prose words
# don't match: an UPPER/Mixed token of 2-8 chars, a separator, then a
# numeric/dotted tail.  HCR-001  STD-1.4  SEC.3.2  CFG_12  POL-4.1.2
#
# The negative-lookahead excludes the CIS-reference keywords
# (CIS / REC / RECOMMENDATION / CONTROL) so a "CIS 1.1.1" reference is
# NEVER mis-parsed as a client control id — that bug made the masking
# step wipe the real CIS ids out of every table row.
_CUSTOM_ID_RE = re.compile(
    r"\b(?!(?:CIS|REC|RECOMMENDATION|CONTROL)\b)"
    r"([A-Z][A-Z0-9]{1,7}[-_. ]\d{1,3}(?:\.\d{1,3}){0,3})\b"
)


@dataclass
class MappingRow:
    custom_id: str
    cis_benchmark: str
    cis_control_ids: str          # comma-joined, de-duped, order-preserved
    context: str                  # trimmed source snippet for audit


@dataclass
class MappingResult:
    rows: list[MappingRow] = field(default_factory=list)
    benchmark_titles: list[str] = field(default_factory=list)
    n_segments: int = 0
    n_cis_ids_total: int = 0
    n_custom_ids_total: int = 0
    warnings: list[str] = field(default_factory=list)

    def as_preview(self, cap: int = 200) -> dict:
        return {
            "benchmark_titles": self.benchmark_titles,
            "row_count": len(self.rows),
            "cis_ids_detected": self.n_cis_ids_total,
            "custom_ids_detected": self.n_custom_ids_total,
            "segments_scanned": self.n_segments,
            "warnings": self.warnings,
            "rows": [
                {
                    "custom_id": r.custom_id,
                    "cis_benchmark": r.cis_benchmark,
                    "cis_control_ids": r.cis_control_ids,
                    "context": r.context,
                }
                for r in self.rows[:cap]
            ],
            "truncated": len(self.rows) > cap,
        }


# ──────────────────────────────────────────────────────────────────
# 1. Text extraction
# ──────────────────────────────────────────────────────────────────

def extract_text(filename: str, data: bytes) -> list[str]:
    """Return ordered text segments from a .docx or .pdf upload.

    .docx: every paragraph + every table row (cells joined by ' | ')
           in document order, so a control table linearises cleanly.
    .pdf : every non-empty line across all pages, in page order.

    Raises ValueError on an unsupported extension or a corrupt file —
    the caller surfaces it as a 400.
    """
    name = (filename or "").lower()
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    raise ValueError(
        "Unsupported file type — upload a .docx or .pdf hardening "
        "standard / HCR document."
    )


def _extract_docx(data: bytes) -> list[str]:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:                                     # pragma: no cover
        raise ValueError("python-docx is not available on the server.")
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ValueError(f"Could not open the .docx ({e}).")

    segments: list[str] = []
    body = doc.element.body
    # Walk the body in true document order so a paragraph that sits
    # between two tables keeps its position relative to them. We map
    # each <w:p>/<w:tbl> child back to the python-docx wrapper.
    para_by_el = {p._element: p for p in doc.paragraphs}
    tbl_by_el = {t._element: t for t in doc.tables}
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_by_el.get(child)
            if p is not None:
                txt = (p.text or "").strip()
                if txt:
                    segments.append(txt)
        elif child.tag == qn("w:tbl"):
            t = tbl_by_el.get(child)
            if t is not None:
                for row in t.rows:
                    cells = [
                        " ".join((c.text or "").split()).strip()
                        for c in row.cells
                    ]
                    line = " | ".join(c for c in cells if c)
                    if line:
                        segments.append(line)
    return segments


def _extract_pdf(data: bytes) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError:                                     # pragma: no cover
        raise ValueError("pypdf is not available on the server.")
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise ValueError(f"Could not open the .pdf ({e}).")

    segments: list[str] = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        for ln in txt.splitlines():
            ln = ln.strip()
            if ln:
                segments.append(ln)
    return segments


# ──────────────────────────────────────────────────────────────────
# 2. Reference detection + pairing
# ──────────────────────────────────────────────────────────────────

def _dedupe(seq: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for s in seq:
        if s not in seen:
            seen.add(s)
            out.append(s)
    return out


def find_cis_references(segments: list[str]) -> MappingResult:
    """Pair every CIS control id in the document with the nearest
    in-scope custom control id, and attach the benchmark title that's
    in effect at that point.

    Algorithm (single forward pass — documents are written top-down):
      * Track the most recently seen custom-control id.
      * Track the most recently seen benchmark title (titles often
        appear once in a section header and apply to the rows below).
      * For every segment, collect CIS ids. If the segment ALSO
        contains a custom id, that custom id owns the CIS ids in the
        same segment (the common table-row case
        "HCR-012 | ... | CIS 1.2.3"). Otherwise the CIS ids attach to
        the most recent custom id seen so far.
      * A CIS id is only emitted when EITHER a benchmark title has
        appeared somewhere in the doc OR the id is strongly prefixed
        ("CIS 1.2.3" / "Recommendation 1.2.3"). This stops a stray
        software version like "Apache 2.4.58" being mis-read as a
        control id in documents that never mention CIS at all.
    """
    result = MappingResult(n_segments=len(segments))

    all_titles: list[str] = []
    for seg in segments:
        for m in _CIS_TITLE_RE.finditer(seg):
            all_titles.append(" ".join(m.group(0).split()))
    result.benchmark_titles = _dedupe(all_titles)
    any_title = bool(result.benchmark_titles)

    # rows keyed by (custom_id, benchmark) so multiple CIS ids under
    # the same control collapse into one row with a joined id list.
    agg: dict[tuple[str, str], MappingRow] = {}
    cur_custom = ""
    cur_title = result.benchmark_titles[0] if result.benchmark_titles else ""

    for seg in segments:
        # Update the in-effect benchmark title when this segment names one.
        tm = _CIS_TITLE_RE.search(seg)
        if tm:
            cur_title = " ".join(tm.group(0).split())

        # Custom id in THIS segment owns same-segment CIS ids.
        cm = _CUSTOM_ID_RE.search(seg)
        seg_custom = cm.group(1) if cm else ""
        if seg_custom:
            cur_custom = seg_custom
            result.n_custom_ids_total += 1

        # Build the text we scan for CIS ids with the contaminating
        # tokens MASKED out first, so their embedded numbers can't be
        # mis-read as control ids:
        #   * every custom-id match ("STD-3.4", "LNX-1.1") — its "3.4"
        #     / "1.1" tail is NOT a CIS id.
        #   * every benchmark title ("CIS Ubuntu Linux 22.04 LTS
        #     Benchmark v1.0.0") — the "22.04" / "1.0.0" are an OS
        #     release + a benchmark version, not control ids.
        scan = seg
        for cmx in _CUSTOM_ID_RE.finditer(seg):
            scan = scan.replace(cmx.group(1), " ")
        for tmx in _CIS_TITLE_RE.finditer(seg):
            scan = scan.replace(tmx.group(0), " CIS Benchmark ")

        # Collect CIS ids from the MASKED text. Strong (prefixed) ids
        # always count; bare dotted ids only count when a benchmark
        # title exists somewhere in the document.
        strong_ids = [m.group(1) for m in _CIS_ID_STRONG_RE.finditer(scan)]
        if any_title:
            loose_ids = [m.group(1) for m in _CIS_ID_RE.finditer(scan)]
        else:
            loose_ids = []
        cis_ids = _dedupe(strong_ids + loose_ids)
        # Belt-and-braces: never let the custom id itself echo back.
        cis_ids = [c for c in cis_ids if c != seg_custom and c != cur_custom]
        if not cis_ids:
            continue
        result.n_cis_ids_total += len(cis_ids)

        owner = seg_custom or cur_custom or ""
        key = (owner, cur_title)
        row = agg.get(key)
        if row is None:
            row = MappingRow(
                custom_id=owner,
                cis_benchmark=cur_title,
                cis_control_ids="",
                context=seg[:300],
            )
            agg[key] = row
        merged = _dedupe(
            ([x for x in row.cis_control_ids.split(", ") if x]) + cis_ids
        )
        row.cis_control_ids = ", ".join(merged)

    # Stable order: by first appearance (dict preserves insertion).
    result.rows = list(agg.values())

    if not result.rows:
        result.warnings.append(
            "No CIS control references were detected. If the document "
            "uses an unusual numbering scheme, check it manually — the "
            "extractor looks for dotted control ids (e.g. 1.2.3) and "
            "'CIS … Benchmark' titles."
        )
    if not any_title and result.rows:
        result.warnings.append(
            "No 'CIS … Benchmark' title was found, so only explicitly "
            "prefixed ids ('CIS 1.2.3', 'Recommendation 1.2.3') were "
            "mapped. Bare dotted numbers were skipped to avoid false "
            "positives — verify nothing was missed."
        )
    if any(not r.custom_id for r in result.rows):
        result.warnings.append(
            "Some CIS references could not be tied to a client control "
            "id (no custom id like 'HCR-001' / 'STD-1.4' was in scope). "
            "Those rows have a blank Custom Control — fill it in "
            "manually from the source document."
        )
    return result


# ──────────────────────────────────────────────────────────────────
# 3. XLSX builder
# ──────────────────────────────────────────────────────────────────

def build_xlsx(result: MappingResult) -> bytes:
    """Render the mapping as a styled .xlsx and return the bytes.

    Sheet 1 "CIS Mapping": one row per (custom control, benchmark).
    Sheet 2 "Benchmarks":  the distinct CIS benchmark titles cited.
    Sheet 3 "Summary":     counts + any extraction warnings.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    hdr_fill = PatternFill("solid", fgColor="1F3A5F")
    hdr_font = Font(bold=True, color="FFFFFF")
    wrap = Alignment(vertical="top", wrap_text=True)
    thin = Side(style="thin", color="D9D9D9")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _formula_safe(v: str) -> str:
        # Excel treats a leading = + - @ as a formula — prefix a quote
        # so a malicious / accidental "=cmd|..." cell stays inert.
        s = "" if v is None else str(v)
        return "'" + s if s[:1] in ("=", "+", "-", "@") else s

    # ── Sheet 1: CIS Mapping
    ws = wb.active
    ws.title = "CIS Mapping"
    headers = ["#", "Custom Control", "CIS Benchmark",
               "CIS Control ID(s)", "Source Context"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = wrap
        cell.border = box
    for i, r in enumerate(result.rows, 1):
        vals = [i, r.custom_id or "(unassigned)", r.cis_benchmark or "(none)",
                r.cis_control_ids, r.context]
        for c, v in enumerate(vals, 1):
            cell = ws.cell(row=i + 1, column=c,
                           value=v if c == 1 else _formula_safe(v))
            cell.alignment = wrap
            cell.border = box
    widths = [5, 22, 46, 26, 70]
    for c, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(c)].width = w
    ws.freeze_panes = "A2"
    if result.rows:
        ws.auto_filter.ref = f"A1:E{len(result.rows) + 1}"

    # ── Sheet 2: Benchmarks
    ws2 = wb.create_sheet("Benchmarks")
    ws2.cell(row=1, column=1, value="CIS Benchmark Titles Referenced")
    ws2["A1"].fill = hdr_fill
    ws2["A1"].font = hdr_font
    ws2.column_dimensions["A"].width = 80
    for i, t in enumerate(result.benchmark_titles, 1):
        ws2.cell(row=i + 1, column=1, value=_formula_safe(t)).alignment = wrap
    if not result.benchmark_titles:
        ws2.cell(row=2, column=1,
                 value="(no 'CIS … Benchmark' title detected)")

    # ── Sheet 3: Summary
    ws3 = wb.create_sheet("Summary")
    rows = [
        ("Mapping rows", len(result.rows)),
        ("Distinct CIS benchmarks cited", len(result.benchmark_titles)),
        ("Total CIS control ids detected", result.n_cis_ids_total),
        ("Total custom control ids detected", result.n_custom_ids_total),
        ("Document segments scanned", result.n_segments),
    ]
    for i, (k, v) in enumerate(rows, 1):
        ws3.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws3.cell(row=i, column=2, value=v)
    ws3.column_dimensions["A"].width = 36
    ws3.column_dimensions["B"].width = 12
    base = len(rows) + 2
    ws3.cell(row=base, column=1, value="Warnings").font = Font(bold=True)
    if result.warnings:
        for j, w in enumerate(result.warnings, 1):
            ws3.cell(row=base + j, column=1,
                     value="• " + _formula_safe(w)).alignment = wrap
        ws3.column_dimensions["A"].width = 90
    else:
        ws3.cell(row=base + 1, column=1, value="• none")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
