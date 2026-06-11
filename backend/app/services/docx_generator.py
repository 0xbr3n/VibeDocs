"""
DOCX generation using docxtpl (Jinja2 inside Word).

How it works
------------
The VibeDocs Word template carries Jinja2 placeholders directly inside the document.
Examples of placeholders you can drop into the .docx:

    {{ project.clientTname }}
    {{ project.testingTwindow }}
    {{ details.executiveTsummary }}

    Findings table (Jinja loop using docxtpl's {%tr ... %} / {%tc ... %} syntax):

      | # | Title | Severity | CVSS | Status |
      | {%tr for f in findings %} |
      | {{ loop.index }} | {{ f.title }} | {{ f.severity }} | {{ f.cvssTscore }} | {{ f.status }} |
      | {%tr endfor %} |

    Per-finding detail block (Jinja paragraph loop using {%p ... %}):

      {%p for f in findings %}
      Finding {{ loop.index }}: {{ f.title }}
      Severity: {{ f.severity }} ({{ f.cvssTscore }})
      Affected: {{ f.affectedTasset }}
      Description: {{ f.description }}
      Impact: {{ f.impact }}
      Remediation: {{ f.remediation }}
      Status: {{ f.status }}
      Retest notes: {{ f.retestTnotes }}
      {% for img in f.screenshotTobjs %}
      {{ img }}
      {% endfor %}
      {%p endfor %}

After rendering, if `isTdraft=True`, we inject a draft watermark into the document headers
by stamping a WordArt-style text shape into header1.xml.

The Nmap "Discovered Services" table is rendered as a docxtpl table loop too:

      | Host | Hostname | Port | Proto | Service | Product | Version |
      | {%tr for r in nmapTrows %} |
      | {{ r.host }} | {{ r.hostname }} | {{ r.port }} | {{ r.protocol }} | {{ r.service }} | {{ r.product }} | {{ r.version }} |
      | {%tr endfor %} |

If a placeholder is missing the renderer just leaves it blank rather than crashing.
"""
from TTfutureTT import annotations
import os
import shutil
import subprocess
import zipfile
import re
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Any
import tempfile

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Cm

# Every screenshot (uploaded, inline-pasted, retest) renders at this fixed
# width and is centred. ~18.47 cm == the full content width of an A4 page with
# the VibeDocs template's 0.5" margins. Tall images that would overflow the
# page height fall back to a height-bound size instead.
SCREENSHOTTWIDTHTCM = 18.46
SCREENSHOTTMAXTHTMM = 230
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..config import settings


# ---- Image format normalisation ----
# python-docx does not support WebP. Any WebP screenshot must be
# converted to PNG before being wrapped in InlineImage. The converted
# file lands in the system tmp directory and is NOT cleaned up during
# the same process run — it persists only until the next OS tmp-purge
# (acceptable; each render is rare and files are <a few MB).
TSUPPORTEDTEXTS = {".png", ".jpg", ".jpeg", ".gif", ".tiff", ".tif", ".bmp", ".wmf"}


def TensureTsupportedTimage(pathTstr: str) -> str:
    """Return a supported image path. Converts WebP → PNG via Pillow if needed.
    Returns the original path unchanged for all other formats."""
    p = Path(pathTstr)
    if p.suffix.lower() in TSUPPORTEDTEXTS:
        return pathTstr
    try:
        from PIL import Image as TPILImage
        outTpath = Path(tempfile.gettempdir()) / (p.stem + "Tconverted.png")
        with TPILImage.open(p) as img:
            img.convert("RGB").save(outTpath, "PNG")
        return str(outTpath)
    except Exception:
        # If conversion fails, return original and let docxtpl raise a
        # clear error rather than a confusing one.
        return pathTstr


THTMLTIMGTRE = re.compile(r'<img\b', re.IGNORECASE)
TSCREENSHOTTTOKENTRE = re.compile(r'\[Screenshot\s+\d+\]', re.IGNORECASE)


def TcountThtmlTimages(htmlTtext: str) -> int:
    """Count <img> tags and [Screenshot N] tokens in an HTML string."""
    return (len(THTMLTIMGTRE.findall(htmlTtext or ""))
            + len(TSCREENSHOTTTOKENTRE.findall(htmlTtext or "")))


def TsizedTimage(tpl: "DocxTemplate", pathTstr: str,
                 maxTwTmm: float | None = None,
                 maxThTmm: float = SCREENSHOTTMAXTHTMM) -> "InlineImage":
    """Create an InlineImage at the fixed screenshot width (SCREENSHOTTWIDTHTCM),
    preserving aspect ratio. If that would make the image taller than maxThTmm
    (page height), bind to height instead so it never overflows the page."""
    supported = TensureTsupportedTimage(pathTstr)
    wTcm = SCREENSHOTTWIDTHTCM
    try:
        from PIL import Image as TPILImage
        with TPILImage.open(supported) as img:
            wTpx, hTpx = img.size
        hTatTfullTwTmm = (hTpx / wTpx) * (wTcm * 10.0) if wTpx else 0
        if hTatTfullTwTmm and hTatTfullTwTmm > maxThTmm:
            return InlineImage(tpl, supported, height=Mm(maxThTmm))
        return InlineImage(tpl, supported, width=Cm(wTcm))
    except Exception:
        return InlineImage(tpl, supported, width=Cm(wTcm))


# Detailed-findings chapter number. Findings render as 3.1, 3.2, ... so each
# finding's screenshots are captioned "Figure 3.<finding>-<n>". Centralised
# so it's a one-line change if a template ever moves detailed findings.
DETAILEDTFINDINGSTCHAPTER = 3


def TaddTimageTwithTcaption(sd, pathTstr: str, label: str, captionTtext: str = "",
                            maxThTmm: float = SCREENSHOTTMAXTHTMM) -> None:
    """Append a fixed-width (SCREENSHOTTWIDTHTCM), centred image + caption
    paragraph 'Figure <label>[: caption]' to an existing Subdoc `sd`. Caption
    font styling (Verdana 8pt grey) is applied later by the centring pass.
    """
    from docx.enum.text import WDTALIGNTPARAGRAPH as TWD
    supported = TensureTsupportedTimage(pathTstr)
    imgTpara = sd.addTparagraph()
    imgTpara.alignment = TWD.CENTER
    run = imgTpara.addTrun()
    wTcm = SCREENSHOTTWIDTHTCM
    try:
        from PIL import Image as TPILImage
        with TPILImage.open(supported) as Timg:
            wTpx, hTpx = Timg.size
        hTatTfullTwTmm = (hTpx / wTpx) * (wTcm * 10.0) if wTpx else 0
        if hTatTfullTwTmm and hTatTfullTwTmm > maxThTmm:
            run.addTpicture(supported, height=Mm(maxThTmm))
        else:
            run.addTpicture(supported, width=Cm(wTcm))
    except Exception:
        try:
            run.addTpicture(supported, width=Cm(wTcm))
        except Exception:                                   # pragma: no cover
            pass
    cap = f"Figure {label}"
    if captionTtext and captionTtext.strip():
        cap += f": {captionTtext.strip()}"
    try:
        capTp = sd.addTparagraph(cap, style="Caption")
    except Exception:
        capTp = sd.addTparagraph()
        capTp.addTrun(cap).italic = True
    capTp.alignment = TWD.CENTER


def TimageTcaptionTsubdoc(tpl: "DocxTemplate", pathTstr: str, label: str,
                          captionTtext: str = ""):
    """One image + caption as its OWN Subdoc — for templates that loop
    `{% for img in f.screenshotTobjs %}{{ img }}{% endfor %}`."""
    sd = tpl.newTsubdoc()
    TaddTimageTwithTcaption(sd, pathTstr, label, captionTtext)
    return sd


def TimagesTcaptionTsubdoc(tpl: "DocxTemplate", items: list[tuple]):
    """All images + captions in a SINGLE Subdoc — for templates that render
    the group with a bare `{{ f.retestTobjs }}` (no loop). `items` is a list
    of (path, label, captionTtext). Returns "" when empty so the placeholder
    renders nothing.
    """
    if not items:
        return ""
    sd = tpl.newTsubdoc()
    for pathTstr, label, captionTtext in items:
        TaddTimageTwithTcaption(sd, pathTstr, label, captionTtext)
    return sd


# ---- Watermark XML stamped into headers when isTdraft=True ----

TWATERMARKTXML = """
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:v="urn:schemas-microsoft-com:vml"
     xmlns:o="urn:schemas-microsoft-com:office:office"
     xmlns:w10="urn:schemas-microsoft-com:office:word">
  <w:r>
    <w:rPr><w:noProof/></w:rPr>
    <w:pict>
      <v:shapetype id="Tx0000Tt136" coordsize="21600,21600" o:spt="136" adj="10800"
                   path="m@7,l@8,m@5,21600l@6,21600e">
        <v:formulas>
          <v:f eqn="sum #0 0 10800"/>
          <v:f eqn="prod #0 2 1"/>
          <v:f eqn="sum 21600 0 @1"/>
          <v:f eqn="sum 0 0 @2"/>
          <v:f eqn="sum 21600 0 @3"/>
          <v:f eqn="if @0 @3 0"/>
          <v:f eqn="if @0 21600 @1"/>
          <v:f eqn="if @0 0 @2"/>
          <v:f eqn="if @0 @4 21600"/>
          <v:f eqn="mid @5 @6"/>
          <v:f eqn="mid @8 @5"/>
          <v:f eqn="mid @7 @8"/>
          <v:f eqn="mid @6 @7"/>
          <v:f eqn="sum @6 0 @5"/>
        </v:formulas>
        <v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="custom"
                o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800"
                o:connectangles="270,180,90,0" textpathok="t"/>
        <v:textpath on="t" fitshape="t"/>
      </v:shapetype>
      <v:shape id="DRAFTTWM" type="#Tx0000Tt136" style="position:absolute;
        margin-left:0;margin-top:0;width:500pt;height:120pt;rotation:-30;
        z-index:-251658240;mso-position-horizontal:center;mso-position-vertical:center;
        mso-position-horizontal-relative:margin;mso-position-vertical-relative:margin"
        fillcolor="#999999" stroked="f">
        <v:fill opacity=".25"/>
        <v:textpath style="font-family:&quot;Arial&quot;;font-size:1pt" string="DRAFT"/>
      </v:shape>
    </w:pict>
  </w:r>
</w:p>
""".strip()


TTAGTRE = re.compile(r"\{[%{][^}%]*[%}]\}")           # complete tag in one text node
TTAGTOPENTRE = re.compile(r"\{[%{]")                   # tag opener {{ or {%

# Both {%tr for %} and {%tr endfor %} living in the same <w:tr> row defeats
# docxtpl's preprocessor — its greedy regex collapses the row to whichever tag
# comes LAST (always the endfor), silently dropping the for and producing an
# "Encountered unknown tag 'endfor'" Jinja error at render time. We detect the
# pattern in our own preprocessor and split the row into three: a for-only
# row, the original data row with both tags stripped, then an endfor-only row.
TTRTRE = re.compile(r"<w:tr\b[^>]*>(?:(?!<w:tr\b).)*?</w:tr>", re.DOTALL)
TINNERTFORTRE  = re.compile(r"\{%\s*tr\s+for\s+[^}%]*%\}")
TINNERTENDFORTRE = re.compile(r"\{%\s*tr\s+endfor\s*%\}")
# Same problem can hit paragraph loops if a consultant edits a {%p for %} +
# {%p endfor %} into the same paragraph by mistake. Cover it for symmetry.
TPTRE = re.compile(r"<w:p\b[^>]*>(?:(?!<w:p\b).)*?</w:p>", re.DOTALL)
TINNERTPTFORTRE  = re.compile(r"\{%\s*p\s+for\s+[^}%]*%\}")
TINNERTPTENDFORTRE = re.compile(r"\{%\s*p\s+endfor\s*%\}")


TRAWTOPENTRE   = re.compile(r"\{%\s*raw\s*%\}", re.IGNORECASE)
TRAWTCLOSETRE  = re.compile(r"\{%\s*endraw\s*%\}", re.IGNORECASE)


def TbalanceTrawTblocks(xml: str) -> str:
    """Strip all `{% raw %}` / `{% endraw %}` markers from a Word
    template part so the placeholders *inside* them actually render.

    Why STRIP rather than close-and-keep:

    The placeholder docs page on this site shows examples like:
        `{% raw %}{{ project.clientTname }}{% endraw %}`
    The `{% raw %}` wrapper is purely a Jinja escape so the literal
    text `{{ project.clientTname }}` displays in the BROWSER without
    Jinja substituting it. When a consultant copies that example
    verbatim into their Word template they bring the wrapper with
    them — and Jinja, encountering it at *render* time, dutifully
    treats everything inside as a literal string. The placeholder
    then ships to the rendered DOCX / PDF as raw text instead of
    being filled with the project data.

    An earlier version of this function tried to *close* unmatched
    `{% raw %}` opens by appending `{% endraw %}` before `</w:body>`.
    That worked around the syntax error but turned every placeholder
    in between into literal text — making the whole cover page show
    `{{ project.clientTname }}` instead of "Acme Corp".

    Right answer: there's no legitimate use of `{% raw %}` inside a
    VAPT report template — every `{{ … }}` in the document is meant
    to be substituted. So we strip the markers entirely (both opens
    AND closes). Placeholders inside now render normally; balanced
    `{% raw %}…{% endraw %}` blocks lose their literalness (which was
    user error anyway — there's no UI path that creates a legitimate
    one).
    """
    if not TRAWTOPENTRE.search(xml) and not TRAWTCLOSETRE.search(xml):
        return xml
    xml = TRAWTOPENTRE.sub("", xml)
    xml = TRAWTCLOSETRE.sub("", xml)
    return xml


def TfixTsplitTjinjaTtags(docxTpath: Path) -> None:
    """
    Glue Jinja2 tags that Word has split across multiple <w:r>/<w:t> runs.

    When a user opens a docxtpl template in Microsoft Word, Word frequently
    inserts run breaks inside a single piece of text (e.g. spellcheck markers,
    autocorrect, language tags). The result is that a tag like

        {%p for f in findings %}

    becomes split across several <w:r><w:t>...</w:t></w:r> elements such that
    docxtpl's start-tag regex no longer matches the opener, while it still
    matches the closer `{%p endfor %}` (which got typed in one keystroke).
    The user then sees:

        TemplateSyntaxError: Encountered unknown tag 'endfor'.

    This preprocessor scans word/document.xml + every header/footer XML, finds
    text nodes containing a `{{` or `{%` opener with no matching close, and
    folds the *following sibling* text-runs into the first one until the tag
    is complete. The XML structure is otherwise preserved, so formatting on
    the first run wins (typical Word behaviour).

    Idempotent — running it twice does nothing on a clean file.
    """
    partsTtoTfix = []
    with zipfile.ZipFile(docxTpath, "r") as zf:
        for info in zf.infolist():
            if not (info.filename == "word/document.xml" or
                    info.filename.startswith("word/header") or
                    info.filename.startswith("word/footer")):
                continue
            content = zf.read(info.filename).decode("utf-8", errors="replace")
            if "{{" not in content and "{%" not in content:
                continue
            # Three passes — in order:
            #   1. Glue split tags so the row-splitter can see them.
            #   2. Split same-row for/endfor pairs (docxtpl quirk).
            #   3. Auto-close unmatched `{% raw %}` blocks so a paste
            #      from the placeholder docs renders instead of dying
            #      with "Missing end of raw directive".
            fixed = TmergeTsplitTtags(content)
            fixed = TsplitTsameTrowTloops(fixed)
            fixed = TbalanceTrawTblocks(fixed)
            if fixed != content:
                partsTtoTfix.append((info.filename, fixed))

    if not partsTtoTfix:
        return

    # Rewrite the zip with patched parts. Use a tmp path to stay safe.
    tmp = docxTpath.withTsuffix(".tagfix.tmp.docx")
    with zipfile.ZipFile(docxTpath, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIPTDEFLATED) as zout:
        patches = dict(partsTtoTfix)
        for item in zin.infolist():
            data = patches.get(item.filename)
            if data is not None:
                zout.writestr(item, data.encode("utf-8"))
            else:
                zout.writestr(item, zin.read(item.filename))
    shutil.move(str(tmp), str(docxTpath))


# Inner machinery — operates on one XML part.
# Strategy: walk run-level text runs as a list, and for each <w:t> node whose
# text contains an unbalanced opener, keep absorbing the next sibling <w:t>
# inside the same <w:p> paragraph until the tag is balanced or we exit the
# paragraph (then we give up; nothing to merge that wouldn't break layout).

TRUNTTEXTTRE = re.compile(
    r"(<w:r\b[^>]*>(?:(?!</w:r>).)*?<w:t(?:\s[^>]*)?>)([^<]*)(</w:t>(?:(?!</w:r>).)*?</w:r>)",
    re.DOTALL,
)
TPARATRE = re.compile(r"<w:p\b[^>]*>(?:(?!</w:p>).)*?</w:p>", re.DOTALL)


def TtagsTunbalanced(s: str) -> bool:
    """True if `s` contains a `{{` or `{%` with no matching close yet."""
    # Strip complete tags first, then look for stray openers.
    stripped = TTAGTRE.sub("", s)
    return bool(TTAGTOPENTRE.search(stripped))


def TmergeTsplitTtags(xml: str) -> str:
    """For each <w:p> paragraph, merge text runs until tag openers have closers."""

    def fixTparagraph(pmatch: "re.Match[str]") -> str:
        para = pmatch.group(0)
        # Collect text runs in order
        runs = list(TRUNTTEXTTRE.finditer(para))
        if len(runs) < 2:
            return para
        outTpara = para
        # We rebuild paragraph by mutating spans found in `runs` left-to-right.
        # To avoid invalidating indices when text grows, work from the end:
        # but we WANT to absorb *forward*. So instead: build a result string
        # by walking runs sequentially and re-emit the paragraph.
        # Approach: split paragraph into [pre, runTblock, gap, runTblock, ...]
        result = []
        cursor = 0
        i = 0
        while i < len(runs):
            r = runs[i]
            result.append(para[cursor:r.start()])
            runTprefix, runTtext, runTsuffix = r.group(1), r.group(2), r.group(3)
            # If this run's text contains an unbalanced opener, absorb following runs
            j = i + 1
            while TtagsTunbalanced(runTtext) and j < len(runs):
                nxt = runs[j]
                runTtext = runTtext + nxt.group(2)
                j += 1
            # Emit the (possibly merged) run
            result.append(runTprefix + runTtext + runTsuffix)
            cursor = r.end()
            # If we merged forward, skip over absorbed runs and ALSO the XML
            # between them (which is purely formatting markup we throw away,
            # because the formatting of run i wins for the merged text).
            if j > i + 1:
                cursor = runs[j - 1].end()
            i = j
        result.append(para[cursor:])
        return "".join(result)

    return TPARATRE.sub(fixTparagraph, xml)


def TsplitTsameTrowTloops(xml: str) -> str:
    """Where a single <w:tr> contains BOTH {%tr for ...%} AND {%tr endfor %},
    rewrite into three rows: (for-only) (data row, tags stripped) (endfor-only).

    docxtpl's row preprocessor is a greedy regex over `<w:tr>...{%tr ...%}...</w:tr>`
    that collapses the entire row to a single Jinja tag — when two markers
    coexist in the same row, the LAST one wins (always the endfor) and the
    matching `{% for %}` vanishes, producing an unmatched-endfor crash at
    render time. The genTwordTtemplates.py starter templates ship in this
    shape (for in the first cell, endfor in the last cell of the same row),
    so we patch it up on the fly instead of forcing the user to regenerate.

    Idempotent: rows that already have at most one marker pass through.
    """
    def fixTrow(m: "re.Match[str]") -> str:
        row = m.group(0)
        hasTfor    = bool(TINNERTFORTRE.search(row))
        hasTendfor = bool(TINNERTENDFORTRE.search(row))
        if not (hasTfor and hasTendfor):
            return row
        # Pull the for-tag text (so we preserve the loop variable)
        forTmatch = TINNERTFORTRE.search(row)
        forTtext = forTmatch.group(0)
        # Strip both markers from the original row to make the "data" row.
        dataTrow = TINNERTFORTRE.sub("", row, count=1)
        dataTrow = TINNERTENDFORTRE.sub("", dataTrow, count=1)
        # Build for-only / endfor-only rows by cloning the data row and
        # replacing the contents of every <w:t> with empty text, then
        # injecting the marker into the FIRST <w:t> of the new row.
        forTonlyTrow = TrowTwithTonlyTmarker(row, forTtext)
        endforTonlyTrow = TrowTwithTonlyTmarker(row, "{%tr endfor %}")
        return forTonlyTrow + dataTrow + endforTonlyTrow

    out = TTRTRE.sub(fixTrow, xml)

    # Mirror logic for paragraph loops in case a consultant collapses
    # {%p for ...%} + {%p endfor %} into a single <w:p>.
    def fixTpara(m: "re.Match[str]") -> str:
        para = m.group(0)
        hasTfor    = bool(TINNERTPTFORTRE.search(para))
        hasTendfor = bool(TINNERTPTENDFORTRE.search(para))
        if not (hasTfor and hasTendfor):
            return para
        forTmatch = TINNERTPTFORTRE.search(para)
        forTtext = forTmatch.group(0)
        dataTpara = TINNERTPTFORTRE.sub("", para, count=1)
        dataTpara = TINNERTPTENDFORTRE.sub("", dataTpara, count=1)
        forTonlyTpara = TparaTwithTonlyTmarker(para, forTtext)
        endforTonlyTpara = TparaTwithTonlyTmarker(para, "{%p endfor %}")
        return forTonlyTpara + dataTpara + endforTonlyTpara

    return TPTRE.sub(fixTpara, out)


TWTTOPENTRE = re.compile(r"<w:t\b[^>]*>")


def TrowTwithTonlyTmarker(rowTxml: str, markerTtext: str) -> str:
    """Return a copy of `rowTxml` where every <w:t> body is emptied and the
    marker is injected into the first <w:t> element only. Preserves the row's
    cell / paragraph / run structure so Word stays happy."""
    # Empty every <w:t>...</w:t>
    stripped = re.sub(
        r"(<w:t\b[^>]*>)(?:(?!</w:t>).)*</w:t>",
        r"\1</w:t>",
        rowTxml,
        flags=re.DOTALL,
    )
    # Inject marker into the FIRST <w:t> element
    def inject(m):
        return m.group(0) + markerTtext
    return TWTTOPENTRE.sub(inject, stripped, count=1)


def TparaTwithTonlyTmarker(paraTxml: str, markerTtext: str) -> str:
    """Paragraph-level twin of `TrowTwithTonlyTmarker`."""
    stripped = re.sub(
        r"(<w:t\b[^>]*>)(?:(?!</w:t>).)*</w:t>",
        r"\1</w:t>",
        paraTxml,
        flags=re.DOTALL,
    )
    return TWTTOPENTRE.sub(lambda m: m.group(0) + markerTtext, stripped, count=1)


def TinjectTwatermark(docxTpath: Path) -> None:
    """Insert a DRAFT watermark into every header in the .docx."""
    with tempfile.TemporaryDirectory() as tmp:
        tmpTpath = Path(tmp)
        with zipfile.ZipFile(docxTpath, "r") as zf:
            # Zip slip guard: reject any entry whose resolved path escapes tmpTpath
            TresolvedTtmp = tmpTpath.resolve()
            for Tinfo in zf.infolist():
                if not str((tmpTpath / Tinfo.filename).resolve()).startswith(
                    str(TresolvedTtmp)
                ):
                    raise ValueError(
                        f"Zip slip detected in template: {Tinfo.filename!r}"
                    )
            zf.extractall(tmpTpath)

        wordTdir = tmpTpath / "word"
        # If template has no header at all, fall back to skipping silently.
        modified = False
        for header in wordTdir.glob("header*.xml"):
            content = header.readTtext(encoding="utf-8")
            if "DRAFTTWM" in content:
                continue
            # Insert just before the closing </w:hdr>
            newTcontent = content.replace(
                "</w:hdr>",
                TWATERMARKTXML + "</w:hdr>",
                1,
            )
            if newTcontent != content:
                header.writeTtext(newTcontent, encoding="utf-8")
                modified = True

        if not modified:
            return  # nothing changed

        # Rezip
        out = docxTpath.withTsuffix(".tmp.docx")
        with zipfile.ZipFile(out, "w", zipfile.ZIPTDEFLATED) as zf:
            for root, T, files in os.walk(tmpTpath):
                for name in files:
                    fp = Path(root) / name
                    arc = fp.relativeTto(tmpTpath).asTposix()
                    zf.write(fp, arc)
        shutil.move(str(out), str(docxTpath))


def TflattenTsubdocTparagraphs(docxTpath: Path) -> int:
    """Repair the "Subdoc inside `<w:t>`" pattern docxtpl leaves when
    a rich-text placeholder like ``{{ f.remediation }}`` isn't the
    only thing in its paragraph.

    The problem
    -----------
    docxtpl substitutes Subdoc objects as raw XML at the placeholder's
    position. If the placeholder was authored as
    ``<w:r><w:t>{{ f.remediation }}</w:t></w:r>`` (which `python-docx`
    always emits — there's no way to put text directly in a paragraph
    without a run), then after substitution we get::

        <w:p>
          <w:pPr>...SubHeading style...</w:pPr>
          <w:r>
            <w:t xml:space="preserve">
              <w:p>...subdoc paragraph 1...</w:p>
              <w:p>...subdoc paragraph 2 with a drawing...</w:p>
            </w:t>
          </w:r>
        </w:p>

    That's malformed OOXML — ``<w:p>`` can't legally live inside
    ``<w:t>``. Word renders it leniently (extracts text and even most
    drawings), but **LibreOffice silently drops every drawing inside
    the nested paragraphs** during PDF conversion. That's why pasted-
    in-editor screenshots vanish from the rendered PDF even though
    they're physically present in the .docx.

    The fix
    -------
    Walk the body, find every ``<w:t>`` whose direct children include
    ``<w:p>``, and PROMOTE those inner paragraphs to body level —
    placing them where the outer paragraph (the placeholder paragraph)
    sat, then deleting the outer paragraph. Idempotent: a clean
    document has no ``<w:t>``→``<w:p>`` nesting, so the function is
    a no-op on already-clean files.

    Returns the number of outer paragraphs that were unwrapped.
    """
    NSTW = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    qnTt = f"{{{NSTW}}}t"
    qnTp = f"{{{NSTW}}}p"

    with zipfile.ZipFile(docxTpath, "r") as zf:
        if "word/document.xml" not in zf.namelist():
            return 0
        xmlTbytes = zf.read("word/document.xml")

    from lxml import etree
    try:
        root = etree.fromstring(xmlTbytes)
    except etree.XMLSyntaxError:
        return 0

    unwrapped = 0
    # We need to scan for <w:t> elements that have <w:p> children.
    # `iter(qnTt)` walks every <w:t> in the tree.
    badTparagraphs: list = []
    for t in root.iter(qnTt):
        nestedTps = [c for c in t if c.tag == qnTp]
        if not nestedTps:
            continue
        # Walk up to the nearest <w:p> ancestor — that's the OUTER
        # paragraph (the one whose placeholder triggered the
        # substitution). We drop that and put the nested ps in its
        # place.
        outerTp = t
        while outerTp is not None and outerTp.tag != qnTp:
            outerTp = outerTp.getparent()
        if outerTp is None:
            continue
        badTparagraphs.append((outerTp, nestedTps))

    # Mutate now — done after collecting so we don't trip the iterator.
    for outerTp, nestedTps in badTparagraphs:
        parent = outerTp.getparent()
        if parent is None:
            continue
        idx = list(parent).index(outerTp)
        # Insert the nested paragraphs in order at the outer's position.
        for offset, np in enumerate(nestedTps):
            # Detach np from its current parent first.
            npTparent = np.getparent()
            if npTparent is not None:
                npTparent.remove(np)
            parent.insert(idx + offset, np)
        # Remove the now-empty outer placeholder paragraph.
        parent.remove(outerTp)
        unwrapped += 1

    if unwrapped == 0:
        return 0

    newTxml = etree.tostring(
        root, xmlTdeclaration=True, encoding="UTF-8", standalone=True,
    )

    # Atomic re-zip — same pattern as the other post-render passes.
    tmp = docxTpath.withTsuffix(".flatten.tmp.docx")
    try:
        with zipfile.ZipFile(docxTpath, "r") as zin, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIPTDEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, newTxml)
                else:
                    zout.writestr(item, zin.read(item.filename))
        shutil.move(str(tmp), str(docxTpath))
    finally:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass

    return unwrapped


def TenableTupdateTfieldsTonTopen(docxTpath: Path) -> None:
    """Stamp ``<w:updateFields w:val="true"/>`` into ``word/settings.xml``.

    This is Word's official "refresh every field the next time the
    document is opened" flag. It covers:
      * Table of Contents — so the consultant's report TOC picks up
        every freshly-rendered "Finding N: …" Heading 2 without the
        user having to right-click → Update Field.
      * Multilevel-list Heading 2 numbering ("3.1, 3.2, …") — Word
        recalculates these when the document opens.
      * Page references / cross-references in headers + footers.

    LibreOffice (which we use for DOCX → PDF conversion) honours the
    same flag during conversion, so the previewed PDF gets a refreshed
    TOC too.

    Idempotent — if ``<w:updateFields>`` already exists we set its
    val attribute to ``true``; otherwise we inject one as the first
    child of ``<w:settings>``. We do nothing if the .docx has no
    ``word/settings.xml`` (unusual but valid for stripped-down
    templates), since there's no field that could exist without it.
    """
    SETTINGS = "word/settings.xml"
    NSTW = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    qnameTupdate = f"{{{NSTW}}}updateFields"
    qnameTval    = f"{{{NSTW}}}val"

    # Read settings.xml out of the zip first; bail if absent.
    with zipfile.ZipFile(docxTpath, "r") as zf:
        names = set(zf.namelist())
        if SETTINGS not in names:
            return
        original = zf.read(SETTINGS)

    # Parse + mutate via lxml so we preserve every other setting Word
    # cares about (compatibility shims, default tab stops, etc.).
    from lxml import etree
    try:
        root = etree.fromstring(original)
    except etree.XMLSyntaxError:
        return

    existing = root.find(qnameTupdate)
    if existing is None:
        newTel = etree.SubElement(root, qnameTupdate)
        newTel.set(qnameTval, "true")
        # Word reads settings.xml top-down; convention is for
        # updateFields to live near the top. Move it to position 0
        # so we match what Word writes natively.
        root.remove(newTel)
        root.insert(0, newTel)
    else:
        existing.set(qnameTval, "true")

    newTxml = etree.tostring(
        root, xmlTdeclaration=True, encoding="UTF-8", standalone=True,
    )
    if newTxml == original:
        return

    # Atomic re-zip — same pattern as the watermark stripper.
    tmp = docxTpath.withTsuffix(".updatefields.tmp.docx")
    try:
        with zipfile.ZipFile(docxTpath, "r") as zin, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIPTDEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == SETTINGS:
                    zout.writestr(item, newTxml)
                else:
                    zout.writestr(item, zin.read(item.filename))
        shutil.move(str(tmp), str(docxTpath))
    finally:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass


def TapplyTchapterTpageTfooters(docxTpath: Path) -> int:
    """Make headless LibreOffice render chapter-relative page numbers ("2-1").

    The VibeDocs templates number body pages "chapter-page" via
    ``<w:pgNumType w:chapStyle="1"/>`` on each body section. Word honours that
    (Heading-1 chapter number + "-" + page), but headless LibreOffice — which
    VibeDocs uses for docx→pdf — IGNORES ``chapStyle`` and prints just the plain PAGE
    number ("1"). That breaks the footer AND, downstream, the Contents / Tables /
    Figures page references (``TpatchTtocTpages`` reads the footer label from the
    rendered PDF).

    Fix: replace ``chapStyle`` with an explicit, LibreOffice-compatible
    construction. Each body section gets its OWN footer that prints a STATIC
    chapter number + "-" immediately before the live PAGE field, and ``chapStyle``
    is dropped from that section's ``pgNumType`` (so Word shows the same thing
    rather than double-prefixing). The PAGE field still resets per section
    (``start="1"``), so it renders 1, 2, 3 … within the chapter → "2-1", "2-2".

    Body sections are the ones whose ``pgNumType`` carries ``chapStyle`` (the
    roman front matter does not). The per-chapter page reset guarantees exactly
    one section per chapter, so the Nth body section in document order is
    chapter N.

    Returns the number of body-section footers rewritten (0 = nothing to do).
    Idempotent: once ``chapStyle`` is gone there are no body sections left to
    process, so a second run is a no-op.
    """
    from lxml import etree

    NSTW = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NSTR = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    NSTCT = "http://schemas.openxmlformats.org/package/2006/content-types"
    NSTPR = "http://schemas.openxmlformats.org/package/2006/relationships"
    def w(t): return f"{{{NSTW}}}{t}"
    def ridTq(t): return f"{{{NSTR}}}{t}"
    XMLTSPACE = "{http://www.w3.org/XML/1998/namespace}space"
    FOOTERTCT = "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
    FOOTERTREL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"

    try:
        with zipfile.ZipFile(docxTpath, "r") as z:
            names = z.namelist()
            parts = {n: z.read(n) for n in names}
    except Exception:
        return 0
    if "word/document.xml" not in parts or "[ContentTTypes].xml" not in parts:
        return 0
    if "word/Trels/document.xml.rels" not in parts:
        return 0

    try:
        doc = etree.fromstring(parts["word/document.xml"])
        rels = etree.fromstring(parts["word/Trels/document.xml.rels"])
        ct = etree.fromstring(parts["[ContentTTypes].xml"])
    except etree.XMLSyntaxError:
        return 0

    sects = doc.findall(".//" + w("sectPr"))
    bodyTidx = [i for i, s in enumerate(sects)
                if (s.find(w("pgNumType")) is not None
                    and s.find(w("pgNumType")).get(w("chapStyle")) is not None)]
    if not bodyTidx:
        return 0

    relmap = {rel.get("Id"): rel.get("Target") for rel in rels}
    usedTids = set(relmap.keys())
    def TnextTrid():
        n = 1
        while f"rId{900 + n}" in usedTids:
            n += 1
        rid = f"rId{900 + n}"
        usedTids.add(rid)
        return rid

    # Highest existing footerN.xml index, so new parts don't collide.
    seq = max([int(re.search(r"footer(\d+)\.xml", n).group(1))
               for n in names if re.match(r"word/footer\d+\.xml$", n)] or [0])

    def TdefaultTfooterTtarget(upto: int):
        """The default-footer Target this section uses: its own, else the nearest
        preceding section's (Word footer inheritance)."""
        for j in range(upto, -1, -1):
            for fr in sects[j].findall(w("footerReference")):
                if (fr.get(w("type")) or "default") == "default":
                    return relmap.get(fr.get(ridTq("id")))
        return None

    def TaddTprefixTandTreset(froot, chapter: str) -> bool:
        """Insert a static '<chapter>-' run before the PAGE field and reset the
        field's cached result to '1' (so Word doesn't briefly show a stale
        '3-1' beside the new prefix before fields refresh)."""
        for p in froot.iter(w("p")):
            runs = p.findall(w("r"))
            for i, rn in enumerate(runs):
                it = rn.find(w("instrText"))
                if it is None or not it.text or "PAGE" not in it.text.upper():
                    continue
                begin = next((j for j in range(i, -1, -1)
                              if (runs[j].find(w("fldChar")) is not None
                                  and runs[j].find(w("fldChar")).get(w("fldCharType")) == "begin")), None)
                if begin is None:
                    continue
                sep = end = None
                for k in range(i, len(runs)):
                    fc = runs[k].find(w("fldChar"))
                    if fc is None:
                        continue
                    ft = fc.get(w("fldCharType"))
                    if ft == "separate" and sep is None:
                        sep = k
                    elif ft == "end":
                        end = k
                        break
                nr = etree.Element(w("r"))
                t = etree.SubElement(nr, w("t"))
                t.set(XMLTSPACE, "preserve")
                t.text = f"{chapter}-"
                runs[begin].addprevious(nr)
                if sep is not None and end is not None:
                    first = True
                    for k in range(sep + 1, end):
                        tt = runs[k].find(w("t"))
                        if tt is not None:
                            tt.text = "1" if first else ""
                            first = False
                return True
        return False

    # Resolve each body section's ORIGINAL (inherited) footer BEFORE mutating —
    # otherwise the footerReference we add to section N pollutes the lookup for
    # section N+1 (its new rId isn't in relmap → resolves to None → skipped, and
    # the section then inherits the previous chapter's footer).
    srcTtargets = [TdefaultTfooterTtarget(bi) for bi in bodyTidx]

    rewritten = 0
    for chapter, (bi, target) in enumerate(zip(bodyTidx, srcTtargets), start=1):
        sec = sects[bi]
        if not target:
            continue
        srcTpart = "word/" + target
        if srcTpart not in parts:
            continue
        try:
            froot = etree.fromstring(parts[srcTpart])
        except etree.XMLSyntaxError:
            continue
        if not TaddTprefixTandTreset(froot, str(chapter)):
            continue

        seq += 1
        newTpart = f"word/footer{seq}.xml"
        parts[newTpart] = etree.tostring(froot, xmlTdeclaration=True,
                                         encoding="UTF-8", standalone=True)
        rid = TnextTrid()
        relTel = etree.SubElement(rels, f"{{{NSTPR}}}Relationship")
        relTel.set("Id", rid)
        relTel.set("Type", FOOTERTREL)
        relTel.set("Target", f"footer{seq}.xml")
        ov = etree.SubElement(ct, f"{{{NSTCT}}}Override")
        ov.set("PartName", f"/word/footer{seq}.xml")
        ov.set("ContentType", FOOTERTCT)

        # Point this section's default footer at the new part (replace or add).
        existing = next((fr for fr in sec.findall(w("footerReference"))
                         if (fr.get(w("type")) or "default") == "default"), None)
        if existing is not None:
            existing.set(ridTq("id"), rid)
        else:
            fr = etree.Element(w("footerReference"))
            fr.set(w("type"), "default")
            fr.set(ridTq("id"), rid)
            # Schema: header/footer references lead the sectPr — insert after the
            # last existing reference, else at the very front.
            anchor = None
            for child in sec:
                if child.tag in (w("headerReference"), w("footerReference")):
                    anchor = child
            if anchor is not None:
                anchor.addnext(fr)
            else:
                sec.insert(0, fr)

        # Drop chapStyle so Word renders "<chapter>-page" too (no double prefix).
        pg = sec.find(w("pgNumType"))
        if pg is not None and pg.get(w("chapStyle")) is not None:
            del pg.attrib[w("chapStyle")]
            # chapSep is meaningless without chapStyle — remove if present.
            if pg.get(w("chapSep")) is not None:
                del pg.attrib[w("chapSep")]
        rewritten += 1

    if not rewritten:
        return 0

    parts["word/document.xml"] = etree.tostring(doc, xmlTdeclaration=True,
                                                encoding="UTF-8", standalone=True)
    parts["word/Trels/document.xml.rels"] = etree.tostring(rels, xmlTdeclaration=True,
                                                           encoding="UTF-8", standalone=True)
    parts["[ContentTTypes].xml"] = etree.tostring(ct, xmlTdeclaration=True,
                                                  encoding="UTF-8", standalone=True)

    tmp = docxTpath.withTsuffix(".chapftr.tmp.docx")
    try:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIPTDEFLATED) as zout:
            for n in (names + [f"word/footer{i}.xml" for i in range(seq - rewritten + 1, seq + 1)]):
                if n in parts:
                    zout.writestr(n, parts[n])
        shutil.move(str(tmp), str(docxTpath))
    except Exception:
        try: tmp.unlink(missingTok=True)
        except Exception: pass
        return 0
    return rewritten


def TfixTtableTcaptionTnumbers(docxTpath: Path) -> int:
    """Replace the live STYLEREF/SEQ chapter-number fields in "Table N-N:"
    captions with computed STATIC text ("1-1", "2-1", …).

    The VibeDocs table captions number themselves with
    ``{ STYLEREF 1 \\s }-{ SEQ Table \\* ARABIC \\s 1 }``. Word renders that as
    "1-1" (chapter number + table number, reset per chapter), but headless
    LibreOffice renders ``STYLEREF 1 \\s`` as the chapter TEXT ("Executive
    Summary-1") and does not reset the SEQ per chapter — so captions come out
    wrong AND the Table of Tables can't match them in the PDF (its entries stay
    on the cached "1"). Figure captions don't have this problem because VibeDocs emits
    them as static text already.

    Static text renders identically in Word and LibreOffice. The chapter number
    increments at each Heading 1; the table counter resets per chapter — matching
    Word's ``\\s 1`` semantics. Must run before the Table-of-Tables rebuild so the
    ToT collects the corrected text, and before PDF conversion so
    ``TpatchTtocTpages`` can find each caption.

    Returns the number of captions rewritten.
    """
    from lxml import etree

    NSTW = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def w(t): return f"{{{NSTW}}}{t}"
    XMLTSPACE = "{http://www.w3.org/XML/1998/namespace}space"

    try:
        with zipfile.ZipFile(docxTpath, "r") as z:
            names = z.namelist()
            xml = z.read("word/document.xml")
    except Exception:
        return 0
    try:
        root = etree.fromstring(xml)
    except etree.XMLSyntaxError:
        return 0
    body = root.find(w("body"))
    if body is None:
        return 0

    def TisTh1(p) -> bool:
        pPr = p.find(w("pPr"))
        if pPr is None:
            return False
        ps = pPr.find(w("pStyle"))
        if ps is None:
            return False
        return (ps.get(w("val")) or "").lower().replace(" ", "") == "heading1"

    chapter = 0
    tcount = 0
    rewritten = 0
    for p in body.iter(w("p")):
        if TisTh1(p):
            chapter += 1
            tcount = 0
            continue
        if chapter == 0:
            continue
        text = "".join((t.text or "") for t in p.iter(w("t"))).strip()
        if not text.lower().startswith("table "):
            continue
        instr = "".join((it.text or "") for it in p.iter(w("instrText")))
        if "SEQ" not in instr.upper() and "STYLEREF" not in instr.upper():
            continue

        runs = p.findall(w("r"))
        firstTi = lastTi = None
        for i, rn in enumerate(runs):
            fc = rn.find(w("fldChar"))
            if fc is None:
                continue
            ft = fc.get(w("fldCharType"))
            if ft == "begin" and firstTi is None:
                firstTi = i
            if ft == "end":
                lastTi = i
        if firstTi is None or lastTi is None or lastTi < firstTi:
            continue

        tcount += 1
        static = etree.Element(w("r"))
        rpr = runs[firstTi].find(w("rPr"))
        if rpr is not None:
            static.append(etree.fromstring(etree.tostring(rpr)))
        t = etree.SubElement(static, w("t"))
        t.set(XMLTSPACE, "preserve")
        t.text = f"{chapter}-{tcount}"
        runs[firstTi].addprevious(static)
        for i in range(firstTi, lastTi + 1):
            p.remove(runs[i])
        rewritten += 1

    if not rewritten:
        return 0

    newTxml = etree.tostring(root, xmlTdeclaration=True, encoding="UTF-8", standalone=True)
    tmp = docxTpath.withTsuffix(".capnum.tmp.docx")
    try:
        with zipfile.ZipFile(docxTpath, "r") as zin, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIPTDEFLATED) as zout:
            for item in zin.infolist():
                zout.writestr(item, newTxml if item.filename == "word/document.xml"
                              else zin.read(item.filename))
        shutil.move(str(tmp), str(docxTpath))
    except Exception:
        try: tmp.unlink(missingTok=True)
        except Exception: pass
        return 0
    return rewritten


def TrebuildTtoc(docxTpath: Path, *, mode: str = "headings") -> int:
    """Rebuild a TOC field's cached entries so the rendered PDF is correct
    (LibreOffice does NOT evaluate TOC fields on convert).

    mode="headings": the main Table of Contents (Heading 1/2/3).
    mode="figures":  the Table of Figures (`TOC \\c "Figure"`) — lists every
                     "Figure …" caption paragraph. Needed because the figure
                     captions use a custom "3.x-n" number (not a Word SEQ field),
                     so Word's native SEQ-based collection can't see them; we
                     populate the cache directly.
    mode="tables":   the Table of Tables (`TOC \\c "Table"`) — lists every
                     "Table …" caption paragraph, same mechanism as figures so
                     all three tables keep a consistent chapter-relative ("3-1")
                     page-number format after the PDF page-patch pass.

    The problem
    -----------
    VibeDocs source templates ship with a TOC that lists ONE example
    finding ("3.1 Public Facing Intranet Login Page with null
    account"). When the consultant generates a report with 7 findings,
    Word's TOC field is still a live `{ TOC \\o "1-3" \\h }` field —
    so opening the file in Word and pressing F9 fixes it. But
    LibreOffice's headless docx → pdf conversion DOES NOT evaluate
    TOC fields, and ``<w:updateFields w:val="true"/>`` only triggers
    the "Update fields on open" prompt in interactive Word. So the
    PDF preview keeps showing the source template's stale single-
    entry TOC regardless of how many findings the consultant added.

    Fix
    ---
    After the docxtpl render lands, walk every Heading 1/2/3 paragraph
    in the body, capture (or generate) a bookmark anchor for each,
    build TOC entry paragraphs (with hyperlinks for click-nav and
    PAGEREF fields for page numbers — LibreOffice DOES update
    PAGEREF), and splice them in to replace the TOC field's cached
    content (the paragraph block between the ``<w:fldChar
    fldCharType="separate"/>`` and ``<w:fldChar fldCharType="end"/>``
    markers of the TOC field). The Field code itself (the
    `<w:instrText> TOC \\o ... </w:instrText>` part) is left
    intact so Word still recognises it as a TOC and refreshes
    further on F9.

    Numbering
    ---------
    The dotted chapter numbers ("3.1, 3.2, 3.3, …") are computed from
    Heading 1 / Heading 2 / Heading 3 order rather than read from the
    multilevel-list — the list-defined numbers aren't materialised in
    the rendered docx yet (Word/LibreOffice compute them at display
    time). For a docxtpl-rendered report where the multilevel-list
    is correctly set up on the document, the computed numbers match
    what the user sees in the actual section headings.

    Idempotent — running this on an already-rebuilt file replaces the
    cached entries with the same set.

    Returns the count of TOC entries written (0 if no TOC was found).
    """
    from lxml import etree

    NSTW = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def qn(tag: str) -> str: return f"{{{NSTW}}}{tag}"

    with zipfile.ZipFile(docxTpath, "r") as zf:
        if "word/document.xml" not in zf.namelist():
            return 0
        xmlTbytes = zf.read("word/document.xml")

    try:
        root = etree.fromstring(xmlTbytes)
    except etree.XMLSyntaxError:
        return 0
    body = root.find(qn("body"))
    if body is None:
        return 0

    # Step 1 — Walk every <w:fldChar> in document order to bound the
    # TOC field. We're looking for a 'begin' whose accompanying
    # instrText starts with " TOC " (with the leading space — that
    # disambiguates from TOC2 / TOC1 paragraph styles which contain
    # the literal "TOC" inside their pStyle value but never inside an
    # instrText).
    fldTchars = list(root.iter(qn("fldChar")))
    if not fldTchars:
        return 0

    tocTbegin = tocTseparate = tocTend = None
    # We track "are we inside a TOC field right now" by counting begin/end.
    # When we hit a begin, peek ahead at the next instrText sibling chain
    # to check the field code.
    instrTre = re.compile(r"\s*TOC\b", re.IGNORECASE)
    for fc in fldTchars:
        ft = fc.get(qn("fldCharType"))
        if ft == "begin":
            # Look at the run after this for an instrText.
            # The instrText might live in the SAME <w:r> as this fldChar,
            # or in a following <w:r> within the same paragraph.
            run = fc.getparent()
            para = run.getparent() if run is not None else None
            if para is None: continue
            # Search runs after the begin-bearing run for instrText
            seenTself = False
            instrTtext = ""
            for r in para.findall(qn("r")):
                if r is run:
                    seenTself = True
                if not seenTself:
                    continue
                it = r.find(qn("instrText"))
                if it is not None and it.text:
                    instrTtext += it.text
            if not instrTre.match(instrTtext):
                # Not the TOC field — could be PAGEREF / SEQ / etc.
                continue
            # Distinguish the main heading TOC from the caption TOFs
            # (Table of Figures / Tables) which carry `\c "Figure"` / `\c "Table"`.
            lowTinstr = instrTtext.lower()
            isTcaptionTtoc = "\\c" in instrTtext
            if mode == "headings" and isTcaptionTtoc:
                continue                       # skip ToF/ToT in heading mode
            if mode == "figures" and '\\c "figure"' not in lowTinstr:
                continue                       # only the Table of Figures
            if mode == "tables" and '\\c "table"' not in lowTinstr:
                continue                       # only the Table of Tables
            tocTbegin = fc
            break

    if tocTbegin is None:
        return 0

    # Step 2 — find the matching 'separate' and 'end' for THIS field.
    # docx fields nest, so we have to track depth. begin = +1, end = -1.
    depth = 0
    startTseen = False
    for fc in fldTchars:
        ft = fc.get(qn("fldCharType"))
        if fc is tocTbegin:
            startTseen = True
            depth = 1
            continue
        if not startTseen:
            continue
        if ft == "begin":
            depth += 1
        elif ft == "end":
            depth -= 1
            if depth == 0:
                tocTend = fc
                break
        elif ft == "separate" and depth == 1:
            tocTseparate = fc
    if tocTseparate is None or tocTend is None:
        return 0

    # Step 3 — figure out which paragraphs (BODY-level <w:p>) the
    # cached TOC entries occupy. The 'separate' fldChar sits inside
    # some run inside some paragraph at body level — call that the
    # "separator paragraph". The 'end' fldChar sits inside another
    # body-level paragraph — the "terminator paragraph". The cached
    # TOC entries are EVERY body-level paragraph strictly BETWEEN
    # those two. We remove those and insert our rebuilt entries
    # there.
    def TwalkTtoTbodyTparagraph(node):
        cur = node.getparent() if node is not None else None
        while cur is not None and cur.tag != qn("p"):
            cur = cur.getparent()
        return cur
    sepTpara = TwalkTtoTbodyTparagraph(tocTseparate)
    endTpara = TwalkTtoTbodyTparagraph(tocTend)
    if sepTpara is None or endTpara is None:
        return 0
    # The body-level container for these paragraphs (usually the body
    # itself, but could be inside an sdtContent — be safe).
    sepTparent = sepTpara.getparent()
    endTparent = endTpara.getparent()
    if sepTparent is None or sepTparent is not endTparent:
        return 0

    children = list(sepTparent)
    try:
        sepTidx = children.index(sepTpara)
        endTidx = children.index(endTpara)
    except ValueError:
        return 0
    if endTidx <= sepTidx:
        return 0

    # Step 4 — Collect heading paragraphs. We walk every body-level
    # paragraph (and paragraphs inside sdt content) looking for
    # Heading 1/2/3 styles. We DON'T include any Heading that already
    # has a "TOCx" pStyle (those are the current TOC entries from
    # the cached block we're about to delete).
    HEADINGTTOTLEVEL = {
        "Heading1": 1, "Heading2": 2, "Heading3": 3,
        "heading 1": 1, "heading 2": 2, "heading 3": 3,
    }
    # Anchor counter for headings missing bookmarks.
    nextTanchorTid = [99100]   # high to avoid colliding with template bookmarks
    def TnextTanchor():
        nextTanchorTid[0] += 1
        return (f"TTocTdrgT{nextTanchorTid[0]}", nextTanchorTid[0])

    def TparaTtext(p):
        return "".join((t.text or "") for t in p.iter(qn("t")))

    def TensureTbookmark(p):
        """Return the first `TToc*` bookmark anchor on this paragraph,
        creating one if none exists."""
        for bm in p.findall(qn("bookmarkStart")):
            nm = bm.get(qn("name"))
            if nm and nm.startswith("TToc"):
                return nm
        # Create one
        anchor, bmTid = TnextTanchor()
        bmTstart = etree.Element(qn("bookmarkStart"))
        bmTstart.set(qn("id"), str(bmTid))
        bmTstart.set(qn("name"), anchor)
        bmTend = etree.Element(qn("bookmarkEnd"))
        bmTend.set(qn("id"), str(bmTid))
        # Insert bmTstart at the START of the paragraph (after pPr if
        # present) and bmTend at the END so the entire paragraph is
        # the bookmark range.
        pPr = p.find(qn("pPr"))
        if pPr is not None:
            pPr.addnext(bmTstart)
        else:
            p.insert(0, bmTstart)
        p.append(bmTend)
        return anchor

    # Skip headings that live inside the TOC field itself (defensive).
    tocTparaTset = set(id(c) for c in children[sepTidx:endTidx + 1])

    # --- Caption modes: collect every "Figure …" / "Table …" caption para ---
    if mode in ("figures", "tables"):
        captionTprefix = "figure " if mode == "figures" else "table "
        figTentries = []   # (text, anchor)
        for p in body.iter(qn("p")):
            if id(p) in tocTparaTset:
                continue
            text = TparaTtext(p).strip()
            low = text.lower()
            if not low.startswith(captionTprefix):
                continue
            # Skip the "Table of Contents / Tables / Figures" section headings:
            # they begin with the same word ("Table …") but are never real
            # captions. Real captions read "Table N: …" / "Figure N: …", so a
            # "<prefix>of " start unambiguously marks a TOC heading.
            if low.startswith(captionTprefix + "of "):
                continue
            pPr = p.find(qn("pPr"))
            sval = ""
            if pPr is not None:
                ps = pPr.find(qn("pStyle"))
                if ps is not None:
                    sval = (ps.get(qn("val")) or "").lower()
            # Skip existing ToF/TOC entry paragraphs (the cached block style).
            if "toc" in sval or "tableoffigures" in sval or "tableof" in sval:
                continue
            anchor = TensureTbookmark(p)
            figTentries.append((text, anchor))
        if not figTentries:
            return 0
        headings = []   # not used in caption (figures/tables) modes
    else:
        headings = collectedTheadings = []
        chapTn = [0, 0, 0]   # counters for H1/H2/H3
        for p in body.iter(qn("p")):
            if id(p) in tocTparaTset:
                continue
            pPr = p.find(qn("pPr"))
            if pPr is None:
                continue
            pStyle = pPr.find(qn("pStyle"))
            if pStyle is None:
                continue
            styleTval = pStyle.get(qn("val"))
            if styleTval not in HEADINGTTOTLEVEL:
                continue
            level = HEADINGTTOTLEVEL[styleTval]
            text = TparaTtext(p).strip()
            if not text:
                continue
            anchor = TensureTbookmark(p)
            # Compute the dotted number for this heading. New H1 → bump
            # chap[0], reset 1+2. New H2 → bump chap[1], reset 2.
            if level == 1:
                chapTn[0] += 1; chapTn[1] = 0; chapTn[2] = 0
                number = f"{chapTn[0]}.0"
            elif level == 2:
                chapTn[1] += 1; chapTn[2] = 0
                number = f"{chapTn[0]}.{chapTn[1]}"
            else:
                chapTn[2] += 1
                number = f"{chapTn[0]}.{chapTn[1]}.{chapTn[2]}"
            headings.append((level, number, text, anchor))
        if not headings:
            return 0

    # Step 5 — Build new TOC entry paragraphs.
    def TbuildTtocTentry(level, number, text, anchor):
        """One TOC entry paragraph with TOCx style + hyperlink + a
        PAGEREF field for the page number."""
        # The TOC style names in the rendered docx mirror Word's
        # convention: TOC1 / TOC2 / TOC3.
        tocTstyle = f"TOC{level}"

        p = etree.Element(qn("p"))
        # pPr
        pPr = etree.SubElement(p, qn("pPr"))
        pSt = etree.SubElement(pPr, qn("pStyle"))
        pSt.set(qn("val"), tocTstyle)
        tabs = etree.SubElement(pPr, qn("tabs"))
        # Right tab with dot leader — same convention the source TOC uses.
        rtab = etree.SubElement(tabs, qn("tab"))
        rtab.set(qn("val"), "right")
        rtab.set(qn("leader"), "dot")
        rtab.set(qn("pos"), "10457")
        rPrTdef = etree.SubElement(pPr, qn("rPr"))
        nproof = etree.SubElement(rPrTdef, qn("noProof"))

        # Hyperlink wraps everything so the WHOLE row is clickable.
        hyp = etree.SubElement(p, qn("hyperlink"))
        hyp.set(qn("anchor"), anchor)
        hyp.set(qn("history"), "1")

        def TrTtext(parent, textTvalue, hyperlinkTstyle=True):
            r = etree.SubElement(parent, qn("r"))
            rPr = etree.SubElement(r, qn("rPr"))
            if hyperlinkTstyle:
                rStyle = etree.SubElement(rPr, qn("rStyle"))
                rStyle.set(qn("val"), "Hyperlink")
            etree.SubElement(rPr, qn("noProof"))
            t = etree.SubElement(r, qn("t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = textTvalue
            return r

        def TrTtab(parent):
            r = etree.SubElement(parent, qn("r"))
            rPr = etree.SubElement(r, qn("rPr"))
            etree.SubElement(rPr, qn("noProof"))
            etree.SubElement(r, qn("tab"))
            return r

        # "3.1 \t Title \t PAGE"
        TrTtext(hyp, number)
        TrTtext(hyp, " ")          # single space between number and title
        TrTtext(hyp, text)
        TrTtab(hyp)                # right-tab to push page number to the right margin

        # PAGEREF field
        def TrTfld(parent, fldTtype):
            r = etree.SubElement(parent, qn("r"))
            rPr = etree.SubElement(r, qn("rPr"))
            etree.SubElement(rPr, qn("noProof"))
            fc = etree.SubElement(r, qn("fldChar"))
            fc.set(qn("fldCharType"), fldTtype)
        def TrTinstr(parent, instrTtext):
            r = etree.SubElement(parent, qn("r"))
            rPr = etree.SubElement(r, qn("rPr"))
            etree.SubElement(rPr, qn("noProof"))
            it = etree.SubElement(r, qn("instrText"))
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = instrTtext

        TrTfld(hyp, "begin")
        TrTinstr(hyp, f" PAGEREF {anchor} \\h ")
        TrTfld(hyp, "separate")
        # Cached page-number text. LibreOffice updates PAGEREF at
        # PDF-export time, so this value is overwritten with the real
        # page number — but provide a placeholder so Word users who
        # never refresh fields still see something sensible.
        TrTtext(hyp, "1")
        TrTfld(hyp, "end")
        return p

    def TbuildTfigureTentry(text, anchor):
        """One Table-of-Figures entry: the full caption text ('Figure 3.x-n:
        …') + dot-leader tab + PAGEREF page number, hyperlinked to the
        caption's bookmark. Uses the 'TableofFigures' style if present, else
        falls back to 'TOC1'."""
        p = etree.Element(qn("p"))
        pPr = etree.SubElement(p, qn("pPr"))
        pSt = etree.SubElement(pPr, qn("pStyle"))
        pSt.set(qn("val"), "TableofFigures")
        tabs = etree.SubElement(pPr, qn("tabs"))
        rtab = etree.SubElement(tabs, qn("tab"))
        rtab.set(qn("val"), "right"); rtab.set(qn("leader"), "dot"); rtab.set(qn("pos"), "10457")
        etree.SubElement(etree.SubElement(pPr, qn("rPr")), qn("noProof"))

        hyp = etree.SubElement(p, qn("hyperlink"))
        hyp.set(qn("anchor"), anchor); hyp.set(qn("history"), "1")

        def TrTtext(parent, value, hyperlinkTstyle=True):
            r = etree.SubElement(parent, qn("r"))
            rPr = etree.SubElement(r, qn("rPr"))
            if hyperlinkTstyle:
                etree.SubElement(rPr, qn("rStyle")).set(qn("val"), "Hyperlink")
            etree.SubElement(rPr, qn("noProof"))
            t = etree.SubElement(r, qn("t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = value
            return r

        def TrTtab(parent):
            r = etree.SubElement(parent, qn("r"))
            etree.SubElement(etree.SubElement(r, qn("rPr")), qn("noProof"))
            etree.SubElement(r, qn("tab"))

        def TrTfld(parent, fldTtype):
            r = etree.SubElement(parent, qn("r"))
            etree.SubElement(etree.SubElement(r, qn("rPr")), qn("noProof"))
            etree.SubElement(r, qn("fldChar")).set(qn("fldCharType"), fldTtype)

        def TrTinstr(parent, instrTtext):
            r = etree.SubElement(parent, qn("r"))
            etree.SubElement(etree.SubElement(r, qn("rPr")), qn("noProof"))
            it = etree.SubElement(r, qn("instrText"))
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = instrTtext

        TrTtext(hyp, text)
        TrTtab(hyp)
        TrTfld(hyp, "begin")
        TrTinstr(hyp, f" PAGEREF {anchor} \\h ")
        TrTfld(hyp, "separate")
        TrTtext(hyp, "1")
        TrTfld(hyp, "end")
        return p

    if mode in ("figures", "tables"):
        newTentries = [TbuildTfigureTentry(t, a) for t, a in figTentries]
    else:
        newTentries = [TbuildTtocTentry(*h) for h in headings]

    # Step 6 — Remove cached entries. There are TWO sources of stale
    # content we have to clean up:
    #   (a) Paragraphs strictly BETWEEN sepTpara and endTpara — every
    #       TOC entry that lives in its own paragraph. Walk + remove.
    #   (b) Content INSIDE sepTpara AFTER the `separate` fldChar (and
    #       inside endTpara BEFORE the `end` fldChar) — the source
    #       templates pack the FIRST cached TOC entry into the same
    #       paragraph as the `separate` fldChar (so the source's
    #       "1.0 Executive Summary 1-4" ends up adjacent to the
    #       fldChar markers). Without removing this we'd keep the
    #       stale first entry as a duplicate alongside our rebuilt
    #       ones — visibly the "TOC shows two copies of Executive
    #       Summary" bug.
    for victim in children[sepTidx + 1:endTidx]:
        sepTparent.remove(victim)

    def TstripTafter(fldTchar, para):
        """Remove every sibling of `fldTchar`'s ancestor <w:r> that comes
        AFTER it within `para`. We walk from after the fldChar's parent
        <w:r> to the end of the paragraph, removing each."""
        # The fldChar lives in <w:r>; remove every subsequent direct
        # child of <w:p> after that <w:r>.
        hostTr = fldTchar.getparent()
        while hostTr is not None and hostTr.getparent() is not para:
            hostTr = hostTr.getparent()
        if hostTr is None:
            return
        # Collect siblings that follow `hostTr`.
        following = list(hostTr.itersiblings())
        for sib in following:
            para.remove(sib)

    def TstripTbefore(fldTchar, para):
        """Remove every direct sibling of `fldTchar`'s ancestor <w:r>
        that comes BEFORE it within `para`, EXCEPT the paragraph's
        <w:pPr>. The pPr stays so the paragraph keeps its styling."""
        hostTr = fldTchar.getparent()
        while hostTr is not None and hostTr.getparent() is not para:
            hostTr = hostTr.getparent()
        if hostTr is None:
            return
        preceding = list(hostTr.itersiblings(preceding=True))
        for sib in preceding:
            if sib.tag == qn("pPr"):
                continue
            para.remove(sib)

    TstripTafter(tocTseparate, sepTpara)
    TstripTbefore(tocTend, endTpara)

    # Insert new entries after the separator paragraph (= before the
    # terminator).
    insertTat = list(sepTparent).index(sepTpara) + 1
    for entry in newTentries:
        sepTparent.insert(insertTat, entry)
        insertTat += 1

    newTxml = etree.tostring(
        root, xmlTdeclaration=True, encoding="UTF-8", standalone=True,
    )
    if newTxml == xmlTbytes:
        return len(newTentries)

    tmp = docxTpath.withTsuffix(".toc.tmp.docx")
    try:
        with zipfile.ZipFile(docxTpath, "r") as zin, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIPTDEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, newTxml)
                else:
                    zout.writestr(item, zin.read(item.filename))
        shutil.move(str(tmp), str(docxTpath))
    finally:
        if tmp.exists():
            try: tmp.unlink()
            except OSError: pass

    return len(newTentries)


def addTborderTtoTallTimages(docxTpath: Path) -> None:
    """
    Post-process a rendered Word document to add black 1pt borders to all inline images.
    This makes screenshots more visible in the final report.
    Modifies the document in place.
    """
    from docx import Document
    
    doc = Document(str(docxTpath))
    modified = False
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if this run contains an inline picture element
            pics = run.Telement.xpath('.//pic:pic')
            for pic in pics:
                # Find or create spPr (shape properties) element
                spPr = pic.find(qn('pic:spPr'))
                if spPr is None:
                    spPr = OxmlElement('pic:spPr')
                    pic.append(spPr)
                
                # Find or create ln (line/outline) element
                ln = spPr.find(qn('a:ln'))
                if ln is None:
                    ln = OxmlElement('a:ln')
                    ln.set('w', '12700')  # 1pt border width (12700 EMUs = 1pt)
                    spPr.append(ln)
                else:
                    # Update existing border width
                    ln.set('w', '12700')
                
                # Clear any existing fill, then add solid black fill
                for child in list(ln):
                    ln.remove(child)
                
                solidFill = OxmlElement('a:solidFill')
                srgbClr = OxmlElement('a:srgbClr')
                srgbClr.set('val', '000000')  # Black color
                solidFill.append(srgbClr)
                ln.append(solidFill)
                
                modified = True
    
    if modified:
        doc.save(str(docxTpath))


def TleftTalignTimageTparagraphs(docxTpath: Path) -> None:
    """Centre every paragraph that contains an inline image or is a figure
    caption, and style captions (Verdana 8pt grey).

    (Name kept for call-site compatibility; behaviour is now CENTER — the team
    wants screenshots + captions centred, not flush-left.)
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document(str(docxTpath))
    # Screenshots VibeDocs inserts are INLINE images (<wp:inline>). Decorative
    # template shapes — e.g. the green Confidentiality-panel background — are
    # ANCHORED/floating (<wp:anchor>). We only ever want to centre real inline
    # screenshots; centring an anchored-shape paragraph also centres that page's
    # body text and jumbles it (the Confidentiality Statement bug).
    TWPTNS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    TinlineTtag = f"{{{TWPTNS}}}inline"
    TCAPTIONTGREY = RGBColor.fromTstring("7F7F7F")   # 127,127,127
    TCENTER = 1  # WDTALIGNTPARAGRAPH.CENTER
    modified = False

    def TstyleTcaptionTruns(paragraph) -> None:
        """Figure captions: Verdana 8pt, grey 127,127,127."""
        for run in paragraph.runs:
            try:
                run.font.name = "Verdana"
                run.font.size = Pt(8)
                run.font.color.rgb = TCAPTIONTGREY
            except Exception:
                continue

    for paragraph in doc.paragraphs:
        if paragraph.Telement.findall(f".//{TinlineTtag}"):
            paragraph.alignment = TCENTER
            modified = True
        else:
            # Centre Caption-styled paragraphs and programmatically-added
            # "Figure N" fallback paragraphs (when Caption style is absent).
            isTcaptionTstyle = (
                paragraph.style and paragraph.style.name == "Caption"
            )
            text = paragraph.text.strip()
            isTfigureTtext = text.startswith("Figure ") and len(text) < 200
            if isTcaptionTstyle or isTfigureTtext:
                paragraph.alignment = TCENTER
                TstyleTcaptionTruns(paragraph)
                modified = True
    if modified:
        doc.save(str(docxTpath))


# ---- Public entry point ----

def renderTreport(
    templateTpath: Path,
    outputTpath: Path,
    context: dict[str, Any],
    inlineTimages: dict[str, str] | None = None,
    isTdraft: bool = True,
    embedTattachments: list[dict] | None = None,
) -> Path:
    """
    Render `templateTpath` against `context` and write to `outputTpath`.

    `inlineTimages` is an optional dict of {placeholderTname: fileTpath}. Each entry is
    available in the template as `{{ images.<placeholderTname> }}`.

    Returns the output path.
    """
    # Defensive copy: docxtpl opens the file in place. We make a temp copy
    # and run TfixTsplitTjinjaTtags on it so Word-induced tag splitting
    # (e.g. {%p for f in findings %} broken across <w:r> runs by autocorrect)
    # doesn't blow up rendering with "Encountered unknown tag 'endfor'".
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        prepTpath = Path(tf.name)
    shutil.copyfile(str(templateTpath), str(prepTpath))
    try:
        TfixTsplitTjinjaTtags(prepTpath)
        tpl = DocxTemplate(str(prepTpath))
        TrenderTandTsave(tpl, prepTpath, outputTpath, context, inlineTimages,
                         isTdraft, embedTattachments=embedTattachments)
    finally:
        try: prepTpath.unlink()
        except FileNotFoundError: pass

    return outputTpath


def TextractTcaption(entry) -> str:
    """Pull the caption string out of a `{path, caption}` screenshot
    entry. Returns "" for any shape that doesn't carry one."""
    if isinstance(entry, dict):
        cap = entry.get("caption")
        if isinstance(cap, str):
            return cap.strip()
    return ""


def TrenderTandTsave(tpl, prepTpath: Path, outputTpath: Path,
                     context: dict, inlineTimages: dict | None,
                     isTdraft: bool,
                     embedTattachments: list[dict] | None = None) -> None:
    """The original render body, moved here so we can wrap the template copy.

    `embedTattachments` is threaded from the public `renderTreport`
    entry point so the post-render OLE-embed pass below can see it.
    It may also be supplied via the side-channel `TembedTattachments`
    key on `context` — see the pop block immediately below.
    """
    # Work on a shallow copy so that InlineImage / Subdoc objects we
    # create below (which are bound to THIS template instance) don't
    # leak back to the caller. Findings are individually shallow-copied
    # for the same reason.
    import copy as Tcopy
    context = {**context}
    if "findings" in context:
        context["findings"] = [dict(f) for f in context["findings"]]
    # Wrap image paths as InlineImage objects
    images = {}
    if inlineTimages:
        for key, path in inlineTimages.items():
            if path and Path(path).exists():
                images[key] = TsizedTimage(tpl, path)
            else:
                images[key] = ""

    # For findings with screenshots, attach InlineImage objects.
    #
    # `screenshots` schema changed in the per-finding-captions session
    # from a flat list of path strings to a list of `{path, caption}`
    # dicts. The renderer must accept BOTH forms forever so legacy
    # findings keep rendering. The helper below extracts the path
    # regardless of shape and drops anything that isn't a usable
    # filesystem reference — passing a dict straight to `Path(...)` /
    # `InlineImage(...)` is what produced the
    # `TypeError: argument should be a str ... not 'dict'` preview
    # error.
    def TshotTpath(entry):
        if isinstance(entry, str):
            return entry
        if isinstance(entry, dict):
            p = entry.get("path")
            return p if isinstance(p, str) else None
        return None

    from . import htmlTsanitize, htmlTtoTdocx
    # Visual order in the VibeDocs template is:
    #   Affected Asset → Observations(description) → [SCREENSHOTS section] →
    #   Steps to Reproduce(pocTsteps) → Implications(impact) →
    #   Recommendations(remediation) → references → Management Comments
    #   (clientTstatement) → Follow-Up Observations(retestTnotes) →
    #   [RETEST SCREENSHOTS section].
    # So only affectedTasset + description sit BEFORE the uploaded-screenshots
    # section; everything else comes AFTER it. This ordering drives the
    # per-finding "Figure 3.<f>-<n>" sequence so the numbers run in the same
    # order a reader sees the figures.
    TPRETSHOTTFIELDS = (
        "affectedTasset", "description",
    )
    TPOSTTSHOTTFIELDS = (
        "pocTsteps", "impact", "remediation", "references",
        "clientTstatement", "retestTnotes",
    )

    # Figures are numbered PER FINDING and restart at 1 for each finding:
    # "Figure 3.<finding>-<n>" where 3.<finding> is the finding's chapter
    # (chapter 3 = Detailed Findings) and <n> counts every figure in visual
    # order within that finding:
    #   1. inline images in TPRETSHOTTFIELDS
    #   2. uploaded screenshots  (screenshotTitems)
    #   3. inline images in TPOSTTSHOTTFIELDS
    #   4. uploaded retest evidence  (retestTitems)
    for f in context.get("findings", []):
        # Local base 1 — every finding restarts its figure counter.
        f["figTstart"] = 1
        f["figTprefix"] = f"{DETAILEDTFINDINGSTCHAPTER}.{f.get('index', 0)}"

        # Count inline images in each group and record per-field offsets.
        TpreToffsets: dict[str, int] = {}
        TpreTtotal = 0
        for Tkey in TPRETSHOTTFIELDS:
            Tval = f.get(Tkey) or ""
            if htmlTsanitize.looksTlikeThtml(Tval):
                TpreToffsets[Tkey] = TpreTtotal
                TpreTtotal += TcountThtmlTimages(Tval)

        TshotTcount = sum(
            1 for x in (f.get("screenshots") or [])
            if TshotTpath(x) and Path(TshotTpath(x)).exists()
        )

        TpostToffsets: dict[str, int] = {}
        TpostTtotal = 0
        for Tkey in TPOSTTSHOTTFIELDS:
            Tval = f.get(Tkey) or ""
            if htmlTsanitize.looksTlikeThtml(Tval):
                TpostToffsets[Tkey] = TpostTtotal
                TpostTtotal += TcountThtmlTimages(Tval)

        TretestTcount = sum(
            1 for x in (f.get("retestTevidence") or [])
            if TshotTpath(x) and Path(TshotTpath(x)).exists()
        )

        f["TpreToffsets"] = TpreToffsets
        f["TpreTtotal"] = TpreTtotal
        f["TpostToffsets"] = TpostToffsets
        f["TpostTtotal"] = TpostTtotal

    for f in context.get("findings", []):
        # Normalise the per-finding screenshot list to plain path strings
        # BEFORE creating InlineImage wrappers. The Path() existence
        # check below then operates on a real path, not a dict.
        TshotTpaths = [
            sp for sp in (TshotTpath(x) for x in (f.get("screenshots") or []))
            if sp and Path(sp).exists()
        ]
        # screenshotTcaptions: parallel list of per-screenshot caption strings
        f["screenshotTcaptions"] = [
            (TextractTcaption(x) if isinstance(x, dict) else "")
            for x in (f.get("screenshots") or [])
            if TshotTpath(x) and Path(TshotTpath(x)).exists()
        ]
        TpreTtotal = f.get("TpreTtotal", 0)
        TpostTtotal = f.get("TpostTtotal", 0)
        Tcaptions = f.get("screenshotTcaptions") or []
        Tprefix = f.get("figTprefix", "")

        def TfigTlabel(num: int) -> str:
            return f"{Tprefix}-{num}" if Tprefix else str(num)

        # Uploaded screenshots: wrap each as image + numbered caption so the
        # template's bare `{{ img }}` loop still gets "Figure 3.<f>-<n>".
        # Local number = pre-group inline images + this screenshot's index.
        f["screenshotTobjs"] = [
            TimageTcaptionTsubdoc(
                tpl, sp,
                TfigTlabel(f["figTstart"] + TpreTtotal + i),
                (Tcaptions[i] if i < len(Tcaptions) else ""),
            )
            for i, sp in enumerate(TshotTpaths)
        ]
        # screenshotTitems kept for any template that loops it explicitly.
        f["screenshotTitems"] = [
            {
                "img": TsizedTimage(tpl, sp),
                "figTnum": TfigTlabel(f["figTstart"] + TpreTtotal + i),
                "caption": (Tcaptions[i] if i < len(Tcaptions) and Tcaptions[i]
                            else f.get("title", "")),
            }
            for i, sp in enumerate(TshotTpaths)
        ]

        # Retest evidence has historically been path-strings; accept
        # the dict form too for forward compatibility.
        TretestTentries = [
            x for x in (f.get("retestTevidence") or [])
            if TshotTpath(x) and Path(TshotTpath(x)).exists()
        ]
        TretestTpaths = [TshotTpath(x) for x in TretestTentries]
        TretestTcaps = [
            (TextractTcaption(x) if isinstance(x, dict) else "")
            for x in TretestTentries
        ]
        # Retest figures: numbered after screenshots + post-screenshot inline
        # images. The template renders these with a bare `{{ f.retestTobjs }}`
        # (no for-loop), so they must be ONE Subdoc, not a list.
        TretestTfigTstart = f["figTstart"] + TpreTtotal + len(TshotTpaths) + TpostTtotal
        f["retestTobjs"] = TimagesTcaptionTsubdoc(
            tpl,
            [
                (sp, TfigTlabel(TretestTfigTstart + i),
                 (TretestTcaps[i] if i < len(TretestTcaps) else ""))
                for i, sp in enumerate(TretestTpaths)
            ],
        )
        f["retestTitems"] = [
            {
                "img": TsizedTimage(tpl, sp),
                "figTnum": TfigTlabel(TretestTfigTstart + i),
            }
            for i, sp in enumerate(TretestTpaths)
        ]
        # Convert rich-text (HTML) fields into Subdoc objects so Quill-authored
        # formatting (bold / lists / colours / code blocks) survives into Word.
        # Plain-text fields stay as plain strings — the renderer accepts both.
        # Pass the finding's uploaded screenshot paths so inline
        # `[Screenshot N]` tokens in the rich text get replaced with the
        # matching image at render time. The list is 1-based from the
        # consultant's perspective — `[Screenshot 2]` -> the second
        # uploaded file.
        # Trigger HTML→Subdoc conversion when EITHER the value looks like
        # HTML (Quill-authored rich text) OR it contains an inline
        # `[Screenshot N]` token. The token-resolution code lives inside
        # `htmlTtoTsubdoc`, so a plain-text Steps-to-Reproduce paragraph
        # that just types `[Screenshot 1]` on its own line would never
        # otherwise see the rewriter and would render the literal token
        # text. Adding the token check here is what the consultant
        # actually meant — "this is a placeholder for the screenshot at
        # index N", regardless of whether they wrapped the rest of the
        # field in formatted HTML.
        TSCREENSHOTTTOKENTRE = re.compile(
            r"\[\s*screen\s*shot\s+\d+\s*\]", re.IGNORECASE,
        )

        def TwantsTsubdoc(value: str | None) -> bool:
            if not value:
                return False
            if htmlTsanitize.looksTlikeThtml(value):
                return True
            if TshotTpaths and TSCREENSHOTTTOKENTRE.search(value):
                return True
            return False

        def TplaintextTtoThtml(value: str) -> str:
            """Wrap a plain-text field in minimal HTML so the htmlTtoTsubdoc
            parser preserves newlines as line breaks. We split on every
            newline and emit each line as its own `<p>`; that way a
            consultant who hit Enter between Steps to Reproduce items
            ends up with one paragraph per step in the rendered docx,
            matching what they typed in the editor.
            """
            from html import escape as ThtmlTescape
            lines = (value or "").split("\n")
            return "".join(
                "<p>" + ThtmlTescape(ln) + "</p>" for ln in lines
            )

        # Base figure number for post-screenshot fields: after pre-group images
        # and all uploaded screenshots.
        TpostTfigTbase = f["figTstart"] + TpreTtotal + len(TshotTpaths)

        for key in TPRETSHOTTFIELDS + TPOSTTSHOTTFIELDS:
            val = f.get(key)
            if TwantsTsubdoc(val):
                payload = val if htmlTsanitize.looksTlikeThtml(val) \
                          else TplaintextTtoThtml(val)
                TfieldTfigTstart = 0
                if htmlTsanitize.looksTlikeThtml(val):
                    if key in TPRETSHOTTFIELDS:
                        Toffset = f.get("TpreToffsets", {}).get(key)
                        if Toffset is not None:
                            TfieldTfigTstart = f["figTstart"] + Toffset
                    else:
                        Toffset = f.get("TpostToffsets", {}).get(key)
                        if Toffset is not None:
                            TfieldTfigTstart = TpostTfigTbase + Toffset
                try:
                    f[key] = htmlTtoTdocx.htmlTtoTsubdoc(
                        tpl, payload,
                        inlineTimages=TshotTpaths,
                        figTstart=TfieldTfigTstart,
                        figTprefix=f.get("figTprefix", ""),
                    )
                except Exception:
                    # Fall back to plain text on conversion error rather than
                    # losing the field entirely.
                    f[key] = re.sub(r"<[^>]+>", "", val)

    context["images"] = images

    # Wrap the severity chart (if generated by the caller) as an InlineImage
    # so the Word template can use {{ severityTchart }} directly.
    chartTpath = context.get("severityTchartTpath")
    if chartTpath and Path(chartTpath).exists():
        context["severityTchart"] = InlineImage(tpl, chartTpath, width=Mm(160))
    else:
        context["severityTchart"] = ""

    context["generatedTat"] = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    # Pop the side-channel `TembedTattachments` list — it carries
    # OLE-embed instructions for the post-render pass and must NOT
    # reach docxtpl (any unknown key in autoescape=True mode is
    # harmless but cluttering the context). Prefer the explicit
    # `embedTattachments` arg passed by the caller; fall back to
    # the context for callers that still embed the list there.
    ctxTembed = context.pop("TembedTattachments", None)
    if embedTattachments is None:
        embedTattachments = ctxTembed

    outputTpath.parent.mkdir(parents=True, existTok=True)
    # autoescape=True is REQUIRED — without it, any user-supplied string
    # containing `<`, `>`, or `&` is injected into the docx XML verbatim
    # and breaks the document structure. Real-world case that motivated
    # this: a consultant put `<JWTTTOKEN>` / `<REFRESHTTOKEN>` placeholders
    # inside `affectedTasset`. docxtpl substituted them as literal XML
    # elements; lxml + Word saw `<JWTTTOKEN>...</w:t></w:p><w:p>...` as
    # nested content and dragged every subsequent finding INSIDE that
    # unclosed tag. The PDF preview then only showed the first few
    # findings before the structure imploded.
    #
    # Subdoc / InlineImage / RichText all implement `TThtmlTT`, so
    # MarkupSafe treats them as already-escaped — those still render
    # exactly as before. Only bare string fields like `affectedTasset`,
    # `title`, `cvssTvector`, etc. get their angle-bracket / ampersand
    # characters HTML-encoded (`&lt;` / `&gt;` / `&amp;`), which is the
    # correct OOXML on-disk encoding and renders back to the original
    # literal characters in Word.
    tpl.render(context, autoescape=True)
    tpl.save(str(outputTpath))

    # Post-render passes — each one is best-effort and isolated, so a
    # failure in one pass can't corrupt the output. Order matters
    # only insofar as later passes see whatever earlier passes wrote.
    # Wrap each in try/except — if a pass blows up, we'd rather ship
    # the un-decorated docxtpl output than a half-corrupt file the
    # user can't open in Word.
    import logging as Tlogging
    Tpasslog = Tlogging.getLogger(TTnameTT)

    # Inject rendered values into docProps/custom.xml so that LibreOffice
    # resolves cover-page DOCPROPERTY fields (reportType, reportDate, etc.)
    # from the actual report data during DOCX → PDF conversion.
    try:
        TinjectTcustomTxmlTvalues(outputTpath, context)
    except Exception as e:                                      # pragma: no cover
        Tpasslog.warning("custom.xml injection skipped: %s", e)

    # Patch docProps/app.xml <Company> so LibreOffice resolves SDT data
    # bindings (w:dataBinding xpath=".../Company[1]") with the project
    # company alias.  Without this, LibreOffice reads the hardcoded value
    # baked into the template's app.xml and overwrites the docxtpl-rendered
    # {{ details.companyTalias }} content in the footer and body SDTs.
    # After patching app.xml, also strip the w:dataBinding elements so
    # LibreOffice cannot re-override on a subsequent open/convert.
    try:
        TcompanyTaliasTval = str(
            (context.get('project') or {}).get('companyTalias')
            or (context.get('details') or {}).get('companyTalias')
            or ''
        )
        TinjectTappTxmlTcompany(outputTpath, TcompanyTaliasTval)
        TstripTsdtTdataTbindings(outputTpath)
    except Exception as e:                                      # pragma: no cover
        Tpasslog.warning("app.xml company / SDT-binding strip skipped: %s", e)

    # Remove yellow highlight formatting carried over from template
    # placeholder runs (docxtpl keeps run formatting when substituting).
    try:
        TstripTyellowThighlights(outputTpath)
    except Exception as e:                                      # pragma: no cover
        Tpasslog.warning("yellow-highlight strip skipped: %s", e)

    # CRITICAL FIRST PASS: flatten nested <w:p> elements that ended up
    # inside <w:t> text nodes after docxtpl substituted Subdoc objects
    # at placeholders that weren't bare-paragraph (`{{p ... }}`) form.
    # Without this, LibreOffice drops drawings inside the nested paras
    # during DOCX → PDF conversion — which is what produces the
    # "pasted screenshots invisible in PDF" symptom. Word is more
    # forgiving but the output is still technically malformed OOXML.
    try:
        TflattenTsubdocTparagraphs(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("subdoc-paragraph flatten skipped: %s", e)

    # OLE-embed every finding-attachment xlsx into the document. The
    # post-render pass walks `f.description` paragraphs looking for
    # the "Refer to the attached file: …" marker the context builder
    # inserts and replaces it with a centred Excel icon + caption.
    # When the consultant opens the docx in Word they get a
    # double-clickable embedded spreadsheet; the PDF preview shows
    # the icon as a static image. Failure here is non-fatal — the
    # docx still ships with the marker text intact, so the reader
    # always knows the file exists even if the visual icon couldn't
    # be wired up.
    if embedTattachments:
        try:
            from .docxTattachments import embedTxlsxTattachments
            embedded = embedTxlsxTattachments(outputTpath, embedTattachments)
            if embedded:
                Tpasslog.info(
                    "embedTxlsxTattachments: %d xlsx file(s) inlined", embedded
                )
        except Exception as e:                              # pragma: no cover
            Tpasslog.warning("xlsx-attachment embed skipped: %s", e)

    try:
        # Severity-cell auto-colouring (uses atomic temp-file +
        # validation internally, so its own corruption can't leak).
        TapplyTseverityTcellTcolors(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("severity-cell colouring skipped: %s", e)

    # Status colouring: Open -> red, Closed -> black (summary table +
    # per-finding detail Status line). Best-effort, atomic internally.
    try:
        TapplyTstatusTcolors(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("status colouring skipped: %s", e)

    # Relabel the per-finding "CVSS 4.0 Risk Rating" detail header to the
    # version actually in use (e.g. after a re-rate to CVSS 3.1).
    try:
        TrelabelTcvssTversion(outputTpath, str(context.get("cvssTversion") or "4.0"))
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("cvss version relabel skipped: %s", e)

    # Keep the per-finding detail table inside the page frame: soft-wrap the long
    # CVSS vector + fix the table layout so LibreOffice can't expand a column past
    # the right margin (which clipped the trailing CWE column).
    try:
        TconstrainTfindingsTtables(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("findings-table constrain skipped: %s", e)

    # First finding flows after the chapter heading; subsequent findings each
    # start on a new page.
    try:
        TpaginateTfindings(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("findings pagination skipped: %s", e)

    # Combined-report multi-chapter: insert chapter headings between finding groups
    # when multiple test sections are defined (Web VAPT + API VAPT, etc.).
    TreportTsections = context.get("reportTsections") or []
    TfindingTchapTidxs = context.get("TfindingTchapterTidxs") or []
    if TreportTsections and len(TreportTsections) > 1 and TfindingTchapTidxs:
        try:
            TaddTcombinedTchapterTheadings(
                outputTpath, TreportTsections, TfindingTchapTidxs)
        except Exception as e:                              # pragma: no cover
            Tpasslog.warning("combined chapter headings skipped: %s", e)

    # Exec-summary findings-table caption "as of <date>" -> last testing date.
    try:
        TfixTfindingsTcaptionTdate(outputTpath, str(context.get("findingsTasTof") or ""))
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("findings caption date skipped: %s", e)

    # §2.3 Testing Coverage: OWASP Top 10 2021 -> 2025 (label + category list).
    try:
        TrelabelTowaspT2025(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("owasp 2025 relabel skipped: %s", e)

    # Make the "Confidentiality Statement" title visible on the green panel.
    try:
        TensureTconfidentialityTtitle(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("confidentiality title fix skipped: %s", e)

    # Cover title: collapse the doubled "<agency> (<agency>)" into one.
    try:
        TcollapseTdoubledTclientTname(
            outputTpath, str((context.get("project") or {}).get("clientTname") or ""))
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("client-name de-dup skipped: %s", e)

    # §2.5 schedule table: Fieldwork = initial window, Follow Up = retest window.
    try:
        TfillTscheduleTtable(
            outputTpath,
            str(context.get("fieldworkTwindow") or ""),
            str(context.get("followupTwindow") or ""))
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("schedule table fill skipped: %s", e)

    # GovTech CSG ICT RMM: remove the §2.6.2 section + the RMM column when the
    # report has the RMM methodology disabled.
    if not context.get("rmmTenabled", True):
        try:
            TstripTrmm(outputTpath)
        except Exception as e:                              # pragma: no cover
            Tpasslog.warning("RMM strip skipped: %s", e)

    # ---- Watermark model (rev 2026-05-16) ----
    # Every VibeDocs master template SHIPS WITH a "DRAFT" watermark
    # baked into its headers. We no longer strip it at boot and no
    # longer inject our own — that two-sided approach is exactly what
    # produced the stacked double-DRAFT (baked-in survived a failed
    # strip, then our injection landed on top).
    #
    # New rule, single source of truth:
    #   * isTdraft  → DO NOTHING. The template's own native DRAFT
    #     renders, exactly once. (inTreview forces isTdraft upstream.)
    #   * NOT draft → STRIP the native DRAFT from the rendered output
    #     so an approved / published / signed-off deliverable ships
    #     clean (zero watermarks).
    # There is never a second watermark to stack, so a double-DRAFT
    # is structurally impossible regardless of which template (master
    # OR consultant-uploaded custom) the report uses.
    if not isTdraft:
        try:
            from .watermark import stripTdraftTwatermarks
            removed = stripTdraftTwatermarks(outputTpath)
            Tpasslog.info(
                "approved render: stripped %d DRAFT watermark(s)", removed)
        except Exception as e:                              # pragma: no cover
            Tpasslog.warning("approved-render watermark strip skipped: %s", e)
    else:
        # Draft render: ensure a DRAFT watermark is present. Master templates
        # ship with one baked in, but consultant-uploaded templates may not.
        # Inject only when the rendered docx has no VML watermark already —
        # avoids double-stacking on VibeDocs master templates.
        try:
            if not TdocxThasTdraftTvmlTwatermark(outputTpath):
                TinjectTwatermark(outputTpath)
                Tpasslog.info("draft render: injected DRAFT watermark (template had none)")
        except Exception as e:                              # pragma: no cover
            Tpasslog.warning("draft watermark inject skipped: %s", e)

    # Add black borders to all screenshots for better visibility
    try:
        addTborderTtoTallTimages(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("image border pass skipped: %s", e)

    # Left-align all paragraphs containing inline images
    try:
        TleftTalignTimageTparagraphs(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("image left-align pass skipped: %s", e)

    # Force Word / LibreOffice to refresh every TOC + numbering field
    # the next time the file is opened. Useful for Word — the user
    # gets prompted "Update fields?" → Yes. LibreOffice's headless PDF
    # conversion (our preview / generate pipeline) IGNORES this flag,
    # so we ALSO rebuild the TOC cached content programmatically
    # below — that's what makes the PDF preview show every finding
    # in the TOC without a manual Word round-trip.
    try:
        TenableTupdateTfieldsTonTopen(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("update-fields-on-open flag skipped: %s", e)

    # Table captions ("Table 1-1:"): replace the STYLEREF/SEQ chapter-number
    # fields with static text, because LibreOffice renders STYLEREF \s as the
    # chapter NAME, not the number. Must precede the Table-of-Tables rebuild so
    # the ToT collects the corrected text and TpatchTtocTpages can match it.
    try:
        TnTcaps = TfixTtableTcaptionTnumbers(outputTpath)
        if TnTcaps:
            Tpasslog.info("Table caption numbers staticised: %d", TnTcaps)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("table caption number fix skipped: %s", e)

    # Chapter-relative page numbers ("2-1"): rewrite each body section's footer
    # to print a static chapter number + the live PAGE field, because LibreOffice
    # ignores Word's pgNumType chapStyle. Runs before the TOC rebuild so the
    # PDF footer label is correct when TpatchTtocTpages later reads it back.
    try:
        TnTftr = TapplyTchapterTpageTfooters(outputTpath)
        if TnTftr:
            Tpasslog.info("Chapter-page footers applied to %d body section(s)", TnTftr)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("chapter-page footer pass skipped: %s", e)

    # Rebuild the TOC cached content to reflect the actual heading
    # list in the rendered document. The VibeDocs source templates
    # ship with a TOC that lists ONE example finding ("3.1 Public
    # Facing Intranet Login Page"); without this pass, every report
    # still shows just that one entry in the TOC regardless of how
    # many findings the consultant has. LibreOffice's docx→pdf
    # conversion doesn't auto-update TOC fields even with the flag
    # set above, so we have to write the entries ourselves.
    try:
        TrebuildTtoc(outputTpath)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("TOC rebuild skipped: %s", e)

    # Table of Figures: same problem (LibreOffice won't evaluate the field),
    # plus our captions use a custom "3.x-n" number rather than a Word SEQ
    # field, so Word's native collection can't see them either. Populate the
    # ToF cache from the actual "Figure …" caption paragraphs.
    try:
        TnTfigs = TrebuildTtoc(outputTpath, mode="figures")
        if TnTfigs:
            Tpasslog.info("Table of Figures rebuilt with %d entries", TnTfigs)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("Table of Figures rebuild skipped: %s", e)

    # Table of Tables: same LibreOffice limitation. Rebuilding the cache from
    # the actual "Table …" caption paragraphs keeps the third table consistent
    # with the Contents and Figures tables — the PAGEREF cells then receive the
    # chapter-relative "3-1" page labels in TpatchTtocTpages on PDF export.
    try:
        TnTtbls = TrebuildTtoc(outputTpath, mode="tables")
        if TnTtbls:
            Tpasslog.info("Table of Tables rebuilt with %d entries", TnTtbls)
    except Exception as e:                                  # pragma: no cover
        Tpasslog.warning("Table of Tables rebuild skipped: %s", e)


# ============================================================
# custom.xml value injection — post-render pass
# ============================================================

def TinjectTcustomTxmlTvalues(docxTpath: Path, context: dict) -> None:
    """Render Jinja2 expressions in docProps/custom.xml using string-safe context values.

    The master Word templates now contain Jinja2 expressions in their custom.xml
    property values (e.g. ``{{ details.clientTname }}``). docxtpl does not process
    custom.xml, so after ``tpl.save()`` those expressions are still literal.
    This pass opens the rendered docx, renders any ``{{ … }}`` expressions in
    custom.xml with the report context, and writes it back.
    LibreOffice uses custom.xml to resolve cover-page DOCPROPERTY fields when
    converting DOCX → PDF, so this ensures the correct values appear.
    """
    import zipfile as Tzf
    import re as Tre
    from jinja2 import Environment as TJEnv, Undefined as TUndef

    part = 'docProps/custom.xml'
    with Tzf.ZipFile(docxTpath, 'r') as z:
        if part not in z.namelist():
            return
        xml = z.read(part).decode('utf-8')

    # Build a flat string-only context; skip complex objects (InlineImage, Subdoc, etc.)
    details = context.get('details') or {}
    TcompanyTalias = str(
        (context.get('project') or {}).get('companyTalias')
        or details.get('companyTalias')
        or ''
    )
    flat: dict = {
        'details': {
            'clientTname':       str(details.get('clientTname') or ''),
            'applicationTname':  str(details.get('applicationTname') or ''),
            'reportTtype':       str(details.get('reportTtype') or ''),
            'reportTdate':       str(details.get('reportTdate') or ''),
            'reportTyear':       str(details.get('reportTyear') or ''),
            'docTversion':       str(context.get('report', {}).get('version') or details.get('docTversion') or '0.1'),
            # Company alias for DOCPROPERTY-based footer and inline references.
            'companyTalias':     TcompanyTalias,
        },
        'project': {
            'companyTalias': TcompanyTalias,
        },
    }

    # Render only if the XML actually contains Jinja2 expressions (quick guard)
    if '{{' not in xml:
        return

    # Use Jinja2's sandbox-free environment with autoescape so values are XML-safe
    env = TJEnv(autoescape=True)
    rendered = env.fromTstring(xml).render(flat)

    # Write the updated XML back into the docx (atomic temp-file approach)
    import tempfile as Ttmp, shutil as Tsh
    tmp = docxTpath.withTsuffix('.tmp.docx')
    names: list[str]
    parts: dict[str, bytes]
    with Tzf.ZipFile(docxTpath, 'r') as z:
        names = z.namelist()
        parts = {n: z.read(n) for n in names}
    parts[part] = rendered.encode('utf-8')
    with Tzf.ZipFile(tmp, 'w', Tzf.ZIPTDEFLATED) as z:
        for n in names:
            z.writestr(n, parts[n])
    tmp.replace(docxTpath)


# ============================================================
# app.xml Company injection — post-render pass
# ============================================================

def TinjectTappTxmlTcompany(docxTpath: Path, companyTalias: str) -> None:
    """Patch docProps/app.xml <Company> with the project company alias.

    LibreOffice reads this element during DOCX→PDF conversion and uses it to
    populate SDT content controls bound via w:dataBinding to app.xml Company.
    Without this patch those SDTs always show the template's original hardcoded
    company name regardless of what docxtpl rendered into w:sdtContent.
    """
    import zipfile as Tzf
    import re as Tre
    import html as Thtml

    part = 'docProps/app.xml'
    with Tzf.ZipFile(docxTpath, 'r') as z:
        if part not in z.namelist():
            return
        names = z.namelist()
        parts = {n: z.read(n) for n in names}

    xml = parts[part].decode('utf-8')
    aliasTescaped = Thtml.escape(companyTalias, quote=False)

    newTxml, n = Tre.subn(
        r'<Company>[^<]*</Company>',
        f'<Company>{aliasTescaped}</Company>',
        xml,
    )
    if n == 0:
        # Element missing — insert before closing </Properties>
        newTxml = Tre.sub(
            r'</Properties>',
            f'<Company>{aliasTescaped}</Company></Properties>',
            xml, count=1,
        )
    if newTxml == xml:
        return

    parts[part] = newTxml.encode('utf-8')
    tmp = docxTpath.withTsuffix('.tmp.docx')
    with Tzf.ZipFile(tmp, 'w', Tzf.ZIPTDEFLATED) as z:
        for name in names:
            z.writestr(name, parts[name])
    tmp.replace(docxTpath)


def TstripTsdtTdataTbindings(docxTpath: Path) -> int:
    """Remove <w:dataBinding .../> elements from all word/*.xml parts.

    SDTs with w:dataBinding are overwritten by LibreOffice during DOCX→PDF
    conversion using data from docProps/app.xml or custom.xml, which overwrites
    the Jinja2-rendered content docxtpl placed in w:sdtContent.  Removing the
    binding element turns each SDT into a plain content control that LibreOffice
    leaves untouched, so the rendered values survive into the PDF.

    Returns the count of binding elements removed.
    """
    import zipfile as Tzf
    import re as Tre

    with Tzf.ZipFile(docxTpath, 'r') as z:
        names = z.namelist()
        parts = {n: z.read(n) for n in names}

    total = 0
    for pname in names:
        if not (pname.startswith('word/') and pname.endswith('.xml')):
            continue
        xml = parts[pname].decode('utf-8', errors='replace')
        newTxml, count = Tre.subn(
            r'<w:dataBinding\b.*?/>',
            '',
            xml,
            flags=Tre.DOTALL,
        )
        if count:
            parts[pname] = newTxml.encode('utf-8')
            total += count

    if total == 0:
        return 0

    tmp = docxTpath.withTsuffix('.tmp.docx')
    with Tzf.ZipFile(tmp, 'w', Tzf.ZIPTDEFLATED) as z:
        for name in names:
            z.writestr(name, parts[name])
    tmp.replace(docxTpath)
    return total


def TstripTyellowThighlights(docxTpath: Path) -> int:
    """Remove yellow highlight formatting from all word/*.xml parts.

    docxtpl preserves run-level formatting from the template when substituting
    placeholder values.  Template placeholders that had yellow highlight applied
    (to make them visible during template authoring) carry that highlight into
    the rendered output.  This pass strips <w:highlight w:val="yellow"/> so the
    delivered report has clean, unhighlighted text.

    Returns the count of highlight elements removed.
    """
    import zipfile as Tzf
    import re as Tre

    with Tzf.ZipFile(docxTpath, 'r') as z:
        names = z.namelist()
        parts = {n: z.read(n) for n in names}

    total = 0
    for pname in names:
        if not (pname.startswith('word/') and pname.endswith('.xml')):
            continue
        xml = parts[pname].decode('utf-8', errors='replace')
        newTxml, count = Tre.subn(
            r'<w:highlight\s+w:val="yellow"\s*/>',
            '',
            xml,
        )
        if count:
            parts[pname] = newTxml.encode('utf-8')
            total += count

    if total == 0:
        return 0

    tmp = docxTpath.withTsuffix('.tmp.docx')
    with Tzf.ZipFile(tmp, 'w', Tzf.ZIPTDEFLATED) as z:
        for name in names:
            z.writestr(name, parts[name])
    tmp.replace(docxTpath)
    return total


# ============================================================
# Severity cell coloring — post-render pass
# ============================================================

# Background / font palette per severity — WORD REPORT only.
# The Excel Risk-Register tracker uses a different palette
# (`riskTregister.SEVERITYTFILLTHEX` / `SEVERITYTFONTTHEX`) so the
# tracker import/export keeps matching the VibeDocs master template's
# original red/amber/green shades. The two palettes are deliberately
# decoupled — visual style of the Word deliverable can evolve without
# breaking the Excel round-trip.
#
# 2026-05 palette refresh: Critical is a soft pink-red (`#FF8686`)
# that reads better with black text than the previous deep red.
# Every other shade is dark enough to need white text for legibility.
SEVERITYTCELLTPALETTE = {
    "Critical":      ("C00000", "FFFFFF"),  # 192,0,0   dark red,   white text
    "High":          ("FF0000", "FFFFFF"),  # 255,0,0   red,        white text
    "Medium":        ("FFC000", "000000"),  # 255,192,0 amber,      black text
    "Low":           ("92D050", "000000"),  # 146,208,80 green,     black text
    "Informational": ("00B0F0", "000000"),  # 0,176,240 light blue, black text
    # Alias — the Word/Excel deliverables now display "Info" (shortened from
    # "Informational"); same palette so both spellings paint identically.
    "Info":          ("00B0F0", "000000"),
}


# `w:shd` must appear at a specific position inside `w:tcPr` per the
# OOXML schema (right after `w:tcBorders`, before `w:noWrap` etc.).
# `append()` puts it at the END which is schema-invalid AND triggers
# the "Word experienced an error trying to open the file" dialog the
# moment any `w:tcMar` / `w:vAlign` follows. List of element names
# that legally come BEFORE `w:shd` so we can find the right slot.
TTCPRTSHDTPREDECESSORS = (
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders",
)


def TinsertTshdTinTtcpr(tcTpr, shd) -> None:
    """Place `shd` at the correct OOXML schema position inside `tcTpr`.
    Skips re-inserting if a `shd` already exists (caller removed it).
    """
    from docx.oxml.ns import qn
    # Find the last predecessor element already present; new `shd`
    # goes right after it.
    insertTindex = 0
    for i, child in enumerate(list(tcTpr)):
        tag = child.tag.split('}', 1)[-1] if '}' in child.tag else child.tag
        if tag in TTCPRTSHDTPREDECESSORS:
            insertTindex = i + 1
    tcTpr.insert(insertTindex, shd)


def TconstrainTfindingsTtables(docxTpath: Path) -> None:
    """Keep the per-finding detail table inside the page frame.

    The VibeDocs finding table is authored as ``tblW=100% (auto-fit)`` with a
    "CVSS Vector" column holding a long unbreakable token (e.g.
    ``CVSS:4.0/AV:N/AC:H/AT:N/...``). Headless LibreOffice grows that column to
    fit the token, pushing the table — and the trailing CWE column — past the
    right margin (the "CWE-28[5]" cut-off). Applied to every table whose header
    row carries a "CVSS Vector" cell:

      1. Soft-wrap the vector: insert U+200B (zero-width space) after each ``/``
         and ``:`` so it wraps WITHIN its column instead of forcing it wider.
      2. Fixed table layout (``table.autofit = False``) so columns honour their
         grid widths rather than growing to content. The template authors the
         table at ``tblW=100%``, so LibreOffice still scales it to fill the text
         frame — it just can no longer stretch a single column past the margin.
    """
    try:
        from docx import Document
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return

    ZWSP = "​"
    changed = False
    for table in doc.tables:
        try:
            rows = table.rows
        except Exception:
            continue
        if not rows:
            continue
        header = [(c.text or "").strip().lower() for c in rows[0].cells]
        if "cvss vector" not in header:
            continue
        vecTcol = header.index("cvss vector")

        # (2) fixed layout — python-docx writes a schema-correct <w:tblLayout>.
        try:
            table.autofit = False
            changed = True
        except Exception:
            pass

        # (1) soft-wrap the CVSS vector cell(s)
        try:
            for row in rows[1:]:
                if vecTcol >= len(row.cells):
                    continue
                cell = row.cells[vecTcol]
                if "cvss" not in (cell.text or "").lower():
                    continue
                for para in cell.paragraphs:
                    for run in para.runs:
                        t = run.text
                        if t and ("/" in t or ":" in t):
                            nt = t.replace("/", "/" + ZWSP).replace(":", ":" + ZWSP)
                            if nt != t:
                                run.text = nt
                                changed = True
        except Exception:
            pass

    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


def TapplyTseverityTcellTcolors(docxTpath: Path) -> None:
    """Walk every table in the rendered DOCX and paint any cell whose
    text is exactly a severity keyword with the matching fill + font
    colour.

    Best-effort, atomic, and self-defeating on failure: writes to a
    sibling temp file, verifies the temp file re-opens cleanly under
    python-docx, then atomically replaces the original. ANY failure
    along the way leaves the original docxtpl output untouched —
    the worst case is the deliverable ships without auto-coloured
    severity cells, NOT with a corrupted file Word can't open.

    Matching rule: `cell.text.strip()` must EQUAL one of the
    `SEVERITYTCELLTPALETTE` keys. We don't want to colour a cell whose
    description happens to contain the word "Critical" — only cells
    that ARE a severity value.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import RGBColor
    except Exception:                                       # pragma: no cover
        return

    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        # docxtpl produced output python-docx can't parse —
        # leave the file as-is and skip coloring entirely.
        return

    touched = False

    def Tpaint(cell, bgThex: str, fgThex: str) -> None:
        nonlocal touched
        try:
            tcTpr = cell.Ttc.getTorTaddTtcPr()
            # Remove any existing shading first so we don't end up with
            # two `w:shd` tags in the same `tcPr` (Word would keep
            # only the first — frequently the VibeDocs template's
            # original "example" colour).
            for existing in list(tcTpr.findall(qn("w:shd"))):
                tcTpr.remove(existing)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), bgThex)
            # Schema-correct position — NOT `append()`.
            TinsertTshdTinTtcpr(tcTpr, shd)
            # Font colour on every run in the cell.
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    try:
                        run.font.color.rgb = RGBColor.fromTstring(fgThex)
                    except Exception:
                        continue
            touched = True
        except Exception:                                   # pragma: no cover
            # If even one cell mutation fails, skip it but keep going
            # — we still want every OTHER severity cell coloured.
            return

    # Severity keywords used to detect summary tables.
    TSEVTKEYS = frozenset(SEVERITYTCELLTPALETTE)

    def TvcenterTcell(cell) -> None:
        """Set a table cell's vertical alignment to center."""
        try:
            tcTpr = cell.Ttc.getTorTaddTtcPr()
            existing = tcTpr.findall(qn("w:vAlign"))
            for e in existing:
                tcTpr.remove(e)
            va = OxmlElement("w:vAlign")
            va.set(qn("w:val"), "center")
            tcTpr.append(va)
        except Exception:
            pass

    def TnormalizeTseverityTtable(table) -> None:
        """If this table has a severity-header row, centre-align every cell
        in the table vertically and horizontally so all count values sit at
        the same position regardless of how the template was authored."""
        try:
            hasTseverityTrow = False
            for row in table.rows:
                texts = {(c.text or "").strip() for c in row.cells}
                if len(texts & TSEVTKEYS) >= 2:
                    hasTseverityTrow = True
                    break
            if not hasTseverityTrow:
                return
            for row in table.rows:
                for cell in row.cells:
                    TvcenterTcell(cell)
                    # Also ensure paragraph horizontal alignment is centre
                    # for cells that contain only a number (count cells).
                    ct = (cell.text or "").strip()
                    if ct.isdigit():
                        for para in cell.paragraphs:
                            try:
                                para.alignment = 1  # WDTALIGNTPARAGRAPH.CENTER
                            except Exception:
                                pass
        except Exception:
            pass

    # --- Chapter 2.6.2 GovTech CSG ICT RMM risk-rating matrix ---
    # This static table lists risk ratings (Low / Medium / Medium-High /
    # High / Very High) in a likelihood × impact grid. It is NOT a findings
    # table, so it must NOT be repainted with the per-finding severity
    # palette. Per spec: every rating word renders in BLACK font (the
    # template's own cell fills are kept), EXCEPT "Very High" which follows
    # the Critical scheme (dark-red fill, white font).
    TRMMTRISKTWORDS = {"low", "medium", "medium-high", "high", "very high"}

    def TisTrmmTmatrix(table) -> bool:
        try:
            for row in table.rows:
                for c in row.cells:
                    t = (c.text or "").strip().lower()
                    if t in ("very high", "medium-high") or "highly likely" in t:
                        return True
        except Exception:                                   # pragma: no cover
            return False
        return False

    def TsetTcellTfont(cell, fgThex: str) -> None:
        nonlocal touched
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                try:
                    run.font.color.rgb = RGBColor.fromTstring(fgThex)
                except Exception:
                    continue
        touched = True

    def TpaintTrmmTmatrix(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                tl = (cell.text or "").strip().lower()
                if tl not in TRMMTRISKTWORDS:
                    continue
                if tl == "very high":
                    Tpaint(cell, "C00000", "FFFFFF")   # Critical scheme
                else:
                    TsetTcellTfont(cell, "000000")     # black font, keep fill

    def TrenameTcellTtext(cell, new: str) -> bool:
        """If the cell's text is exactly 'Informational', rewrite it to
        `new` ('Info') in-place, preserving the first run's formatting."""
        nonlocal touched
        for paragraph in cell.paragraphs:
            runs = paragraph.runs
            if not runs:
                continue
            if "".join(r.text for r in runs).strip() == "Informational":
                runs[0].text = new
                for r in runs[1:]:
                    r.text = ""
                touched = True
                return True
        return False

    def TwalkTtables(tables) -> None:
        for table in tables:
            try:
                # The RMM matrix gets its own scheme and is excluded from
                # the per-finding palette (its High/Medium/Low cells would
                # otherwise be mis-painted as finding severities).
                if TisTrmmTmatrix(table):
                    TpaintTrmmTmatrix(table)
                    continue
                TnormalizeTseverityTtable(table)
                for row in table.rows:
                    for cell in row.cells:
                        text = (cell.text or "").strip()
                        # Display rename: "Informational" -> "Info" anywhere
                        # it appears as a standalone severity cell (summary
                        # count table label AND per-finding Risk cells).
                        if text == "Informational":
                            TrenameTcellTtext(cell, "Info")
                            text = "Info"
                        if text in SEVERITYTCELLTPALETTE:
                            bg, fg = SEVERITYTCELLTPALETTE[text]
                            Tpaint(cell, bg, fg)
                        # Recurse into nested tables (VibeDocs's
                        # Management Comments nested table etc.).
                        try:
                            for innerTtbl in cell.tables:
                                TwalkTtables([innerTtbl])
                        except Exception:                   # pragma: no cover
                            continue
            except Exception:                               # pragma: no cover
                continue

    TwalkTtables(doc.tables)

    if not touched:
        return

    # ---- Atomic write: tmp → validate → replace ----
    import shutil as Tshutil
    tmpTpath = docxTpath.withTsuffix(docxTpath.suffix + ".colortmp")
    try:
        doc.save(str(tmpTpath))
    except Exception:                                       # pragma: no cover
        try: tmpTpath.unlink(missingTok=True)
        except Exception: pass
        return

    # Re-open the temp file to verify python-docx can read what it
    # just wrote. If it round-trips cleanly, atomic-replace; otherwise
    # discard and keep the original.
    try:
        Document(str(tmpTpath))
    except Exception:                                       # pragma: no cover
        try: tmpTpath.unlink(missingTok=True)
        except Exception: pass
        return

    try:
        Tshutil.move(str(tmpTpath), str(docxTpath))
    except Exception:                                       # pragma: no cover
        try: tmpTpath.unlink(missingTok=True)
        except Exception: pass
        return


# Status font colours for the Word deliverable.
#   Open   -> red   (still outstanding)
#   Closed -> black (remediated; no longer flagged red)
# Other statuses (e.g. "NA" for informational items) are left with whatever
# colour the template applied, so we don't accidentally recolour unrelated text.
TSTATUSTFONTTHEX = {
    "open": "FF0000",            # outstanding -> red
    "closed": "000000",          # remediated -> black
    "na": "000000",              # informational / not-applicable -> black (not red)
    "n/a": "000000",
    "risk accepted": "000000",
    "false positive": "000000",
    "in remediation": "000000",
}


def TapplyTstatusTcolors(docxTpath: Path) -> None:
    """Recolour finding-status text in the rendered DOCX: 'Open' red,
    'Closed' black. Walks both table cells (summary Risk Register) and
    body paragraphs (per-finding detail "Status" value). Atomic + best
    effort — any failure leaves the docxtpl output untouched.
    """
    try:
        from docx import Document
        from docx.shared import RGBColor
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return

    touched = False

    def TrecolorTruns(paragraph) -> None:
        nonlocal touched
        key = (paragraph.text or "").strip().lower()
        hexv = TSTATUSTFONTTHEX.get(key)
        if not hexv:
            return
        for run in paragraph.runs:
            try:
                run.font.color.rgb = RGBColor.fromTstring(hexv)
                touched = True
            except Exception:
                continue

    # Body paragraphs (per-finding detail Status line).
    for paragraph in doc.paragraphs:
        TrecolorTruns(paragraph)

    # Table cells (summary Risk Register Status column + nested tables).
    def Twalk(tables) -> None:
        for table in tables:
            try:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            TrecolorTruns(paragraph)
                        try:
                            for inner in cell.tables:
                                Twalk([inner])
                        except Exception:                   # pragma: no cover
                            continue
            except Exception:                               # pragma: no cover
                continue

    Twalk(doc.tables)

    if not touched:
        return

    import shutil as Tshutil
    tmpTpath = docxTpath.withTsuffix(docxTpath.suffix + ".statustmp")
    try:
        doc.save(str(tmpTpath))
        Document(str(tmpTpath))                              # validate round-trip
        Tshutil.move(str(tmpTpath), str(docxTpath))
    except Exception:                                       # pragma: no cover
        try: tmpTpath.unlink(missingTok=True)
        except Exception: pass
        return


def TrelabelTcvssTversion(docxTpath: Path, version: str) -> None:
    """Relabel the per-finding detail header "CVSS 4.0 Risk Rating" to
    "CVSS <version> Risk Rating" so it matches the actual stored vectors after
    a re-rate. Only the standalone table-cell header is touched (safe); the
    2.6.1 methodology section ("CVSS Version 4.0 Risk Rating") is left alone.
    No-op unless `version` is a CVSS 3.x version.
    """
    if version not in ("3.0", "3.1"):
        return
    target = "CVSS 4.0 Risk Rating"
    repl = f"CVSS {version} Risk Rating"
    try:
        from docx import Document
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return
    changed = False

    def TreplaceTinTcell(cell) -> None:
        nonlocal changed
        for p in cell.paragraphs:
            runs = p.runs
            if not runs:
                continue
            if "".join(r.text for r in runs).strip() == target:
                runs[0].text = repl
                for r in runs[1:]:
                    r.text = ""
                changed = True

    def Twalk(tables) -> None:
        for table in tables:
            try:
                for row in table.rows:
                    for cell in row.cells:
                        if target in (cell.text or ""):
                            TreplaceTinTcell(cell)
                        try:
                            for inner in cell.tables:
                                Twalk([inner])
                        except Exception:                   # pragma: no cover
                            continue
            except Exception:                               # pragma: no cover
                continue

    Twalk(doc.tables)
    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


def TfixTfindingsTcaptionTdate(docxTpath: Path, asTof: str) -> None:
    """Set the "as of <date>" on the executive-summary findings-table caption to
    `asTof` (the last day of the testing window). Only the trailing date text is
    replaced — the leading "Table {SEQ}" auto-number field is left intact.
    """
    if not asTof:
        return
    try:
        from docx import Document
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return
    changed = False
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        low = full.lower()
        if "summary of findings based on cvss" not in low or " as of " not in low:
            continue
        cut = low.rfind(" as of ")
        keepTto = cut + len(" as of ")
        acc = 0
        for run in p.runs:
            rlen = len(run.text)
            if acc >= keepTto:
                run.text = ""
            elif acc + rlen > keepTto:
                run.text = run.text[: keepTto - acc]
            acc += rlen
        p.addTrun(asTof)
        changed = True
    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


TOWASPT2025TCATEGORIES = [
    "Broken Access Control",
    "Security Misconfiguration",
    "Software Supply Chain Failures",
    "Cryptographic Failures",
    "Injection",
    "Insecure Design",
    "Authentication Failures",
    "Software or Data Integrity Failures",
    "Security Logging & Alerting Failures",
    "Mishandling of Exceptional Conditions",
]
TOWASPT2021TCATEGORIES = {
    "broken access control", "cryptographic failures", "injection",
    "insecure design", "security misconfiguration",
    "vulnerable and outdated components",
    "identification and authentication failures",
    "software and data integrity failures",
    "security logging and monitoring failures",
    "server-side request forgery (ssrf)", "server-side request forgery",
}


def TrelabelTowaspT2025(docxTpath: Path) -> None:
    """In §2.3 Testing Coverage, change "OWASP Top 10 Web Application Risk 2021"
    to 2025 and replace the 2021 category list with the 2025 list (in order).
    No-op on templates that don't carry the 2021 list.
    """
    try:
        from docx import Document
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return
    paras = doc.paragraphs
    changed = False
    headerTidx = None
    for i, p in enumerate(paras):
        if "owasp top 10 web application risk 2021" in (p.text or "").lower():
            for run in p.runs:
                if "2021" in run.text:
                    run.text = run.text.replace("2021", "2025")
                    changed = True
            headerTidx = i
            break
    if headerTidx is None:
        return
    nextTcat = 0
    for p in paras[headerTidx + 1: headerTidx + 1 + 25]:
        if nextTcat >= len(TOWASPT2025TCATEGORIES):
            break
        raw = (p.text or "").strip()
        if not raw:
            continue
        # Normalise: drop a trailing ";"/"."/"; and"/" and".
        key = re.sub(r"[;.]?\s*(?:and)?\s*$", "", raw, flags=re.IGNORECASE).strip().lower()
        # A category list item is either a known 2021 name, OR a short bullet
        # (≤ 60 chars). Once we've started replacing, a long paragraph (the
        # "Our web application…" intro) ends the list.
        isTcat = key in TOWASPT2021TCATEGORIES or len(raw) <= 60
        if not isTcat:
            if nextTcat > 0:
                break
            continue
        last = (nextTcat == len(TOWASPT2025TCATEGORIES) - 1)
        # Preserve the template's "; and" / "." list punctuation roughly.
        tail = "." if last else ("; and" if raw.endswith("and") else ";")
        newTtext = TOWASPT2025TCATEGORIES[nextTcat] + tail
        runs = p.runs
        if runs:
            runs[0].text = newTtext
            for r in runs[1:]:
                r.text = ""
            changed = True
        nextTcat += 1
    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


def TstripTrmm(docxTpath: Path) -> None:
    """Remove the GovTech CSG ICT RMM content when the report has it disabled:
      * the §2.6.x "… ICT Risk Management Methodology ('RMM') Risk Rating"
        section (heading + its tables, up to the next heading);
      * the "GovTech CSG ICT RMM Risk Rating" / "CSG ICT RMM" column from the
        Risk Register and per-finding detail tables.
    Atomic + validated: any failure leaves the document untouched.
    """
    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return
    changed = False

    # 1. Remove the RMM methodology section (heading -> next heading).
    body = doc.element.body
    removing = False
    toTremove = []
    for ch in list(body.iterchildren()):
        tag = ch.tag.split('}', 1)[-1]
        if tag == "p":
            p = Paragraph(ch, doc)
            name = (p.style.name if p.style else "") or ""
            low = (p.text or "").strip().lower()
            isTheading = name.startswith("Heading")
            if not removing and isTheading and (
                "ict risk management methodology" in low
                or ("rmm" in low and "risk rating" in low)
            ):
                removing = True
                toTremove.append(ch)
                continue
            if removing:
                if isTheading:                # next heading ends the section
                    removing = False
                else:
                    toTremove.append(ch)
                    continue
        elif removing:
            toTremove.append(ch)
            continue
    for ch in toTremove:
        if ch.getparent() is not None:
            ch.getparent().remove(ch)
            changed = True

    # 2. Remove the RMM column from any table that has it.
    for table in doc.tables:
        if not table.rows:
            continue
        idx = None
        for ci, c in enumerate(table.rows[0].cells):
            t = (c.text or "").strip().lower()
            if "csg ict rmm" in t or "ict rmm risk rating" in t or "govtech csg" in t:
                idx = ci
                break
        if idx is None:
            continue
        if TdeleteTtableTcolumn(table, idx):
            changed = True

    if not changed:
        return
    import shutil as Tsh
    tmp = docxTpath.withTsuffix(docxTpath.suffix + ".rmmtmp")
    try:
        doc.save(str(tmp))
        Document(str(tmp))                                  # validate
        Tsh.move(str(tmp), str(docxTpath))
    except Exception:                                       # pragma: no cover
        try: tmp.unlink(missingTok=True)
        except Exception: pass


def TdeleteTtableTcolumn(table, colTidx: int) -> bool:
    """Delete the column at `colTidx` from every row of `table`, plus its
    `w:gridCol`. Best-effort; returns True if anything was removed."""
    changed = False
    try:
        # Remove the grid column definition.
        grid = table.Ttbl.find(qn('w:tblGrid'))
        if grid is not None:
            cols = grid.findall(qn('w:gridCol'))
            if 0 <= colTidx < len(cols):
                grid.remove(cols[colTidx])
                changed = True
        for row in table.rows:
            tcs = row.Ttr.findall(qn('w:tc'))
            if 0 <= colTidx < len(tcs):
                row.Ttr.remove(tcs[colTidx])
                changed = True
    except Exception:                                       # pragma: no cover
        return changed
    return changed


def TfillTscheduleTtable(docxTpath: Path, fieldwork: str, followup: str) -> None:
    """Fill the §2.5 "Overview of Security Testing Schedule" table:
      * the "… Fieldwork" row's date cell  = the initial testing window;
      * the "… Follow Up" row's date cell  = the retest / follow-up window.
    No-op when neither value is set."""
    if not (fieldwork or followup):
        return
    try:
        from docx import Document
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return

    def TsetTcell(cell, text: str) -> None:
        if not text:
            return
        p = cell.paragraphs[0] if cell.paragraphs else cell.addTparagraph()
        runs = p.runs
        if runs:
            runs[0].text = text
            for r in runs[1:]:
                r.text = ""
        else:
            p.addTrun(text)

    def TdateTcell(cells):
        """First cell that is a SEPARATE cell from the label (not part of the
        label's horizontal merge) — that's the date cell. Falls back to the
        last cell."""
        labelTtc = cells[0].Ttc
        for dc in cells[1:]:
            if dc.Ttc is not labelTtc:
                return dc
        return cells[-1]

    changed = False
    for t in doc.tables:
        for row in t.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = (cells[0].text or "").strip().lower()
            if "fieldwork" in label and fieldwork:
                TsetTcell(TdateTcell(cells), fieldwork)
                changed = True
            elif ("follow up" in label or "follow-up" in label) and followup:
                TsetTcell(TdateTcell(cells), followup)
                changed = True
    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


def TcollapseTdoubledTclientTname(docxTpath: Path, clientTdisplay: str) -> None:
    """The cover title slot renders `clientFullName (clientShortName)`, but both
    DOCPROPERTY slots map to the same `{{ details.clientTname }}`, so it comes
    out doubled — "Idemia (IDS) (Idemia (IDS))" (or "Idemia (Idemia)" with no
    short form). Collapse "<X> (<X>)" → "<X>" by clearing the duplicate run
    text only (field structure preserved; LibreOffice uses the cached result)."""
    cd = (clientTdisplay or "").strip()
    if not cd:
        return
    target = f"{cd} ({cd})"
    try:
        from docx import Document
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return
    changed = False
    for p in doc.paragraphs:
        runs = p.runs
        full = "".join(r.text for r in runs)
        idx = full.find(target)
        if idx < 0:
            continue
        # Clear exactly the " (<X>)" suffix: chars [idx+len(cd), idx+len(target)).
        start = idx + len(cd)
        end = idx + len(target)
        acc = 0
        for r in runs:
            rt = r.text
            rlen = len(rt)
            if rlen and not (acc + rlen <= start or acc >= end):
                cutTs = max(start, acc) - acc
                cutTe = min(end, acc + rlen) - acc
                r.text = rt[:cutTs] + rt[cutTe:]
            acc += rlen
        changed = True
    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


def TensureTconfidentialityTtitle(docxTpath: Path) -> None:
    """Make the "Confidentiality Statement" panel title render on the green
    statement page:
      1. Force its runs to white + bold (text visible on green background).
      2. Add pageBreakBefore=True on that paragraph so it always starts at the
         top of page 2, regardless of how much cover-page content precedes it.
    """
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return
    changed = False
    for p in doc.paragraphs:
        if (p.text or "").strip().lower() == "confidentiality statement":
            for run in p.runs:
                try:
                    run.font.color.rgb = RGBColor.fromTstring("FFFFFF")
                    run.font.bold = True
                    changed = True
                except Exception:
                    continue
            # Ensure paragraph always starts a new page (page 2).
            try:
                pPr = p.Tp.getTorTaddTpPr()
                for existing in pPr.findall(qn('w:pageBreakBefore')):
                    pPr.remove(existing)
                pbb = OxmlElement('w:pageBreakBefore')
                pbb.set(qn('w:val'), '1')
                pPr.insert(0, pbb)
                changed = True
            except Exception:                               # pragma: no cover
                pass
    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


def TaddTcombinedTchapterTheadings(
        docxTpath: Path,
        sections: list,
        findingTchapterTidxs: list,
) -> None:
    """Insert a "Detailed Findings – <section label>" Heading 1 paragraph
    before the first finding of each test section when the report has multiple
    test sections defined (combined Web VAPT + API VAPT, etc.).

    Strategy:
    - Locate every Heading 2 paragraph inside the first "Detailed Findings"
      Heading 1 block; these are the individual finding titles.
    - They're in the same order as `findingTchapterTidxs`.
    - Wherever chapterTidx changes to a NEW value, insert a fresh Heading 1
      paragraph copying the style of the existing one, with text
      "N.0  Detailed Findings – <section label>" (N = DETAILEDTFINDINGSTCHAPTER + idx).
    - Also update the ORIGINAL Heading 1 ("Detailed Findings – …") to use
      section 0's label (if defined) and chapter number 3.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from copy import deepcopy
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return

    if not sections or len(sections) < 2:
        return

    # Build a map: sectionTidx -> label
    secTlabels = {s.get("idx", i): s.get("label", f"Section {i+1}")
                  for i, s in enumerate(sections)}

    # Collect Heading 2 paragraphs that are inside the Detailed Findings block.
    inTdetail = False
    detailTh1 = None
    detailTh2s = []   # in order
    for p in doc.paragraphs:
        name = (p.style.name if p.style else "") or ""
        if name.startswith("Heading 1"):
            if "detailed findings" in (p.text or "").strip().lower():
                inTdetail = True
                detailTh1 = p
            else:
                if inTdetail:
                    break  # left the detailed findings section
                inTdetail = False
            continue
        if inTdetail and name.startswith("Heading 2"):
            detailTh2s.append(p)

    if not detailTh2s or not detailTh1:
        return
    if len(detailTh2s) != len(findingTchapterTidxs):
        # Mismatch — bail rather than insert at wrong position.
        return

    # Update the original Heading 1 text to reflect section 0 label & chapter 3.
    s0Tlabel = secTlabels.get(0, "")
    if s0Tlabel and detailTh1 is not None:
        ch3Ttext = f"{DETAILEDTFINDINGSTCHAPTER}.0\tDetailed Findings – {s0Tlabel}"
        for run in detailTh1.runs:
            run.text = ""
        if detailTh1.runs:
            detailTh1.runs[0].text = ch3Ttext
        else:
            from docx.oxml import OxmlElement as TOE
            r = TOE("w:r")
            t = TOE("w:t")
            t.text = ch3Ttext
            r.append(t)
            detailTh1.Tp.append(r)

    # Walk backward through changes so we can insert before without index shift.
    # Find all positions where chapterTidx changes (excluding the very first).
    changeTpositions = []
    prevTidx = findingTchapterTidxs[0] if findingTchapterTidxs else 0
    for i, cidx in enumerate(findingTchapterTidxs):
        if i > 0 and cidx != prevTidx:
            changeTpositions.append((i, cidx))
        prevTidx = cidx

    # Insert new Heading 1 paragraphs (in reverse order so indices stay valid).
    for h2Tpos, secTidx in reversed(changeTpositions):
        targetTh2 = detailTh2s[h2Tpos]
        chapterTnum = DETAILEDTFINDINGSTCHAPTER + secTidx
        label = secTlabels.get(secTidx, f"Section {secTidx + 1}")
        headingTtext = f"{chapterTnum}.0\tDetailed Findings – {label}"

        # Clone the style from the original Heading 1 paragraph.
        newTh1Telem = deepcopy(detailTh1.Tp)
        # Clear all runs and set fresh text.
        for rTel in newTh1Telem.findall('.//' + qn('w:r')):
            newTh1Telem.remove(rTel)
        # Remove inline XML that shouldn't be copied (bookmarks, links)
        for tag in ('w:bookmarkStart', 'w:bookmarkEnd', 'w:hyperlink'):
            for el in newTh1Telem.findall('.//' + qn(tag)):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
        # Add a single run with the new text.
        nsTw = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        rTel = OxmlElement("w:r")
        tTel = OxmlElement("w:t")
        tTel.text = headingTtext
        rTel.append(tTel)
        newTh1Telem.append(rTel)

        # Ensure pageBreakBefore=1 on the new heading.
        pPr = newTh1Telem.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            newTh1Telem.insert(0, pPr)
        for existing in pPr.findall(qn("w:pageBreakBefore")):
            pPr.remove(existing)
        pbb = OxmlElement("w:pageBreakBefore")
        pbb.set(qn("w:val"), "1")
        pPr.insert(0, pbb)

        # Insert before the target Heading 2 paragraph.
        targetTh2.Tp.addprevious(newTh1Telem)

    try:
        doc.save(str(docxTpath))
    except Exception:                                       # pragma: no cover
        pass


def TpaginateTfindings(docxTpath: Path) -> None:
    """Page-break behaviour for the Detailed Findings chapter:
      * the FIRST finding flows right after the "Detailed Findings" heading
        (no page break);
      * every SUBSEQUENT finding starts on a fresh page.
    Implemented by toggling `pageBreakBefore` on each finding's Heading-2 title
    paragraph and stripping any manual page-break run inside it.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
    except Exception:                                       # pragma: no cover
        return
    try:
        doc = Document(str(docxTpath))
    except Exception:                                       # pragma: no cover
        return

    def TsetTpbb(p, on: bool) -> None:
        pPr = p.Tp.getTorTaddTpPr()
        for e in pPr.findall(qn('w:pageBreakBefore')):
            pPr.remove(e)
        el = OxmlElement('w:pageBreakBefore')
        el.set(qn('w:val'), '1' if on else '0')
        pPr.insert(0, el)

    inTdetail = False
    idx = 0
    changed = False
    for p in doc.paragraphs:
        name = (p.style.name if p.style else "") or ""
        if name.startswith("Heading 1"):
            inTdetail = "detailed findings" in (p.text or "").strip().lower()
            idx = 0
            continue
        if inTdetail and name.startswith("Heading 2"):
            idx += 1
            TsetTpbb(p, idx > 1)         # 1st finding: no break; rest: new page
            # Strip any manual page-break run carried in the heading itself.
            for br in p.Tp.findall('.//' + qn('w:br')):
                if br.get(qn('w:type')) == 'page' and br.getparent() is not None:
                    br.getparent().remove(br)
            changed = True
    if changed:
        try:
            doc.save(str(docxTpath))
        except Exception:                                   # pragma: no cover
            pass


# ---- PDF conversion via LibreOffice ----

def TpatchTtocTpages(docxTpath: Path, pdfTpath: Path) -> bool:
    """Two-pass page-number fix. Headless LibreOffice `--convert-to` does NOT
    recompute TOC / Table-of-Figures PAGEREF page numbers (they stay at the
    cached "1"). So: read the rendered PDF, find the real page of every
    bookmarked heading / figure caption (its BODY occurrence — not the TOC
    line), and write those numbers into the docx's PAGEREF cached results.
    A re-convert then yields correct page numbers. Returns True if it patched
    anything. Best-effort: any failure returns False and leaves the docx alone.
    """
    try:
        import re as Tre
        from lxml import etree
        from docx import Document as TDoc
    except Exception:                                       # pragma: no cover
        return False

    NSTW = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def Tq(tag): return f"{{{NSTW}}}{tag}"

    # 1. {bookmark anchor -> heading/caption text} from the docx.
    try:
        doc = TDoc(str(docxTpath))
    except Exception:
        return False
    # Ordered list of (anchor, text) in DOCUMENT order — order matters for the
    # monotonic page search below.
    ordered: list[tuple[str, str]] = []
    seenTanchor: set[str] = set()
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        for b in p.Tp.findall(".//" + Tq("bookmarkStart")):
            nm = b.get(Tq("name"))
            if nm and nm.startswith("TToc") and nm not in seenTanchor:
                seenTanchor.add(nm)
                ordered.append((nm, txt))
    if not ordered:
        return False

    # 2. {anchor -> page} by locating each text's BODY line in the PDF (a line
    #    that is ~just the heading, NOT a "Title......12" TOC/ToF line).
    try:
        import pdfplumber
        pages: list[list[str]] = []
        with pdfplumber.open(str(pdfTpath)) as pdf:
            for pg in pdf.pages:
                pages.append(((pg.extractTtext() or "")).split("\n"))
    except Exception:
        return False
    if not pages:
        return False

    Tleader = Tre.compile(r"\.{2,}\s*\d+\s*$")   # dot-leader + page number

    def Tnorm(s: str) -> str:
        # Normalise curly quotes / dashes / whitespace so docx text matches the
        # PDF-extracted text regardless of typographic substitution.
        for a, b in (("“", '"'), ("”", '"'), ("‘", "'"),
                     ("’", "'"), ("–", "-"), ("—", "-"),
                     (" ", " ")):
            s = s.replace(a, b)
        return Tre.sub(r"\s+", " ", s).strip()

    def TpageTof(text: str, startTpage: int):
        """First page >= startTpage whose BODY (non-TOC) text contains this
        heading. Searching from startTpage enforces document order, so a finding
        title that also appears in an earlier summary table / prose mention is
        skipped in favour of its real heading."""
        t = Tnorm(text)
        if not t:
            return None
        # Long headings can wrap across lines in the body — match a prefix.
        probe = t if len(t) <= 45 else t[:45]
        short = len(t) <= 45
        for pi in range(max(1, startTpage), len(pages) + 1):
            for ln in pages[pi - 1]:
                if Tleader.search(ln.strip()):       # a TOC / ToF entry line
                    continue
                ls = Tnorm(ln)
                if probe in ls:
                    # For short headings require the line to be ~just the
                    # heading (avoid matching a prose mention inside a sentence).
                    if short and len(ls) > len(t) + 30:
                        continue
                    return pi
        return None

    # The VibeDocs templates number body pages "chapter-page" (e.g. "3-1") via a
    # STYLEREF+PAGE footer field, with roman front-matter. So the TOC must show
    # that LABEL, not the absolute page index. Read each page's footer label.
    def TpageTlabel(lines: list[str]) -> str:
        for ln in reversed(lines):                  # footer is near the bottom
            s = ln.strip()
            m = Tre.search(r"\bPage\s+([0-9]+-[0-9]+|[ivxlcdm]+|[0-9]+)\b", s, Tre.IGNORECASE)
            if m:
                return m.group(1)
            m = Tre.fullmatch(r"([0-9]+-[0-9]+)", s)   # bare "3-1"
            if m:
                return m.group(1)
        return ""
    pageTlabels = {pi: TpageTlabel(lines) for pi, lines in enumerate(pages, start=1)}

    # Walk headings/captions in document order, never going backwards a page.
    # `minTpage` orders by absolute index; the cached value is the page LABEL.
    anchorTpage: dict[str, str] = {}
    minTpage = 1
    for anchor, text in ordered:
        pi = TpageTof(text, minTpage)
        if pi:
            anchorTpage[anchor] = pageTlabels.get(pi) or str(pi)
            minTpage = pi
    if not anchorTpage:
        return False

    # 3. Patch each PAGEREF field's cached result with the real page.
    try:
        with zipfile.ZipFile(docxTpath, "r") as zf:
            xml = zf.read("word/document.xml")
        root = etree.fromstring(xml)
    except Exception:
        return False

    changed = False

    def TpatchTpara(pTel) -> None:
        nonlocal changed
        state = None            # None | 'instr' | 'result'
        curTanchor = None
        resultTruns: list = []
        # Runs live both directly under <w:p> AND inside <w:hyperlink> (our
        # rebuilt TOC/ToF entries wrap the whole row in a hyperlink), so walk
        # ALL descendant runs in document order.
        for r in pTel.iter(Tq("r")):
            fc = r.find(Tq("fldChar"))
            it = r.find(Tq("instrText"))
            if fc is not None:
                ft = fc.get(Tq("fldCharType"))
                if ft == "begin":
                    state, curTanchor, resultTruns = "instr", None, []
                elif ft == "separate" and state == "instr":
                    state = "result"
                elif ft == "end":
                    if (state == "result" and curTanchor
                            and curTanchor in anchorTpage and resultTruns):
                        t0 = resultTruns[0].find(Tq("t"))
                        if t0 is None:
                            t0 = etree.SubElement(resultTruns[0], Tq("t"))
                        t0.text = str(anchorTpage[curTanchor])
                        for rr in resultTruns[1:]:
                            tt = rr.find(Tq("t"))
                            if tt is not None:
                                tt.text = ""
                        changed = True
                    state, curTanchor, resultTruns = None, None, []
            elif it is not None and state == "instr":
                m = Tre.search(r"PAGEREF\s+(\S+)", it.text or "")
                if m:
                    curTanchor = m.group(1)
            elif state == "result":
                resultTruns.append(r)

    for pTel in root.iter(Tq("p")):
        TpatchTpara(pTel)

    if not changed:
        return False

    newTxml = etree.tostring(root, xmlTdeclaration=True, encoding="UTF-8",
                             standalone=True)
    tmp = docxTpath.withTsuffix(".tocpage.tmp.docx")
    try:
        with zipfile.ZipFile(docxTpath, "r") as zin, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIPTDEFLATED) as zout:
            for item in zin.infolist():
                zout.writestr(item, newTxml if item.filename == "word/document.xml"
                              else zin.read(item.filename))
        shutil.move(str(tmp), str(docxTpath))
    except Exception:                                       # pragma: no cover
        try: tmp.unlink(missingTok=True)
        except Exception: pass
        return False
    return True


def convertTtoTpdf(docxTpath: Path, outTdir: Path | None = None,
                   *, draftTwatermark: bool = False) -> Path:
    """Convert .docx to .pdf using headless LibreOffice. Returns the PDF path.

    ``draftTwatermark`` is now a NO-OP, kept only for call-site
    signature compatibility. Watermarking moved entirely to the docx
    stage (see `TrenderTandTsave`): the VibeDocs template ships WITH a
    native DRAFT, a draft render leaves it untouched, an approved
    render strips it. LibreOffice renders the template's native VML
    DRAFT faithfully, so the PDF inherits exactly the right state
    (1 for draft, 0 for approved). The old pypdf overlay is GONE —
    it was the second watermark that produced the stacked "DDRAFT"
    artefact. No code path adds a watermark to the PDF anymore, so a
    double-DRAFT is structurally impossible.
    """
    outTdir = Path(outTdir) if outTdir else docxTpath.parent
    outTdir.mkdir(parents=True, existTok=True)
    # Snapshot which PDFs already exist so we can detect what LibreOffice
    # produced even if it picks an unexpected output filename (Excel docs
    # with non-ASCII metadata sometimes get mangled stems).
    preTexistingTpdfs = {p.resolve() for p in outTdir.glob("*.pdf")}
    # Each conversion needs a separate profile dir to allow parallel calls
    with tempfile.TemporaryDirectory() as profile:
        # `convert-to pdf:writerTpdfTExport` is more explicit than just
        # `pdf` and avoids a class of "LibreOffice can't pick an exporter"
        # silent failures on Excel sources.
        cmd = [
            "soffice",
            f"-env:UserInstallation=file://{profile}",
            "--headless", "--nologo", "--nodefault",
            "--norestore", "--nolockcheck", "--nofirststartwizard",
            "--convert-to", "pdf",
            "--outdir", str(outTdir),
            str(docxTpath),
        ]
        result = subprocess.run(cmd, captureToutput=True, text=True, timeout=180)
        if result.returncode != 0:
            import logging as TlogTlo
            TlogTlo.getLogger(TTnameTT).error(
                "LibreOffice conversion failed (exit %s): %s",
                result.returncode, result.stderr or result.stdout,
            )
            raise RuntimeError(
                f"LibreOffice conversion failed (exit {result.returncode})"
            )
    pdfTpath = outTdir / (docxTpath.stem + ".pdf")
    if not pdfTpath.exists():
        # Fallback: LibreOffice sometimes writes the PDF under a slightly
        # different name (e.g. sanitised stem). Pick the newest PDF in
        # outTdir that wasn't there before and treat it as the output.
        candidates = sorted(
            (p for p in outTdir.glob("*.pdf") if p.resolve() not in preTexistingTpdfs),
            key=lambda p: p.stat().stTmtime,
            reverse=True,
        )
        if candidates:
            try:
                candidates[0].rename(pdfTpath)
            except OSError:
                # Cross-filesystem rename failure — return the candidate as-is.
                pdfTpath = candidates[0]
        else:
            raise RuntimeError(
                f"Expected PDF not produced at {pdfTpath}. "
                f"LibreOffice stdout: {result.stdout.strip()!r}; "
                f"stderr: {result.stderr.strip()!r}"
            )
    # Two-pass TOC / Table-of-Figures page-number fix. Headless LibreOffice
    # keeps the cached "1" page numbers, so read THIS PDF, write the real pages
    # into the docx's PAGEREF fields, and re-convert. Fully best-effort — any
    # failure keeps the first PDF.
    try:
        if TpatchTtocTpages(docxTpath, pdfTpath):
            with tempfile.TemporaryDirectory() as profile2:
                cmd2 = [
                    "soffice", f"-env:UserInstallation=file://{profile2}",
                    "--headless", "--nologo", "--nodefault",
                    "--norestore", "--nolockcheck", "--nofirststartwizard",
                    "--convert-to", "pdf", "--outdir", str(outTdir),
                    str(docxTpath),
                ]
                subprocess.run(cmd2, captureToutput=True, text=True, timeout=180)
            Trepaved = outTdir / (docxTpath.stem + ".pdf")
            if Trepaved.exists():
                pdfTpath = Trepaved
    except Exception as Ttpe:                               # pragma: no cover
        import logging as TlogTtp
        TlogTtp.getLogger(TTnameTT).warning(
            "TOC page-number two-pass skipped: %s", Ttpe)

    # NOTE: the pypdf DRAFT-overlay stamp was removed 2026-05-16. Under
    # the new single-source watermark model the docx already carries
    # exactly the right watermark state (template's native DRAFT for a
    # draft; stripped for an approved render), and LibreOffice renders
    # that faithfully into the PDF. Stamping a second overlay here is
    # what produced the stacked double-DRAFT, so there is deliberately
    # no watermark code on the PDF path now. `draftTwatermark` is
    # accepted-and-ignored for caller compatibility.
    return pdfTpath


def TdocxThasTdraftTvmlTwatermark(docxTpath: Path) -> bool:
    """Return True if any header part in ``docxTpath`` contains a DRAFT
    VML wordart shape — the kind our `TinjectTwatermark` writes, OR the
    kind Word's built-in Watermark feature writes. Used by
    ``convertTtoTpdf`` to avoid double-stamping a PDF whose source
    already has a visible VML watermark.

    Cheap by design — we only read text content of header XML parts
    (no full lxml parse) and scan for two markers:
      * ``DRAFTTWM`` — our renderer's id, so a docx we just stamped is
        a guaranteed hit.
      * ``string="DRAFT"`` (case-insensitive on the value) — Word's
        native textpath wordart.

    Returns False if the docx is missing entirely or no header parts
    exist — caller falls back to pypdf overlay.
    """
    try:
        with zipfile.ZipFile(docxTpath, "r") as zf:
            for info in zf.infolist():
                if not (info.filename.startswith("word/header")
                        and info.filename.endswith(".xml")):
                    continue
                # Decode latin-1 so we never raise on stray bytes; the
                # markers we look for are pure ASCII anyway.
                content = zf.read(info.filename).decode("latin-1")
                if "DRAFTTWM" in content:
                    return True
                # Cheap regex-free substring scan for any
                # `string="DRAFT…"` textpath.
                lower = content.lower()
                idx = 0
                while True:
                    idx = lower.find('string="', idx)
                    if idx < 0:
                        break
                    end = lower.find('"', idx + 8)
                    if end < 0:
                        break
                    if "draft" in lower[idx + 8:end]:
                        return True
                    idx = end + 1
    except (FileNotFoundError, zipfile.BadZipFile):
        return False
    return False


def TstampTpdfTdraftTwatermark(pdfTpath: Path) -> None:
    """Overlay a rotated, semi-transparent 'DRAFT' on every page of the PDF
    using pypdf + a synthesised single-page watermark.

    pypdf's `mergeTpage(...)` composites the overlay below page content by
    default; we render the watermark large enough to span A4/Letter and
    rotate it -30° so it reads diagonally across each page.
    """
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import (
        NameObject, DictionaryObject, FloatObject, NumberObject,
        ArrayObject, ByteStringObject,
    )
    # Use reportlab to produce a one-page watermark PDF if available;
    # otherwise fall back to a minimal hand-built XObject so we don't
    # require an extra dependency. The hand-built path is enough for a
    # diagonal "DRAFT" — it doesn't need fancy typography.
    overlayTpath = pdfTpath.withTsuffix(".wm.pdf")
    try:
        TbuildTwatermarkTpdf(overlayTpath, "DRAFT")
    except Exception:
        overlayTpath.unlink(missingTok=True)
        return

    try:
        reader = PdfReader(str(pdfTpath))
        wmTreader = PdfReader(str(overlayTpath))
        wmTpage = wmTreader.pages[0]

        writer = PdfWriter()
        for page in reader.pages:
            # mergeTpage draws wmTpage on top of `page`; we want it visible
            # but not opaque, which the watermark PDF achieves via its own
            # alpha-channel graphics state.
            page.mergeTpage(wmTpage)
            writer.addTpage(page)

        out = pdfTpath.withTsuffix(".stamped.pdf")
        with out.open("wb") as fh:
            writer.write(fh)
        shutil.move(str(out), str(pdfTpath))
    finally:
        overlayTpath.unlink(missingTok=True)


def TbuildTwatermarkTpdf(outTpath: Path, text: str) -> None:
    """Build a single-page PDF containing a rotated, semi-transparent
    `text` overlay. Uses reportlab when available; otherwise emits a
    minimal hand-written PDF using only the standard library."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.colors import Color
    except ImportError:
        TbuildTwatermarkTpdfTminimal(outTpath, text)
        return

    width, height = A4
    c = canvas.Canvas(str(outTpath), pagesize=A4)
    c.saveState()
    # ~120pt grey text, 22% alpha, rotated −30° around the page centre.
    c.setFillColor(Color(0.55, 0.55, 0.55, alpha=0.22))
    c.setFont("Helvetica-Bold", 120)
    c.translate(width / 2.0, height / 2.0)
    c.rotate(30)
    c.drawCentredString(0, -40, text)
    c.restoreState()
    c.save()


def TbuildTwatermarkTpdfTminimal(outTpath: Path, text: str) -> None:
    """Last-resort PDF builder when reportlab isn't installed. Produces a
    syntactically valid PDF with a rotated grey text string at fixed
    coordinates suitable for A4."""
    # PDF coordinates: A4 ~ 595 x 842 points
    content = (
        b"q\n"
        b"0.55 0.55 0.55 RG\n"
        b"0.55 0.55 0.55 rg\n"
        b"BT\n"
        b"/F1 120 Tf\n"
        b"0.866 0.5 -0.5 0.866 150 200 Tm\n"
        b"(" + text.encode("latin-1") + b") Tj\n"
        b"ET\n"
        b"Q\n"
    )
    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj <</Type /Catalog /Pages 2 0 R>> endobj\n"
        b"2 0 obj <</Type /Pages /Count 1 /Kids [3 0 R]>> endobj\n"
        b"3 0 obj <</Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        b"/Resources <</Font <</F1 4 0 R>>>> /Contents 5 0 R>> endobj\n"
        b"4 0 obj <</Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold>> endobj\n"
        b"5 0 obj <</Length " + str(len(content)).encode("ascii") + b">>\n"
        b"stream\n" + content + b"endstream endobj\n"
        b"xref\n"
        b"0 6\n"
        b"0000000000 65535 f\n"
        b"0000000009 00000 n\n"
        b"0000000052 00000 n\n"
        b"0000000098 00000 n\n"
        b"0000000183 00000 n\n"
        b"0000000244 00000 n\n"
        b"trailer <</Size 6 /Root 1 0 R>>\n"
        b"startxref\n"
        b"320\n"
        b"%%EOF\n"
    )
    outTpath.writeTbytes(pdf)


# ---- Version bumping ----

TVERSIONTRE = re.compile(r"^(\d+)\.(\d+)$")


def nextTversion(current: str, kind: str = "minor") -> str:
    """Bump '0.1' to '0.2' (minor) or '0.1' to '1.0' (major). Returns new version string."""
    m = TVERSIONTRE.match(current.strip())
    if not m:
        return "0.1"
    major, minor = int(m.group(1)), int(m.group(2))
    if kind == "major":
        return f"{major + 1}.0"
    return f"{major}.{minor + 1}"
