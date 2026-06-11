"""
Generate sample generic Word templates for each VAPT type.

Run with:  docker compose exec app python -m app.gen_word_templates
            (or locally: cd backend && python -m app.gen_word_templates)

These are STARTER templates - replace them with your real custom-branded .docx
files later. Make sure your real templates contain the same Jinja placeholders so
the renderer keeps working.

Key placeholders the renderer fills in:
  {{ project.client_name }}
  {{ project.testing_window }}
  {{ project.scope_description }}
  {{ report.name }}
  {{ report.version }}
  {{ template.scope_of_work }}
  {{ template.methodology }}
  {{ details.executive_summary }}      # plus anything else in details
  {{ total_findings }}
  {{ severity_counts.Critical }}, .High, .Medium, .Low, .Informational

  Findings summary table:
      {%tr for f in findings %} ... {%tr endfor %}

  Per-finding detail block:
      {%p for f in findings %} ... {%p endfor %}

  Nmap discovered services (Infra only):
      {%tr for r in nmap_rows %} ... {%tr endfor %}
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from .config import settings


SEVERITY_COLORS = {
    "Critical": "C00000",
    "High": "E97132",
    "Medium": "FFC000",
    "Low": "00B050",
    "Informational": "4472C4",
}


def _set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_heading(doc, text, level=1, color="1F4E79"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18 - level * 2)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_section_break(doc):
    doc.add_paragraph()


def _add_kv_table(doc, pairs):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Light List Accent 1"
    table.autofit = False
    for i, (label, value) in enumerate(pairs):
        c1 = table.rows[i].cells[0]
        c2 = table.rows[i].cells[1]
        c1.width = Cm(5)
        c2.width = Cm(12)
        c1.text = label
        c2.text = value
        for run in c1.paragraphs[0].runs:
            run.bold = True
    return table


def _build_common(doc: Document, template_name: str):
    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("VibeDocs")
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor.from_string("7C5CFC")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(template_name)
    r.bold = True; r.font.size = Pt(24)

    for _ in range(4):
        doc.add_paragraph()

    _add_kv_table(doc, [
        ("Client", "{{ project.client_name }}"),
        ("Project", "{{ project.name }}"),
        ("Report", "{{ report.name }}"),
        ("Version", "v{{ report.version }}"),
        ("Sector", "{{ project.sector }}"),
        ("Testing window", "{{ project.testing_window }}"),
        ("Generated", "{{ generated_at }}"),
    ])

    doc.add_page_break()


def _build_report_details_section(doc):
    _add_heading(doc, "1. Report Details", 1)
    _add_kv_table(doc, [
        ("Tester(s)", "{{ details.tester_names | join(', ') }}"),
        ("User roles tested", "{{ details.user_roles_tested | join(', ') }}"),
        ("URLs / IPs in scope", "{{ project.scope_targets | join(', ') }}"),
        ("Client contact(s)", "{{ details.client_contacts | join(', ') }}"),
        ("Testing window", "{{ project.testing_window }}"),
    ])
    doc.add_paragraph()


def _build_scope_section(doc):
    _add_heading(doc, "2. Scope of Work", 1)
    doc.add_paragraph("{{ template.scope_of_work }}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Scope description (project-specific):").bold = True
    doc.add_paragraph("{{ project.scope_description }}")


def _build_methodology_section(doc):
    _add_heading(doc, "3. Methodology", 1)
    doc.add_paragraph("{{ template.methodology }}")


def _build_executive_summary(doc):
    _add_heading(doc, "4. Executive Summary", 1)
    doc.add_paragraph("{{ details.executive_summary }}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Total findings: {{ total_findings }} "
        "(Critical: {{ severity_counts.Critical }}, "
        "High: {{ severity_counts.High }}, "
        "Medium: {{ severity_counts.Medium }}, "
        "Low: {{ severity_counts.Low }}, "
        "Informational: {{ severity_counts.Informational }})"
    ).bold = True


def _build_findings_summary_table(doc):
    _add_heading(doc, "5. Summary of Findings", 1)
    table = doc.add_table(rows=2, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    headers = ["#", "Title", "Severity", "CVSS", "Status"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        _set_cell_bg(hdr[i], "1F4E79")
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string("FFFFFF")

    # Loop row using docxtpl {%tr ... %} syntax.
    row = table.rows[1].cells
    row[0].text = "{%tr for f in findings %}{{ f.index }}"
    row[1].text = "{{ f.title }}"
    row[2].text = "{{ f.severity }}"
    row[3].text = "{{ f.cvss_score }}"
    row[4].text = "{{ f.status }}{%tr endfor %}"


def _build_findings_detail_section(doc):
    _add_heading(doc, "6. Detailed Findings", 1)

    # We open a paragraph-loop using {%p for f in findings %}
    doc.add_paragraph("{%p for f in findings %}")

    # Per-finding header
    h = doc.add_paragraph()
    r = h.add_run("Finding {{ f.index }}: {{ f.title }}")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string("1F4E79")

    # Severity badge
    badge = doc.add_paragraph()
    badge.add_run("Severity: ").bold = True
    badge.add_run("{{ f.severity }}")
    badge.add_run("    CVSS: ").bold = True
    badge.add_run("{{ f.cvss_score }} ({{ f.cvss_vector }})")
    badge.add_run("    Status: ").bold = True
    badge.add_run("{{ f.status }}")

    _add_kv_table(doc, [
        ("Affected asset", "{{ f.affected_asset }}"),
    ])

    doc.add_paragraph()
    doc.add_paragraph().add_run("Description").bold = True
    doc.add_paragraph("{{ f.description }}")
    doc.add_paragraph().add_run("Impact").bold = True
    doc.add_paragraph("{{ f.impact }}")
    doc.add_paragraph().add_run("Steps to Reproduce").bold = True
    doc.add_paragraph("{{ f.poc_steps }}")

    doc.add_paragraph().add_run("Evidence").bold = True
    doc.add_paragraph("{% for img in f.screenshot_objs %}{{ img }}{% endfor %}")

    doc.add_paragraph().add_run("Remediation").bold = True
    doc.add_paragraph("{{ f.remediation }}")

    doc.add_paragraph().add_run("References").bold = True
    doc.add_paragraph("{{ f.references }}")

    # Retest / follow-up section - always shown so client + tester can fill on retest
    doc.add_paragraph()
    rt = doc.add_paragraph()
    rt.add_run("Retest / Follow-up").bold = True
    doc.add_paragraph("{{ f.retest_notes }}")
    doc.add_paragraph("{% for img in f.retest_objs %}{{ img }}{% endfor %}")
    doc.add_paragraph().add_run("Client statement").bold = True
    doc.add_paragraph("{{ f.client_statement }}")

    doc.add_paragraph("{%p endfor %}")


def _build_nmap_section(doc):
    _add_heading(doc, "Appendix A: Discovered Services (Nmap)", 1)
    doc.add_paragraph(
        "The following hosts, ports, and services were enumerated during the engagement."
    )
    table = doc.add_table(rows=2, cols=7)
    table.style = "Light Grid Accent 1"
    headers = ["Host", "Hostname", "Port", "Proto", "Service", "Product", "Version"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True

    row = table.rows[1].cells
    row[0].text = "{%tr for r in nmap_rows %}{{ r.host }}"
    row[1].text = "{{ r.hostname }}"
    row[2].text = "{{ r.port }}"
    row[3].text = "{{ r.protocol }}"
    row[4].text = "{{ r.service }}"
    row[5].text = "{{ r.product }}"
    row[6].text = "{{ r.version }}{%tr endfor %}"


def _build_template(template_name: str, *, include_nmap: bool = False) -> Document:
    doc = Document()
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _build_common(doc, template_name)
    _build_report_details_section(doc)
    _build_scope_section(doc)
    _build_methodology_section(doc)
    _build_executive_summary(doc)
    _build_findings_summary_table(doc)
    doc.add_page_break()
    _build_findings_detail_section(doc)
    if include_nmap:
        doc.add_page_break()
        _build_nmap_section(doc)
    return doc


# Mapping of generated output filename → VibeDocs source-template
# basename(s) inside `report-templates/`. Each list of candidates is
# tried in order; the first existing file wins. The transformer at
# `tools.build_vibedocs_wapt_template.transform` is structurally
# generic — every VibeDocs template ships with the same "Detailed
# Findings" → Heading 2 → labelled-SubHeadings layout, so the same
# walker handles WAPT, MAPT, NPT, NVA, SCR, CPT, etc.
_VIBEDOCS_SOURCES: dict[str, list[str]] = {
    "web_vapt_template.docx": [
        "Security Assessment XXX WAPT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX WAPT Draft Report v0.1 (Template) .docx",
    ],
    "infra_va_template.docx": [
        "Security Assessment XXX NVA Draft Report v0.1 (Template).docx",
        "Security Assessment XXX NVA Draft Report v0.1 (Template) .docx",
        # WVA falls in the same VA family but targets web — kept as a
        # tertiary fallback for environments where only WVA is present.
        "Security Assessment XXX WVA Draft Report v0.1 (Template).docx",
    ],
    "infra_vapt_template.docx": [
        "Security Assessment XXX NPT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX NPT Draft Report v0.1 (Template) .docx",
    ],
    "api_vapt_template.docx": [
        # No dedicated API draft in the bundle — closest layout is the
        # WAPT template (Web App). Consultants can swap it for an
        # API-specific master by uploading via the admin UI.
        "Security Assessment XXX WAPT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX WAPT Draft Report v0.1 (Template) .docx",
    ],
    "thick_client_pt_template.docx": [
        # Same situation — Thick Client uses the WAPT layout until a
        # dedicated TCAPT template ships.
        "Security Assessment XXX WAPT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX WAPT Draft Report v0.1 (Template) .docx",
    ],
    "mobile_pt_template.docx": [
        "Security Assessment XXX MAPT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX MAPT Draft Report v0.1 (Template) .docx",
    ],
    "source_code_review_template.docx": [
        "Security Assessment XXX SCR Draft Report v0.1 (Template).docx",
        "Security Assessment XXX SCR Draft Report v0.1 (Template) .docx",
    ],
    "aws_cloud_vapt_template.docx": [
        "Security Assessment XXX CPT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX CPT Draft Report v0.1 (Template) .docx",
    ],
    "azure_cloud_vapt_template.docx": [
        "Security Assessment XXX CPT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX CPT Draft Report v0.1 (Template) .docx",
    ],
    # April-2026 Kiosk PT bundle. Without this entry the Kiosk
    # template fell through to the simple plain-paragraph builder
    # because no `_VIBEDOCS_SOURCES` row pointed at the new file.
    "kiosk_pt_template.docx": [
        "Security Assessment XXX Kiosk-PT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX Kiosk-PT Draft Report v0.1 (Template) .docx",
    ],
    # April-2026 OT / ICS bundle. The OT source ships under the
    # `OT Draft Report` stem (no "ICS" in the filename) — keep both
    # the (Template)-suffixed and trailing-space variants so the
    # lookup survives the inconsistent naming the VibeDocs authoring
    # tool sometimes emits.
    "ot_vapt_template.docx": [
        "Security Assessment XXX OT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX OT Draft Report v0.1 (Template) .docx",
    ],
    # April-2026 Wi-Fi PT bundle. Note this source ships WITHOUT the
    # "(Template)" suffix that the rest of the family uses — list a
    # couple of variants so a rename in either direction still
    # resolves to the same canonical file.
    "wifi_pt_template.docx": [
        "Security Assessment XXX Wifi-PT Draft Report v0.1.docx",
        "Security Assessment XXX Wifi-PT Draft Report v0.1 (Template).docx",
        "Security Assessment XXX Wifi-PT Draft Report v0.1 (Template) .docx",
    ],
}


def _try_vibedocs_template(out_filename: str, out_path: Path) -> bool:
    """If a VibeDocs report template matching `out_filename` is bundled
    in `report-templates/`, transform it into a docxtpl-ready master
    and write to `out_path`. Returns True on success, False if no
    suitable source is present (caller falls back to the simple
    builder).

    The transformer at `tools.build_vibedocs_wapt_template.transform`
    is structurally generic — every VibeDocs template in the bundle
    follows the same "Detailed Findings → Heading 2 → labelled
    SubHeadings" layout, so the same walker handles WAPT / MAPT /
    NPT / NVA / SCR / CPT.
    """
    # Path resolution is delegated to `_vibedocs_source_path` so the
    # same dev / container layout fallback applies to both
    # `_try_vibedocs_template` and the diagnose endpoint.
    src = _vibedocs_source_path(out_filename)
    if src is None:
        return False
    try:
        from .tools.build_vibedocs_wapt_template import transform
        transform(src, out_path)
        return True
    except Exception as e:                                # pragma: no cover
        print(f"  VibeDocs template transform failed for "
              f"{out_filename} ({e}); falling back to simple builder.")
        return False


# Kept for backward compatibility — old callers (and any external
# scripts) referenced `_try_vibedocs_wapt` specifically. New code
# should use `_try_vibedocs_template`.
def _try_vibedocs_wapt(out_path: Path) -> bool:
    return _try_vibedocs_template("web_vapt_template.docx", out_path)


SPECS = [
    ("web_vapt_template.docx",          "Web Application VAPT Report",                  False),
    ("infra_va_template.docx",          "Infrastructure Vulnerability Assessment Report", True),
    ("infra_vapt_template.docx",        "Infrastructure VAPT Report",                    True),
    ("api_vapt_template.docx",          "API Penetration Test Report",                   False),
    ("thick_client_pt_template.docx",   "Thick Client Penetration Test Report",          False),
    ("mobile_pt_template.docx",         "Mobile Application Penetration Test Report",    False),
    ("source_code_review_template.docx","Source Code Review Report",                     False),
    ("aws_cloud_vapt_template.docx",    "AWS Cloud VAPT Report",                         False),
    ("azure_cloud_vapt_template.docx",  "Azure Cloud VAPT Report",                       False),
    # No dedicated VibeDocs source for these three VAPT types yet —
    # the simple builder produces a valid (if plainer) .docx so
    # consultants can still generate reports from the picker. Replace
    # with proper VibeDocs templates by dropping the matching file
    # into `report-templates/` and adding it to `_VIBEDOCS_SOURCES`.
    ("wifi_pt_template.docx",           "Wi-Fi Penetration Test Report",                 True),
    ("kiosk_pt_template.docx",          "Kiosk Penetration Test Report",                 False),
    ("ot_vapt_template.docx",           "OT / ICS VAPT Report",                          True),
]


def _candidate_template_roots() -> list[Path]:
    """Return every directory that might hold the bundled VibeDocs
    report-template sources, in priority order.

    Why a list rather than one canonical path: the source tree layout
    differs across the two environments this code has to handle.

    - Dev (running directly from the repo on the host):
        ``backend/app/gen_word_templates.py`` -> three parents up is
        the repo root, where ``report-templates/`` lives as a sibling
        of ``backend/``.

    - Container (image built by ``backend/Dockerfile``):
        ``app/`` is copied to ``/app/app`` and ``report-templates`` is
        copied (and bind-mounted in compose) to ``/app/report-templates``.
        So the templates are TWO parents up, not three.

    The earlier single-path version walked three parents in both
    environments, which resolved to ``/`` inside the container — and
    ``/report-templates`` doesn't exist, so the boot-time regenerator
    silently fell back to the simple plain-paragraph builder on every
    deploy. The fix is to try every plausible root and return the first
    one that contains a known VibeDocs source basename.
    """
    here = Path(__file__).resolve()
    candidates: list[Path] = []
    # Walk up to 4 parents — covers `backend/app/` (dev), `/app/app/`
    # (container), and any unforeseen layout. Stop at the filesystem
    # root so we don't accidentally append `/` twice on Windows / WSL.
    for n in (2, 3, 4):
        try:
            candidates.append(here.parents[n - 1])
        except IndexError:
            continue
    # De-duplicate while preserving order.
    seen: set[str] = set()
    uniq: list[Path] = []
    for c in candidates:
        s = str(c)
        if s not in seen:
            seen.add(s)
            uniq.append(c)
    return uniq


def _vibedocs_source_path(out_filename: str) -> Path | None:
    """Resolve the on-disk path of the VibeDocs source template that
    backs `out_filename`, or None if no bundled source exists.

    Tries every candidate in `_VIBEDOCS_SOURCES[out_filename]` against
    every root in `_candidate_template_roots()`. The first existing
    file wins. This handles BOTH the dev source-tree layout and the
    Docker container layout — see `_candidate_template_roots` for the
    detailed reason the path differs.
    """
    basenames = _VIBEDOCS_SOURCES.get(out_filename, [])
    if not basenames:
        return None
    for root in _candidate_template_roots():
        templates_dir = root / "report-templates"
        if not templates_dir.is_dir():
            continue
        for basename in basenames:
            p = templates_dir / basename
            if p.exists():
                return p
    return None


def main(force_overwrite_simple: bool = True):
    """Regenerate every default Word template under `TEMPLATE_DIR`.

    Resolution rule, per output file:
      1. If a VibeDocs source is bundled in `report-templates/`, run
         the structurally-generic transformer and write the result.
         This is the path that produces the VibeDocs-house-style
         (Heading 2 + 6-column risk table + labelled SubHeadings)
         layout the team expects.
      2. Otherwise, if the file doesn't exist, fall back to the
         simple plain-paragraph builder so the system still has
         SOMETHING to render.
      3. If `force_overwrite_simple` is True (the default), an
         existing file generated by the simple builder is replaced
         once a VibeDocs source becomes available. This is what
         makes the boot-time hook in `main.py` actually upgrade
         old deployments — without it, a stale simple template
         from a previous deploy would persist forever and the
         picker would keep showing the wrong layout.

    Admin-uploaded replacements use UUID-stamped filenames
    (see `templates.py:replace_template_docx`), so they live at a
    different path and survive this regeneration untouched.
    """
    out_dir = Path(settings.TEMPLATE_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)

    # WATERMARK MODEL (rev 2026-05-16): templates now DELIBERATELY KEEP
    # their native VibeDocs "DRAFT" wordart. We no longer strip it at
    # boot and no longer inject our own at render time. The single
    # source of truth lives in `docx_generator._render_and_save`:
    #   * draft render    → leave the template's native DRAFT (1 mark)
    #   * approved render → strip the native DRAFT (0 marks)
    # Stripping here at boot would defeat that — a draft report would
    # then have NO watermark at all. So `_strip()` is intentionally a
    # NO-OP now. It's kept (returning 0) purely so the existing call
    # sites + the `(stripped N DRAFT)` log line don't need touching;
    # the behavioural change is centralised to this one stub.
    strip_draft_watermarks = None  # type: ignore  # disabled by design

    def _strip(path: Path) -> int:   # noqa: ARG001 — kept for call-site parity
        # Intentionally does nothing — templates must retain their
        # native DRAFT watermark under the strip-on-approve model.
        return 0

    # Templates that have been manually curated / imported from a better
    # reference source. Do NOT auto-regenerate these if they already exist —
    # re-running the generic transformer would overwrite carefully placed
    # placeholders (client_owner, tester_names, header/footer tokens, etc.)
    # that the source VibeDocs .docx doesn't carry. A new deploy starts from
    # scratch (file absent) and gets a freshly-built version on first boot;
    # subsequent boots just preserve the curated copy.
    PRESERVE_IF_EXISTS = {
        # All of the below have been manually curated with correct Jinja2
        # placeholders (affected_asset, tester_names, client_name, report_date,
        # testing_window, version history). Adding to this set prevents the
        # VibeDocs-source transformer from overwriting the hand-placed tags
        # on every boot. A fresh deploy (file absent) still gets a generated
        # baseline on first boot; subsequent boots preserve the curated copy.
        "web_vapt_template.docx",
        "api_vapt_template.docx",
        "thick_client_pt_template.docx",
        "infra_vapt_template.docx",
        "infra_va_template.docx",
        "mobile_pt_template.docx",
        "kiosk_pt_template.docx",
        "wifi_pt_template.docx",
        "ot_vapt_template.docx",
        "aws_cloud_vapt_template.docx",
        "azure_cloud_vapt_template.docx",
        "source_code_review_template.docx",
    }

    summary: dict[str, str] = {}
    for fname, title, with_nmap in SPECS:
        path = out_dir / fname
        source = _vibedocs_source_path(fname)
        if source is not None:
            # If this is a manually-curated template and the file already
            # exists, preserve it — don't let the generic transformer clobber
            # the hand-placed placeholders.
            if fname in PRESERVE_IF_EXISTS and path.exists():
                print(f"  Preserving curated template (exists): {fname}")
                summary[fname] = "preserved"
                continue
            # VibeDocs source available — always regenerate. Existing
            # canonical file (if any) is overwritten because the
            # transformer produces a deterministic output: re-running
            # against the same source yields the same bytes.
            if _try_vibedocs_template(fname, path):
                removed = _strip(path)
                tail = f" (stripped {removed} DRAFT)" if removed else ""
                # Second pass: inject report-level metadata placeholders
                # (severity counts, document version/date/type, project name)
                # that the generic transformer doesn't add on its own.
                try:
                    from .tools.inject_template_metadata import process as _inject_meta
                    meta_result = _inject_meta(path)
                    if meta_result.get("changed"):
                        tail += f" + meta({len(meta_result['changed'])})"
                except Exception as _me:
                    print(f"    metadata injection failed for {fname}: {_me}")
                print(f"  Wrote (VibeDocs-derived from {source.name}): {path}{tail}")
                summary[fname] = f"vibedocs:{source.name}"
                continue
            # Transformer failed — fall through to simple builder.
        if path.exists() and not force_overwrite_simple:
            print(f"  Skipping (exists): {fname}")
            summary[fname] = "skipped"
            continue
        doc = _build_template(title, include_nmap=with_nmap)
        doc.save(str(path))
        # Belt-and-braces strip on the simple fallback path too — costs
        # nothing if the file already has no watermark.
        _strip(path)
        print(f"  Wrote (simple fallback): {path}")
        summary[fname] = "simple"

    # Final sweep: strip watermarks from any admin-uploaded UUID-stamped
    # files in TEMPLATE_DIR that the regenerator didn't touch above.
    # These are templates the admin replaced via the UI (filenames like
    # `web_vapt__<uuid8>.docx`); their headers can also carry baked-in
    # DRAFT wordart if the admin uploaded a stock VibeDocs template
    # before the upload-time stripper shipped. Idempotent — running it
    # twice on a clean file removes nothing the second time.
    if strip_draft_watermarks is not None:
        for path in sorted(out_dir.iterdir()):
            if not path.is_file() or path.suffix.lower() != ".docx":
                continue
            if path.name in {fname for fname, _, _ in SPECS}:
                continue   # canonical files were handled in the loop above
            removed = _strip(path)
            if removed:
                print(f"  Stripped {removed} DRAFT watermark(s) from {path.name}")

    print("Done. VibeDocs sources will be used automatically wherever "
          "a matching template is bundled in `report-templates/`.")
    return summary


if __name__ == "__main__":
    main()
