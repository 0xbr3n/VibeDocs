"""Parse finding statuses out of an edited Word report.

When a consultant downloads the generated .docx, edits a finding's Status line
(e.g. changes "Open" to "Closed" or "False Positive") and uploads the file back
into VibeDocs, this module reads each finding's Status value from the document so the
VibeDocs records can be updated to match — without the consultant re-entering them in
the app.

Strategy
--------
The detailed-findings section renders each finding as a heading (its title)
followed by labelled fields, including a "Status" label and its value. We:
  1. Read every body paragraph (the detail section lives in body paragraphs,
     not tables).
  2. Locate each VibeDocs finding's heading by matching its title (tolerant of the
     auto-numbering Word prepends, e.g. "3.1 ").
  3. Within each finding's region (heading -> next finding's heading) find the
     "Status" value and map it to the FindingStatus enum's canonical spelling.

Only STATUS is read — severity/score are intentionally left untouched (a finding
can be a "False Positive" while keeping its original severity).
"""
from __future__ import annotations

import re
from pathlib import Path

# Canonical FindingStatus values (must match models.FindingStatus values).
_CANON_STATUS = {
    "open": "Open",
    "closed": "Closed",
    "risk accepted": "Risk Accepted",
    "riskaccepted": "Risk Accepted",
    "false positive": "False Positive",
    "falsepositive": "False Positive",
    "n/a": "N/A",
    "na": "N/A",
    "not applicable": "N/A",
    "in remediation": "In Remediation",
    "inremediation": "In Remediation",
    "remediation": "In Remediation",
}


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def _canon_status(value: str) -> str | None:
    v = _norm(value)
    if not v:
        return None
    if v in _CANON_STATUS:
        return _CANON_STATUS[v]
    # tolerate trailing punctuation / extra words ("Closed." / "Status: Closed")
    v2 = re.sub(r"[^a-z/ ]", "", v).strip()
    return _CANON_STATUS.get(v2)


def _title_matches(para_norm: str, title_norm: str) -> bool:
    """True if a heading paragraph corresponds to a finding title. Tolerant of
    Word's auto-numbering prefix ("3.1 <title>") and trailing whitespace."""
    if not title_norm:
        return False
    if para_norm == title_norm:
        return True
    # "3.1 lack of security headers" endswith "lack of security headers"
    if para_norm.endswith(title_norm) and len(para_norm) - len(title_norm) <= 10:
        return True
    return False


def _find_status_in_region(region: list[str]) -> str | None:
    """Find the Status value in a finding's paragraph region."""
    for i, p in enumerate(region):
        pn = _norm(p)
        # "Status: Closed" on one line.
        m = re.match(r"^status\s*[:\-]\s*(.+)$", pn)
        if m:
            st = _canon_status(m.group(1))
            if st:
                return st
        # "Status" label on its own line -> value is the next non-empty line.
        if pn == "status":
            for q in region[i + 1:]:
                if _norm(q):
                    st = _canon_status(q)
                    if st:
                        return st
                    break   # next non-empty wasn't a status -> stop
    return None


# Canonical Severity enum values, keyed by the display word seen in the report
# (including the "Info" house-style abbreviation).
_CANON_SEV = {
    "critical": "Critical",
    "high": "High",
    "medium": "Medium",
    "low": "Low",
    "info": "Informational",
    "informational": "Informational",
}


def _canon_severity(value: str) -> str | None:
    return _CANON_SEV.get(_norm(value))


def parse_finding_severities(docx_path: Path,
                             findings: list[tuple[int, str]]) -> dict[int, str]:
    """Read each finding's severity out of the report, returning
    ``{finding_id: canonical_severity}``.

    Two strategies, whichever yields matches:
      1. The per-finding DETAIL table (header "CVSS … Risk Rating") whose first
         data cell is the severity. Correlated to its finding by the heading
         that precedes it (document order). This is the findings-loop output, so
         it is always populated.
      2. The Risk Register summary table (Risk + Description columns), matched on
         the Description = title.
    """
    from docx import Document
    doc = Document(str(docx_path))
    result = _sev_from_detail_tables(doc, findings)
    if not result:
        result = _sev_from_risk_register(doc, findings)
    return result


def _sev_from_detail_tables(doc, findings: list[tuple[int, str]]) -> dict[int, str]:
    """Walk the body in order; whenever a finding heading is seen, the next
    'CVSS … Risk Rating' table's first data cell gives that finding's severity.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    title_lookup = [(fid, _norm(t)) for fid, t in findings]
    result: dict[int, str] = {}
    current_fid: int | None = None
    used: set[int] = set()

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split('}', 1)[-1]
        if tag == "p":
            text = _norm(Paragraph(child, doc).text)
            if not text:
                continue
            for fid, tn in title_lookup:
                if fid in used or not tn:
                    continue
                if text == tn or (text.endswith(tn) and len(text) - len(tn) <= 10):
                    current_fid = fid
                    break
        elif tag == "tbl" and current_fid is not None:
            tbl = Table(child, doc)
            if not tbl.rows:
                continue
            hdr0 = _norm(tbl.rows[0].cells[0].text)
            if "risk rating" in hdr0 and len(tbl.rows) > 1:
                sev = _canon_severity(tbl.rows[1].cells[0].text)
                if sev:
                    result[current_fid] = sev
                    used.add(current_fid)
                    current_fid = None
    return result


def _sev_from_risk_register(doc, findings: list[tuple[int, str]]) -> dict[int, str]:
    remaining = list(findings)
    result: dict[int, str] = {}
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [_norm(c.text) for c in table.rows[0].cells]
        if "risk" not in headers or "description" not in headers:
            continue
        ri, di = headers.index("risk"), headers.index("description")
        for row in table.rows[1:]:
            cells = row.cells
            if max(ri, di) >= len(cells):
                continue
            sev = _canon_severity(cells[ri].text)
            if not sev:
                continue
            title_norm = _norm(cells[di].text)
            if not title_norm:
                continue
            for idx, (fid, title) in enumerate(remaining):
                tn = _norm(title)
                if tn and (title_norm == tn or title_norm.startswith(tn) or tn in title_norm):
                    result[fid] = sev
                    remaining.pop(idx)
                    break
    return result


def parse_finding_statuses(docx_path: Path,
                           findings: list[tuple[int, str]]) -> dict[int, str]:
    """Return ``{finding_id: canonical_status}`` for findings whose Status was
    found in the document. ``findings`` is ``[(id, title), ...]`` in report order.
    """
    from docx import Document
    doc = Document(str(docx_path))
    paras = [(p.text or "").strip() for p in doc.paragraphs]
    n = len(paras)

    # Locate each finding's heading position, scanning forward in report order
    # so duplicate titles map to successive sections.
    positions: list[tuple[int, int | None]] = []
    search_from = 0
    for fid, title in findings:
        tn = _norm(title)
        pos = None
        for i in range(search_from, n):
            if _title_matches(_norm(paras[i]), tn):
                pos = i
                break
        positions.append((fid, pos))
        if pos is not None:
            search_from = pos + 1

    result: dict[int, str] = {}
    for idx, (fid, pos) in enumerate(positions):
        if pos is None:
            continue
        end = n
        for _, later_pos in positions[idx + 1:]:
            if later_pos is not None:
                end = later_pos
                break
        status = _find_status_in_region(paras[pos:end])
        if status:
            result[fid] = status
    return result
