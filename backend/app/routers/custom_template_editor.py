"""
Custom Template Editor API

Endpoints for uploading client Word templates, marking placeholders visually,
and managing the approval workflow.
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse
from sqlalchemy.orm import Session
from pathlib import Path
from docx import Document
import re
import hashlib
import shutil
import logging
from datetime import datetime

logger = logging.getLogger(__name__)
from typing import Optional

from ..database import get_db
from ..models import CustomTemplate, TemplatePlaceholder, User, TemplateStatus, AuditLog, Role
from ..auth import get_current_user, require_admin
from ..config import settings

router = APIRouter(prefix="/api/templates/custom", tags=["custom_templates"])

# Persisted volume mount in docker-compose: /data/uploads is a named volume,
# but /data/custom_templates is NOT — files written there are lost on every
# container rebuild. Anchor under UPLOAD_DIR so they survive.
CUSTOM_TEMPLATES_DIR = Path(settings.UPLOAD_DIR) / "custom_templates"
CUSTOM_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

_MAX_CUSTOM_TEMPLATE_BYTES = 30 * 1024 * 1024   # 30 MB

# Legacy ephemeral location used by earlier builds. Kept as a read-only
# fallback so previously uploaded files still resolve until the user
# re-uploads them.
_LEGACY_CUSTOM_TEMPLATES_DIR = Path("/data/custom_templates")


def _resolve_docx_path(template: CustomTemplate) -> Optional[Path]:
    """Find the actual .docx for a template, tolerating earlier installs that
    stored files under /data/custom_templates (no volume mount) or that have
    since been moved into the persisted UPLOAD_DIR location.
    """
    candidates: list[Path] = []
    if template.docx_path:
        candidates.append(Path(template.docx_path))
        # Fallback: same basename, but under the new persisted directory.
        candidates.append(CUSTOM_TEMPLATES_DIR / Path(template.docx_path).name)
        # Fallback: same basename, under the legacy directory.
        candidates.append(_LEGACY_CUSTOM_TEMPLATES_DIR / Path(template.docx_path).name)
    for c in candidates:
        if c.exists():
            return c
    return None


# ===== Template Upload & Management =====

@router.post("/upload")
async def upload_custom_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    description: Optional[str] = Form(None),
    template_type: str = Form("web_vapt"),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """
    Upload a new custom Word template.
    Returns template_id for the editor UI.
    """
    # Validate file type
    if not (file.filename or "").endswith('.docx'):
        raise HTTPException(400, "Only .docx files accepted")
    
    # Save the uploaded file first
    from datetime import datetime
    _chunks: list[bytes] = []
    _total = 0
    while True:
        _chunk = await file.read(65536)
        if not _chunk:
            break
        _total += len(_chunk)
        if _total > _MAX_CUSTOM_TEMPLATE_BYTES:
            raise HTTPException(413, "Template file exceeds the 30 MB upload limit.")
        _chunks.append(_chunk)
    file_content = b"".join(_chunks)
    file_hash = hashlib.md5(file_content).hexdigest()
    # Sanitise the filename — strip directory separators and any non-safe
    # characters so the on-disk path never escapes CUSTOM_TEMPLATES_DIR.
    raw_name = Path(file.filename).name
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in raw_name)[:160]
    safe_filename = f"{datetime.utcnow().timestamp()}_{safe_name}"
    file_path = CUSTOM_TEMPLATES_DIR / safe_filename
    
    # Write file to disk
    with open(file_path, 'wb') as f:
        f.write(file_content)

    # Inject Jinja2 expressions so report-details values populate correctly.
    try:
        from ..tools.inject_jinja2_into_templates import process_template
        process_template(file_path)
    except Exception as _inj_err:
        logger.warning("Jinja2 injection failed on template editor upload: %s", _inj_err)

    # ============================================
    # INSERT THE AUTO-DETECTION CODE HERE ↓↓↓
    # ============================================
    
    # AUTO-DETECT PLACEHOLDERS from the Word document.
    # We accept BOTH styles so the user is never caught between regimes:
    #   • Legacy flat uppercase tokens: {{CLIENT_NAME}}
    #   • Canonical Jinja dotted paths used by the generator:
    #     {{ project.client_name }}, {{ sections.executive_summary }},
    #     {{ severity_counts.Critical }}, {{ generated_at }}
    # The captured key is normalised to lowercase + dots so the picker
    # downstream can render both styles uniformly.
    _PLACEHOLDER_RE = re.compile(
        r"\{\{\s*([A-Za-z_][A-Za-z0-9_.]*)\s*\}\}"
    )
    try:
        doc = Document(file_path)
        placeholders_found: dict[str, dict] = {}

        def _record(key: str, location: dict, sample: str) -> None:
            norm = key.strip().lower()
            entry = placeholders_found.setdefault(
                norm, {"locations": [], "sample_text": sample[:100]}
            )
            entry["locations"].append(location)

        # Search all paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            for match in _PLACEHOLDER_RE.findall(para.text):
                _record(match,
                        {"paragraph": para_idx, "text": para.text[:100]},
                        para.text)

        # Also search tables
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for match in _PLACEHOLDER_RE.findall(cell.text):
                        _record(match,
                                {"table": table_idx, "text": cell.text[:100]},
                                cell.text)

    except Exception as e:
        logger.warning("Could not auto-detect placeholders: %s", e)
        placeholders_found = {}
    
    # ============================================
    # END OF AUTO-DETECTION CODE ↑↑↑
    # ============================================
    
    # Create template record with auto-detected placeholders.
    # Local/standalone mode: the single built-in user is also the reviewer,
    # so its templates are born approved + public (immediately usable, no
    # review queue). Normal admins still upload as `draft` and go through
    # the review flow — only the `is_local` singleton auto-approves.
    _local = bool(getattr(user, "is_local", False))
    template = CustomTemplate(
        name=name,
        description=description,
        template_type=template_type,
        docx_filename=file.filename,
        docx_path=str(file_path),
        docx_hash=file_hash,
        uploaded_by_id=user.id,
        placeholder_map=placeholders_found,
        status=(TemplateStatus.approved if _local else TemplateStatus.draft),
        is_public=_local,
    )
    db.add(template)
    db.flush()  # Get the ID
    db.commit()
    
    return {
        "template_id": template.id,
        "name": template.name,
        "redirect_url": f"/templates/edit/{template.id}"
    }


@router.get("/pending")
async def list_pending_templates(
    db: Session = Depends(get_db),
    user: User = Depends(require_admin)
):
    """List all templates pending admin review."""
    templates = db.query(CustomTemplate).filter(
        CustomTemplate.status == TemplateStatus.pending_review
    ).order_by(CustomTemplate.created_at.desc()).all()

    return [
        {
            "id": t.id,
            "name": t.name,
            "description": t.description,
            "template_type": t.template_type,
            "uploaded_by": t.uploaded_by.username,
            "created_at": t.created_at.isoformat(),
            "placeholder_count": len(t.placeholder_map or {})
        }
        for t in templates
    ]


@router.get("/{template_id}")
async def get_template(
    template_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Get template metadata and placeholder mappings."""
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404, "Template not found")
    
    # Check access
    if template.uploaded_by_id != user.id and user.role != Role.admin and not template.is_public:
        raise HTTPException(403, "Access denied")
    
    resolved = _resolve_docx_path(template)
    return {
        "id": template.id,
        "name": template.name,
        "description": template.description,
        "template_type": template.template_type,
        "status": template.status.value,
        "placeholder_map": template.placeholder_map or {},
        "is_public": template.is_public,
        "uploaded_by": template.uploaded_by.username,
        "created_at": template.created_at.isoformat(),
        "docx_filename": template.docx_filename,
        # True when the .docx is gone from disk (typical after a container
        # rebuild that wiped /data/custom_templates). UI uses this to show
        # a "re-upload" prompt instead of a useless preview error.
        "file_missing": resolved is None,
    }


@router.get("/{template_id}/diagnose")
async def diagnose_template(
    template_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Inspect a custom .docx template and surface common authoring
    mistakes WITHOUT having to run a full render preview.

    Reports:
      • `placeholders` — every `{{ … }}` and `{% … %}` token found in
        the document's main body + headers + footers, de-duplicated.
        The consultant can compare this against the canonical list in
        the placeholder modal to spot typos.
      • `raw_blocks_present` — True if any `{% raw %}` or
        `{% endraw %}` markers were found. These typically come from
        users copy-pasting the placeholder docs and don't belong in
        a Word template (the renderer strips them automatically; this
        is just a hint).
      • `looks_split` — True when a `{{` opener was detected with no
        matching `}}` close on the same text run. Indicates Word has
        fragmented a placeholder; the renderer's `_fix_split_jinja_tags`
        will glue these back at render time.
    """
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404, "Template not found")
    if template.uploaded_by_id != user.id and user.role != Role.admin:
        raise HTTPException(403, "Access denied")
    src = _resolve_docx_path(template)
    if not src:
        raise HTTPException(410, "Template file is missing on disk.")

    import zipfile, re as _re
    placeholders: set[str] = set()
    raw_open = raw_close = 0
    looks_split = False
    PLACE_RE = _re.compile(r"\{\{\s*([A-Za-z_][A-Za-z0-9_.]*)\s*\}\}")
    OPEN_RE  = _re.compile(r"\{\{")
    CLOSE_RE = _re.compile(r"\}\}")
    try:
        with zipfile.ZipFile(str(src)) as zf:
            for name in zf.namelist():
                if not (name == "word/document.xml" or
                        name.startswith("word/header") or
                        name.startswith("word/footer")):
                    continue
                xml = zf.read(name).decode("utf-8", errors="replace")
                # Strip XML tags so Word's run-splitting doesn't hide
                # the literal text we're looking for.
                text = _re.sub(r"<[^>]+>", "", xml)
                for m in PLACE_RE.finditer(text):
                    placeholders.add(m.group(1))
                raw_open  += len(_re.findall(r"\{%\s*raw\s*%\}",    xml, _re.I))
                raw_close += len(_re.findall(r"\{%\s*endraw\s*%\}", xml, _re.I))
                opens  = len(OPEN_RE.findall(xml))
                closes = len(CLOSE_RE.findall(xml))
                if opens != closes:
                    looks_split = True
    except Exception as e:
        logger.exception("Could not inspect template %s: %s", template_id, e)
        raise HTTPException(500, "Could not inspect the template file.")

    return {
        "template_id": template.id,
        "filename":    template.docx_filename,
        "placeholders": sorted(placeholders),
        "placeholder_count": len(placeholders),
        "raw_blocks_present":   bool(raw_open or raw_close),
        "raw_open_count":  raw_open,
        "raw_close_count": raw_close,
        "looks_split": looks_split,
        "notes": [
            *(["Stripped automatically at render time: {% raw %} / "
               "{% endraw %} markers found. Recommended: open the .docx "
               "and delete them — they prevent the placeholders inside "
               "from being substituted."]
              if raw_open or raw_close else []),
            *(["Some {{ ... }} or {% ... %} openers don't have matching "
               "closers on the same text run. This is usually Word "
               "splitting a tag across spell-check markers; the "
               "renderer auto-glues these. If your template still "
               "doesn't render the placeholder, try re-typing it from "
               "scratch in Word."]
              if looks_split else []),
        ],
    }


@router.get("/{template_id}/download")
async def download_template(
    template_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Download the Word template file."""
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404)

    if template.uploaded_by_id != user.id and user.role != Role.admin and not template.is_public:
        raise HTTPException(403, "Access denied")

    file_path = _resolve_docx_path(template)
    if not file_path:
        raise HTTPException(
            404,
            "Template file is no longer on disk. The container's ephemeral "
            "storage was cleared on a redeploy — re-upload the .docx to restore it.",
        )

    # If we resolved to a fallback path, persist it so the next request is a
    # direct hit instead of paying the fallback cost again.
    if str(file_path) != template.docx_path:
        template.docx_path = str(file_path)
        db.commit()

    return FileResponse(
        path=file_path,
        filename=template.docx_filename or f"template_{template_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.post("/{template_id}/placeholders")
async def save_placeholder_mappings(
    template_id: int,
    mappings: dict,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """
    Save placeholder mappings from the visual editor.
    
    Expected format:
    {
        "client_name": {"paragraph": 3, "text_sample": "Client Name: __"},
        "project_title": {"paragraph": 5, "text_sample": "Project: __"},
        "findings_table": {"paragraph": 12, "text_sample": "Findings:"}
    }
    """
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404)
    
    if template.uploaded_by_id != user.id and user.role != Role.admin:
        raise HTTPException(403)

    # Validate required placeholders
    required_keys = ["client_name", "project_title", "findings_table", "executive_summary"]
    missing = [k for k in required_keys if k not in mappings]
    if missing:
        return {
            "success": False,
            "error": f"Missing required placeholders: {', '.join(missing)}",
            "missing_keys": missing
        }
    
    # Save mappings
    template.placeholder_map = mappings
    template.updated_at = datetime.utcnow()
    db.commit()
    
    return {
        "success": True,
        "message": "Placeholder mappings saved",
        "template_id": template_id
    }


@router.post("/{template_id}/reupload")
async def reupload_template_file(
    template_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Restore the .docx for an existing template whose file was lost (e.g. a
    redeploy wiped /data/custom_templates). Keeps the existing template row,
    its placeholder mappings, and the approval status — only the binary is
    replaced. Caller must be the original uploader or an admin."""
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404, "Template not found")
    if template.uploaded_by_id != user.id and user.role != Role.admin:
        raise HTTPException(403, "Only the original uploader or an admin can re-upload")
    if not (file.filename or "").endswith(".docx"):
        raise HTTPException(400, "Only .docx files accepted")

    _chunks: list[bytes] = []
    _total = 0
    while True:
        _chunk = await file.read(65536)
        if not _chunk:
            break
        _total += len(_chunk)
        if _total > _MAX_CUSTOM_TEMPLATE_BYTES:
            raise HTTPException(413, "Template file exceeds the 30 MB upload limit.")
        _chunks.append(_chunk)
    file_content = b"".join(_chunks)
    file_hash = hashlib.md5(file_content).hexdigest()
    raw_name = Path(file.filename).name
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in raw_name)[:160]
    safe_filename = f"{datetime.utcnow().timestamp()}_{safe_name}"
    file_path = CUSTOM_TEMPLATES_DIR / safe_filename
    with open(file_path, "wb") as f:
        f.write(file_content)

    # Inject Jinja2 expressions so report-details values populate correctly.
    try:
        from ..tools.inject_jinja2_into_templates import process_template
        process_template(file_path)
    except Exception as _inj_err:
        logger.warning("Jinja2 injection failed on template editor reupload: %s", _inj_err)

    template.docx_path = str(file_path)
    template.docx_filename = file.filename
    template.docx_hash = file_hash
    template.updated_at = datetime.utcnow()
    db.commit()

    return {
        "ok": True,
        "template_id": template.id,
        "message": "Template file restored",
    }


@router.post("/{template_id}/submit")
async def submit_for_review(
    template_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Submit template for admin review."""
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404)
    
    if template.uploaded_by_id != user.id:
        raise HTTPException(403)
    
    if template.status != TemplateStatus.draft and template.status != TemplateStatus.rejected:
        raise HTTPException(400, "Template already submitted or approved")
    
    # Validate has placeholder mappings
    if not template.placeholder_map:
        raise HTTPException(400, "Cannot submit: No placeholders marked yet")
    
    template.status = TemplateStatus.pending_review
    template.updated_at = datetime.utcnow()
    # AuditLog row is the source-of-truth for the in-app notification
    # bell — see routers/notifications.py. We log the template name +
    # type so the dropdown can render a useful one-liner without a
    # second DB hit.
    db.add(AuditLog(
        actor_id=user.id,
        action="template.review.requested",
        object_type="custom_template",
        object_id=template.id,
        detail={
            "template_name": template.name,
            "template_type": template.template_type,
        },
    ))
    db.commit()

    return {
        "success": True,
        "message": "Template submitted for review",
        "status": "pending_review"
    }


# ===== Admin Approval Workflow =====

@router.post("/{template_id}/review")
async def review_template(
    template_id: int,
    action: str = Form(...),  # "approve" or "reject"
    notes: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    user: User = Depends(require_admin)
):
    """Admin approves or rejects a template."""
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404)
    
    if template.status != TemplateStatus.pending_review:
        raise HTTPException(400, "Template not pending review")
    
    if action == "approve":
        template.status = TemplateStatus.approved
        template.is_public = True  # Make available to all
    elif action == "reject":
        template.status = TemplateStatus.rejected
    else:
        raise HTTPException(400, "Invalid action")
    
    template.reviewed_by_id = user.id
    template.reviewed_at = datetime.utcnow()
    template.review_notes = notes
    # Notify the original uploader that their template was decided on.
    # The submitter id lives on the template row; the notifications
    # router treats this row as a one-recipient event keyed on that id.
    db.add(AuditLog(
        actor_id=user.id,
        action="template.review.decided",
        object_type="custom_template",
        object_id=template.id,
        detail={
            "template_name": template.name,
            "decision": template.status.value,    # "approved" | "rejected"
            "notes": notes or "",
            "submitter_id": template.uploaded_by_id,
        },
    ))
    db.commit()

    # Best-effort approval email to the uploader. Rejections still
    # surface in-app via the audit notification + the rejected status
    # on the template card; we only ship an email for the happy path.
    # Routed through `notify_user` so the uploader's email-opt-out
    # applies and self-trigger is short-circuited.
    if action == "approve" and template.uploaded_by_id:
        from ..services.notifier import notify_user
        from ..services.url_helpers import absolute_url
        uploader = db.get(User, template.uploaded_by_id)
        notify_user(
            db, uploader, "custom_template_approved", {
                "user": uploader,
                "reviewer_username": user.username,
                "template_name": template.name,
                "review_notes": notes or "",
                "template_url": absolute_url(
                    f"/templates/edit/{template.id}"),
            },
            actor_user_id=user.id,
        )

    return {
        "success": True,
        "status": template.status.value,
        "message": f"Template {action}d successfully"
    }


@router.delete("/{template_id}")
async def delete_template(
    template_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Delete a template (owner or admin only)."""
    template = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404)
    
    if template.uploaded_by_id != user.id and user.role != Role.admin:
        raise HTTPException(403)

    # Delete file (silent no-op if the disk file is already gone — the DB row
    # is authoritative, and we'd rather succeed than block deletion on a
    # missing artefact from an earlier ephemeral install).
    file_path = _resolve_docx_path(template)
    if file_path and file_path.exists():
        try:
            file_path.unlink()
        except OSError:
            pass

    db.delete(template)
    db.commit()
    
    return {"success": True, "message": "Template deleted"}


# ===== Template List (for selection in reports) =====

@router.get("/")
async def list_available_templates(
    template_type: Optional[str] = None,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """List templates: user's own drafts + approved public templates.

    Returns enough metadata for the card UI to render a colourful status
    badge, the uploader's display name (so consultants can tell apart
    multiple copies of the same template uploaded by different people),
    and a placeholder-count chip. Previously the response was missing
    `status` and `placeholder_map`, which is why cards rendered an
    "UNDEFINED" status pill and always showed "No placeholders marked yet".
    """
    from sqlalchemy import or_, and_

    query = db.query(CustomTemplate).filter(
        or_(
            CustomTemplate.uploaded_by_id == user.id,  # User's own templates (any status)
            and_(
                CustomTemplate.status == TemplateStatus.approved,  # Public approved templates
                CustomTemplate.is_public == True
            )
        )
    )

    if template_type:
        query = query.filter(CustomTemplate.template_type == template_type)

    templates = query.order_by(CustomTemplate.created_at.desc()).all()

    out: list[dict] = []
    for t in templates:
        up = t.uploaded_by
        placeholder_count = len(t.placeholder_map or {})
        out.append({
            "id": t.id,
            "name": t.name,
            "description": t.description,
            "template_type": t.template_type,
            "status": t.status.value if t.status else "draft",
            "is_public": bool(t.is_public),
            "placeholder_count": placeholder_count,
            "uploaded_by_id": t.uploaded_by_id,
            "uploaded_by": up.username if up else None,
            "uploaded_by_full_name": (up.full_name or up.username) if up else None,
            "uploaded_by_role": up.role.value if (up and up.role) else None,
            "is_mine": t.uploaded_by_id == user.id,
            "created_at": t.created_at.isoformat() if t.created_at else None,
            "updated_at": t.updated_at.isoformat() if t.updated_at else None,
        })
    return out
