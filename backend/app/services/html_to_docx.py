"""
Convert sanitised HTML (from the Quill rich-text editor) into a docxtpl
Subdoc so that bold/italic/lists/colours/code blocks survive the trip
from the browser into the generated Word document.

Coverage:
  * Block-level: <p>, <div>, <br>, <h1>-<h6>, <blockquote>, <ul>, <ol>, <li>,
                 <pre>
  * Inline:      <b>, <strong>, <i>, <em>, <u>, <s>, <strike>, <code>,
                 <span style="color:..."> / "background-color:..." /
                 "font-weight:..." / "text-decoration:underline"
  * Hyperlinks:  rendered as blue underlined text + a parenthetical URL.
                 (Native XML hyperlinks need rId injection which is fragile
                 inside Subdocs; visible URLs satisfy the same use case for
                 a printed report.)

Anything outside the allow-list above is rendered as plain text.

The renderer is *forgiving*: malformed HTML, missing closing tags, etc.
won't raise — we just stop processing that branch. The sanitiser layer is
responsible for shape; this layer is responsible for rendering.
"""
from __future__ import annotations
import base64
import io
import logging
import re
import html.parser
from typing import Optional

from docx.shared import RGBColor, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Fixed width for every inline-pasted screenshot (matches docx_generator's
# SCREENSHOT_WIDTH_CM). Images are centred along with their captions.
_INLINE_IMG_WIDTH_CM = 18.47

log = logging.getLogger(__name__)

# Recognises the Quill-style data URL Quill emits when a screenshot is
# pasted into the editor: `data:image/png;base64,iVBOR…`. Restricted to
# image MIME types — the sanitizer rejects everything else, but we
# guard here too in case someone calls the renderer directly.
_DATA_IMG_RE = re.compile(
    r"^data:image/(?P<ext>png|jpe?g|gif|webp|bmp);base64,(?P<b64>[A-Za-z0-9+/=\s]+)$",
    re.IGNORECASE,
)


BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
              "blockquote", "pre", "li"}
LIST_TAGS = {"ul", "ol"}
INLINE_TAGS = {"b", "strong", "i", "em", "u", "s", "strike",
               "code", "kbd", "samp", "span", "a", "sub", "sup", "mark"}


# Inline screenshot reference token. Consultants type
# `[Screenshot 2]` inside the Steps to Reproduce / Observations
# rich-text fields to embed the Nth uploaded screenshot at that exact
# spot in the rendered document. The number is 1-based and refers to
# the index in `f.screenshots` for the surrounding finding.
#
# We accept a few common variants so the consultant can type whatever
# feels natural — case-insensitive "screenshot" / "screen shot", with
# or without surrounding whitespace, optional trailing colon. The
# capture group `n` is the 1-based index.
_SCREENSHOT_REF_RE = re.compile(
    r"\[\s*screen\s*shot\s+(?P<n>\d+)\s*\]",
    re.IGNORECASE,
)


def _apply_code_para_style(para) -> None:
    """Set gray background shading and left/right indent on a code-block paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    for old in pPr.findall(qn('w:ind')):
        pPr.remove(old)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')   # 0.25 inch left indent
    ind.set(qn('w:right'), '360')  # 0.25 inch right indent
    pPr.append(ind)


def html_to_subdoc(tpl, html_text: str,
                   inline_images: list[str] | None = None,
                   fig_start: int = 0,
                   fig_prefix: str = ""):
    """Render `html_text` into a fresh Subdoc bound to `tpl`. Returns
    the Subdoc which the caller can drop into the docxtpl context.

    When `inline_images` is provided, any `[Screenshot N]` text token
    encountered while parsing is replaced with the corresponding image
    from the list (1-based). Missing / out-of-range references keep
    the literal token text so the consultant can see they have a
    dangling reference in the deliverable rather than the renderer
    silently swallowing it.

    When `fig_start` > 0, each inline pasted `<img>` gets a "Figure N"
    caption paragraph inserted immediately after it, numbered from
    `fig_start` upward. Pass 0 (default) to suppress captions.
    """
    sd = tpl.new_subdoc()
    if not html_text or not html_text.strip():
        sd.add_paragraph("")
        return sd
    renderer = _DocRenderer(sd, inline_images=inline_images, fig_start=fig_start,
                            fig_prefix=fig_prefix)
    parser = _HTMLToDocxParser(renderer)
    try:
        parser.feed(html_text)
        parser.close()
    except Exception:
        # Last-resort fallback: dump the text content into one paragraph.
        sd.add_paragraph(re.sub(r"<[^>]+>", "", html_text))
    renderer.finalize()
    return sd


# ============================================================
# Internal: stateful renderer + parser
# ============================================================

class _RunStyle:
    """Inline style accumulator while walking the HTML tree."""
    __slots__ = ("bold", "italic", "underline", "strike",
                 "monospace", "color", "highlight", "size_pt")

    def __init__(self) -> None:
        self.bold = False
        self.italic = False
        self.underline = False
        self.strike = False
        self.monospace = False
        self.color: Optional[str] = None
        self.highlight: Optional[str] = None
        self.size_pt: Optional[float] = None

    def clone(self) -> "_RunStyle":
        c = _RunStyle()
        c.bold = self.bold; c.italic = self.italic
        c.underline = self.underline; c.strike = self.strike
        c.monospace = self.monospace
        c.color = self.color; c.highlight = self.highlight
        c.size_pt = self.size_pt
        return c


def _parse_color(value: str) -> Optional[str]:
    """Accept #rrggbb / #rgb / rgb(r,g,b). Returns 6-hex string or None."""
    if not value:
        return None
    v = value.strip().lower()
    m = re.match(r"^#([0-9a-f]{6})$", v)
    if m: return m.group(1).upper()
    m = re.match(r"^#([0-9a-f]{3})$", v)
    if m:
        a, b, c = m.group(1)
        return (a + a + b + b + c + c).upper()
    m = re.match(r"^rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)$", v)
    if m:
        r, g, b = (int(x) & 0xFF for x in m.groups())
        return f"{r:02X}{g:02X}{b:02X}"
    return None


def _parse_style(decl: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for part in (decl or "").split(";"):
        if ":" not in part:
            continue
        k, v = part.split(":", 1)
        out[k.strip().lower()] = v.strip()
    return out


class _DocRenderer:
    """Drives docxtpl Subdoc paragraph/run creation."""

    def __init__(self, subdoc,
                 inline_images: list[str] | None = None,
                 fig_start: int = 0,
                 fig_prefix: str = "") -> None:
        self.subdoc = subdoc
        # Per-finding figure label prefix, e.g. "3.1" -> captions read
        # "Figure 3.1-<n>". Empty -> legacy plain "Figure <n>".
        self._fig_prefix = fig_prefix
        self._current_para = None
        self._style_stack: list[_RunStyle] = [_RunStyle()]
        # Stack of ("ul"|"ol", current_index)
        self._list_stack: list[tuple[str, int]] = []
        # Track if current paragraph is inside a heading (so finalize() knows)
        self._para_style: Optional[str] = None
        # 1-based list of file paths the consultant uploaded for the
        # surrounding finding. Used by `add_text` to resolve inline
        # `[Screenshot N]` tokens. Empty list disables the lookup but
        # leaves the literal token in place.
        self._inline_images: list[str] = list(inline_images or [])
        # Figure numbering for inline pasted images. 0 = disabled (no captions).
        self._fig_start = fig_start
        self._fig_count = 0   # images seen so far in this subdoc
        # When an image caption paragraph was just added, store it here so
        # that any text immediately following the <img> in the same HTML
        # paragraph gets appended to "Figure N:" rather than starting a new
        # paragraph.  Cleared by open_paragraph() or after one use.
        self._pending_caption_para = None

    # ----- paragraph lifecycle -----

    def open_paragraph(self, style: Optional[str] = None) -> None:
        # Close any open paragraph first.
        self.close_paragraph()
        # A new block element starts; any pending caption text that would have
        # belonged to the previous image's caption is no longer relevant.
        self._pending_caption_para = None
        try:
            self._current_para = self.subdoc.add_paragraph(style=style) if style \
                else self.subdoc.add_paragraph()
        except KeyError:
            # Style doesn't exist in this template; fall back.
            self._current_para = self.subdoc.add_paragraph()
        # Force left alignment so finding content never inherits the template's
        # "Normal" justified style — printed reports read better left-aligned.
        try:
            self._current_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass
        self._para_style = style

    def close_paragraph(self) -> None:
        self._current_para = None
        self._para_style = None

    def ensure_paragraph(self) -> None:
        if self._current_para is None:
            self.open_paragraph()

    def finalize(self) -> None:
        # Nothing special; placeholder for future trailing cleanup.
        pass

    # ----- run output -----

    def push_style(self, **overrides) -> None:
        new_style = self._style_stack[-1].clone()
        for k, v in overrides.items():
            setattr(new_style, k, v)
        self._style_stack.append(new_style)

    def pop_style(self) -> None:
        if len(self._style_stack) > 1:
            self._style_stack.pop()

    def add_text(self, text: str) -> None:
        if not text:
            return
        # Replace inline `[Screenshot N]` tokens with the matching
        # uploaded screenshot before the rest of the run-styling runs.
        # We split the text around every match and emit text + image
        # alternately. Tokens with no matching file (out-of-range N or
        # missing path) keep their literal text — that way the
        # consultant sees a broken reference rather than an invisible
        # drop, and can fix it on the next pass.
        if self._inline_images and _SCREENSHOT_REF_RE.search(text):
            last = 0
            for m in _SCREENSHOT_REF_RE.finditer(text):
                if m.start() > last:
                    self._emit_styled_text(text[last:m.start()])
                n = int(m.group("n"))
                path = (self._inline_images[n - 1]
                        if 1 <= n <= len(self._inline_images) else None)
                if path:
                    self._add_image_from_path(path)
                else:
                    # Out-of-range / missing — keep the literal token.
                    self._emit_styled_text(m.group(0))
                last = m.end()
            if last < len(text):
                self._emit_styled_text(text[last:])
            return
        self._emit_styled_text(text)

    def _emit_styled_text(self, text: str) -> None:
        """Internal: write `text` with the current style. Split out from
        `add_text` so the inline-screenshot resolver can call it on the
        text fragments around each token."""
        if not text:
            return
        # If a caption paragraph was just created for an inline image, discard
        # any text immediately following it — Word reports only show "Figure N",
        # no user-supplied caption text.
        if self._pending_caption_para is not None:
            self._pending_caption_para = None
            return
        self.ensure_paragraph()
        style = self._style_stack[-1]
        run = self._current_para.add_run(text)
        run.bold = style.bold
        run.italic = style.italic
        run.underline = style.underline
        if style.strike:
            run.font.strike = True
        if style.monospace:
            run.font.name = "Consolas"
        if style.color:
            try:
                run.font.color.rgb = RGBColor.from_string(style.color)
            except (ValueError, TypeError):
                pass
        if style.highlight:
            # We approximate highlight as cell shading via font color contrast;
            # docx supports w:highlight but only with a fixed palette. Skip
            # for now — Quill rarely uses it.
            pass
        if style.size_pt:
            try:
                from docx.shared import Pt
                run.font.size = Pt(style.size_pt)
            except Exception:
                pass

    def _add_image_from_path(self, path: str) -> None:
        """Embed an image from a filesystem path as a fresh paragraph.

        Used for inline `[Screenshot N]` token resolution. We open the
        file on disk (NOT a data URL like `add_image_data_url`) and
        emit it on its own paragraph so the picture isn't squashed
        between text runs in a tight list item. Width matches the
        140mm convention used elsewhere in the generator
        ([services/docx_generator.py](services/docx_generator.py)).
        """
        import os
        if not path or not os.path.exists(path):
            log.warning("[Screenshot N] target not found: %s", path)
            return
        # Close any open paragraph so the image gets its own line
        # rather than sitting inline next to surrounding text. Wrap in
        # try/except — if the subdoc rejects a fresh paragraph for some
        # reason we'd rather fall through and lose the image than crash
        # the whole render.
        try:
            self.close_paragraph()
            para = self.subdoc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(path, width=Cm(_INLINE_IMG_WIDTH_CM))
        except Exception as e:                              # pragma: no cover
            log.warning("Inline screenshot embed failed for %s: %s", path, e)
            return
        # Figure caption — same pattern as add_image_data_url
        self._fig_count += 1
        if self._fig_start:
            fig_num = self._fig_start + self._fig_count - 1
            _lbl = (f"{self._fig_prefix}-{fig_num}"
                    if self._fig_prefix else str(fig_num))
            self.close_paragraph()
            try:
                cap_para = self.subdoc.add_paragraph(f"Figure {_lbl}", style="Caption")
            except (KeyError, Exception):
                cap_para = self.subdoc.add_paragraph()
                cap_run = cap_para.add_run(f"Figure {_lbl}")
                cap_run.italic = True
            self._pending_caption_para = cap_para

    def add_break(self) -> None:
        if self._current_para is None:
            self.open_paragraph()
            return
        self._current_para.add_run().add_break()

    def open_pre_paragraph(self) -> None:
        """Open a new code-block paragraph with gray background and Consolas indentation."""
        self.close_paragraph()
        para = self.subdoc.add_paragraph()
        self._current_para = para
        self._para_style = 'pre'
        _apply_code_para_style(para)
        try:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    def add_pre_text(self, text: str) -> None:
        """Add text inside a <pre> block, splitting newlines into separate styled paragraphs."""
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                self.open_pre_paragraph()
            if self._current_para is None:
                self.open_pre_paragraph()
            style = self._style_stack[-1]
            run = self._current_para.add_run(line)
            run.font.name = "Consolas"
            run.italic = True
            if style.bold:
                run.bold = True

    def add_image_data_url(self, data_url: str, *,
                            width_mm: Optional[float] = None) -> None:
        """Decode a `data:image/...;base64,...` URL and embed it as an
        inline picture inside the current paragraph. Logs + skips on
        decode failure rather than raising — a broken pasted image
        shouldn't tank the whole render."""
        m = _DATA_IMG_RE.match((data_url or "").strip())
        if not m:
            log.debug("skipping non-image data URL in finding HTML")
            return
        try:
            payload = re.sub(r"\s+", "", m.group("b64"))
            blob = base64.b64decode(payload)
        except Exception as e:                          # pragma: no cover
            log.warning("Could not decode pasted image data URL: %s", e)
            return
        self.ensure_paragraph()
        if self._current_para is not None:
            self._current_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = self._current_para.add_run()
        try:
            # Fixed width (SCREENSHOT_WIDTH_CM), centred. Tall images that would
            # overflow the page height fall back to a height-bound size.
            _max_h = 230.0
            _w_arg = Cm(_INLINE_IMG_WIDTH_CM)
            _h_arg = None
            try:
                from PIL import Image as _PILImage
                with _PILImage.open(io.BytesIO(blob)) as _img:
                    _pw, _ph = _img.size
                _h_at_full = (_ph / _pw) * (_INLINE_IMG_WIDTH_CM * 10.0) if _pw else 0
                if _h_at_full and _h_at_full > _max_h:
                    _w_arg = None
                    _h_arg = Mm(_max_h)
            except Exception:
                pass
            if _h_arg:
                run.add_picture(io.BytesIO(blob), height=_h_arg)
            else:
                run.add_picture(io.BytesIO(blob), width=_w_arg)
        except Exception as e:                          # pragma: no cover
            log.warning("docx add_picture failed for inline screenshot: %s", e)
            return
        # Figure caption — only when the caller supplied a starting figure number.
        self._fig_count += 1
        if self._fig_start:
            fig_num = self._fig_start + self._fig_count - 1
            _lbl = (f"{self._fig_prefix}-{fig_num}"
                    if self._fig_prefix else str(fig_num))
            self.close_paragraph()
            try:
                cap_para = self.subdoc.add_paragraph(f"Figure {_lbl}", style="Caption")
            except (KeyError, Exception):
                cap_para = self.subdoc.add_paragraph()
                cap_run = cap_para.add_run(f"Figure {_lbl}")
                cap_run.italic = True
            self._pending_caption_para = cap_para

    # ----- list management -----

    def enter_list(self, kind: str) -> None:
        self._list_stack.append((kind, 0))

    def exit_list(self) -> None:
        if self._list_stack:
            self._list_stack.pop()

    def open_list_item(self) -> None:
        if not self._list_stack:
            self.open_paragraph()
            return
        kind, idx = self._list_stack[-1]
        if kind == "ol":
            idx += 1
            self._list_stack[-1] = (kind, idx)
            prefix = f"{idx}. "
        else:
            prefix = "• "  # bullet •
        # "List Paragraph" exists in every template; falls back to plain if not.
        # We add the numbering/bullet as a plain text prefix so it always renders,
        # regardless of whether the template has a numbering.xml definition.
        self.open_paragraph(style="List Paragraph")
        if self._current_para is not None:
            self._current_para.add_run(prefix)


class _HTMLToDocxParser(html.parser.HTMLParser):
    def __init__(self, renderer: _DocRenderer) -> None:
        super().__init__(convert_charrefs=True)
        self.r = renderer
        self._skip_text = 0           # >0 inside a tag we want to ignore
        self._open_blocks: list[str] = []
        self._anchor_href: str = ""
        self._anchor_text_parts: list[str] = []  # text seen inside current <a>

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        attrs_d = {k: (v or "") for k, v in attrs}
        # Style attribute -> inline overrides
        style = _parse_style(attrs_d.get("style", ""))

        if tag == "br":
            self.r.add_break()
            return

        # Inline screenshot. Quill pastes screenshots as
        # <img src="data:image/png;base64,…">. The HTML sanitizer
        # only lets through image data URLs (and http(s) URLs that we
        # ignore here — remote refs would require a fetch that we
        # don't want to do at render time).
        if tag == "img":
            src = attrs_d.get("src", "")
            width_attr = attrs_d.get("width", "") or ""
            width_mm: Optional[float] = None
            # If the editor recorded a pixel width, scale to a reasonable
            # mm width — ~3.78 px/mm so a 600px screenshot becomes
            # ~158mm. Capped so a tiny icon doesn't render half-page
            # and a huge paste doesn't overflow.
            try:
                if width_attr and width_attr.isdigit():
                    width_mm = max(40.0, min(170.0, int(width_attr) / 3.78))
            except Exception:
                width_mm = None
            self.r.add_image_data_url(src, width_mm=width_mm)
            # <img> is a void element — html.parser never emits an end
            # tag for it, so we deliberately do NOT push to
            # `_open_blocks` (would leave a phantom entry that corrupts
            # later balance checks).
            return

        if tag in BLOCK_TAGS:
            heading_style = None
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                heading_style = f"Heading {tag[1]}"
            elif tag == "blockquote":
                heading_style = "Intense Quote"
            elif tag == "pre":
                # treat pre as a gray-background Consolas code block
                self.r.push_style(monospace=True)
                self._open_blocks.append("pre")
                self.r.open_pre_paragraph()
                return
            elif tag == "li":
                self.r.open_list_item()
                self._open_blocks.append("li")
                return
            self.r.open_paragraph(style=heading_style)
            self._open_blocks.append(tag)
            return

        if tag in LIST_TAGS:
            self.r.enter_list(tag)
            self._open_blocks.append(tag)
            return

        if tag in INLINE_TAGS:
            overrides: dict = {}
            if tag in ("b", "strong"): overrides["bold"] = True
            if tag in ("i", "em"):     overrides["italic"] = True
            if tag == "u":             overrides["underline"] = True
            if tag in ("s", "strike"): overrides["strike"] = True
            if tag in ("code", "kbd", "samp"): overrides["monospace"] = True
            if tag == "a":
                overrides["underline"] = True
                overrides["color"] = "1F4E79"  # blue
            # Style-based overrides win over tag-default
            if style:
                if style.get("font-weight") in ("bold", "700", "800", "900"):
                    overrides["bold"] = True
                if style.get("font-style") == "italic":
                    overrides["italic"] = True
                td = style.get("text-decoration", "")
                if "underline" in td: overrides["underline"] = True
                if "line-through" in td: overrides["strike"] = True
                color = _parse_color(style.get("color", ""))
                if color: overrides["color"] = color
                bg = _parse_color(style.get("background-color", ""))
                if bg: overrides["highlight"] = bg
            self.r.push_style(**overrides)
            self._open_blocks.append(tag)
            # For <a>, we'll print the link text now and the URL after close.
            if tag == "a":
                self._anchor_href = attrs_d.get("href", "")
                self._anchor_text_parts = []  # reset text collector
            return

        if tag in ("script", "style", "iframe", "object", "embed"):
            # Should be impossible after sanitize() but defence-in-depth.
            self._skip_text += 1
            self._open_blocks.append("_skip")
            return

        # Unknown tag - just don't do anything; text inside still flows
        self._open_blocks.append("_passthrough")

    def handle_startendtag(self, tag: str, attrs) -> None:
        # Self-closing form (<img/>, <br/>) — funnel through the regular
        # start handler. `<img>` short-circuits before pushing to the
        # block stack, so there's nothing to balance afterwards.
        self.handle_starttag(tag, attrs)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if not self._open_blocks:
            return
        top = self._open_blocks.pop()

        if top == "_skip":
            self._skip_text = max(0, self._skip_text - 1)
            return
        if top == "_passthrough":
            return
        if top == "pre":
            self.r.pop_style()
            self.r.close_paragraph()
            return
        if top == "li":
            self.r.close_paragraph()
            return
        if top in LIST_TAGS:
            self.r.exit_list()
            return
        if top in BLOCK_TAGS:
            self.r.close_paragraph()
            return
        if top in INLINE_TAGS:
            if top == "a" and self._anchor_href:
                href = self._anchor_href
                link_text = "".join(self._anchor_text_parts)
                self.r.pop_style()
                # Only append the URL when it isn't already visible in the link
                # text — avoids duplicates when the user types the URL as the
                # display text (the common reference pattern in VAPT reports).
                if href not in link_text:
                    self.r.add_text(f" ({href})")
                self._anchor_href = ""
                self._anchor_text_parts = []
                return
            self.r.pop_style()

    def handle_data(self, data: str) -> None:
        if self._skip_text:
            return
        # Collect text visible inside the current <a> so we can skip the
        # parenthetical URL if the href already appears in the link text.
        if self._anchor_href and "a" in self._open_blocks:
            self._anchor_text_parts.append(data)
        # Collapse runs of whitespace (Word handles its own spacing); preserve
        # single spaces between words and inside <pre> blocks.
        if "pre" in self._open_blocks:
            # Preserve whitespace and split newlines into separate code paragraphs
            self.r.add_pre_text(data)
        else:
            collapsed = re.sub(r"[ \t\r\n\f\v]+", " ", data)
            if collapsed:
                self.r.add_text(collapsed)
